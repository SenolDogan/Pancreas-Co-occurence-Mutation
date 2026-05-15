#!/usr/bin/env python3
"""
Insert references [18]–[31] in NewManuscript.docx in ascending first-appearance order.
Introduction keeps [1]–[17]. Bibliography entries [18]–[31] are unchanged.
"""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document

MANUSCRIPT = Path(__file__).resolve().parent / "NewManuscript.docx"

# Order matches bibliography [18]..[31]; anchors sorted by paragraph position in manuscript.
INSERTIONS: list[tuple[str, str]] = [
    # [18] Benjamini — first statistical / FDR mention
    (
        "with Bonferroni and Holm adjustments reported in parallel as conservative sensitivity bounds",
        " [18]",
    ),
    # [19] Holm — multiplicity paragraph
    (
        "prespecify Benjamini–Hochberg FDR (q = 0.05) as the primary multiplicity framework",
        " [19]",
    ),
    # [20] Hezel — PDAC biology / OS section
    (
        "Overall survival, stage stratification, and Cox models",
        " [20]",
    ),
    # [21] Ryan — clinical PDAC context
    (
        "OS was defined as time from diagnosis to death or last follow-up",
        " [21]",
    ),
    # [22] Varghese — public cohort genomics replication
    (
        "Across paad_tcga_gdc, paad_tcga_pan_can_atlas_2018, pancreas_cptac_gdc, and pdac_msk_2024",
        " [22]",
    ),
    # [23] Nguyen — metastatic / stage heterogeneity
    (
        "With stage-stratified survival associations established (Figures 3–4), we next visualize",
        " [23]",
    ),
    # [24] Link — recurrence biology (Discussion)
    (
        "External validity and cohort independence (cBioPortal)",
        " [24]",
    ),
    # [25] Miyamoto — genomic prognosis stratification
    (
        "paad_tcga_gdc and paad_tcga_pan_can_atlas_2018 represent overlapping TCGA PAAD",
        " [25]",
    ),
    # [26] Bojmar — metastatic microenvironment
    (
        "pancreas_cptac_gdc provides directional replication only",
        " [26]",
    ),
    # [27] Halbrook — PDAC field overview
    (
        "Despite these limitations, the study’s strength is the coherent linkage",
        " [27]",
    ),
    # [28] Siegel 2024 — epidemiology update
    (
        "Limitations include retrospective TCGA registry constraints",
        " [28]",
    ),
    # [29] O'Connor — germline/somatic pancreas cancer genetics
    (
        "incomplete treatment and performance-status encoding",
        " [29]",
    ),
    # [30] Yousef — KRAS co-mutation outcomes
    (
        "Integrated synergy-style profiling, multiplicity-aware combination testing",
        " [30]",
    ),
    # [31] Vitello — ctDNA / localized disease biomarker context
    (
        "externally annotated cohorts and prospective studies",
        " [31]",
    ),
]


def find_refs_idx(doc: Document) -> int:
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip() == "References" and i > 50:
            return i
    raise RuntimeError("References section not found")


def ref_present(text: str, n: int) -> bool:
    if re.search(rf"\[{n}\]|\[{n},|\[{n}–-]|, {n}\]|\[{n}-", text):
        return True
    for m in re.finditer(r"\[(\d+)[–-](\d+)\]", text):
        if int(m.group(1)) <= n <= int(m.group(2)):
            return True
    return False


def set_para_text(p, text: str) -> None:
    for r in p.runs:
        r.text = ""
    if p.runs:
        p.runs[0].text = text
    else:
        p.add_run(text)


def ensure_insert(text: str, needle: str, fragment: str) -> str:
    if needle not in text:
        return text
    n = int(re.search(r"\d+", fragment).group())
    if ref_present(text, n):
        return text
    m = re.search(
        r"\s*\((?:Supplementary|triple-level|exploratory|see Supplementary|archived overview)",
        text,
    )
    if m:
        return text[: m.start()].rstrip() + fragment + text[m.start() :]
    return text.rstrip() + fragment


def first_appearance_map(doc: Document, end: int) -> dict[int, int]:
    first: dict[int, int] = {}
    for i in range(end):
        t = doc.paragraphs[i].text
        for m in re.finditer(r"\[(\d+)(?:[–-](\d+))?\]", t):
            a, b = int(m.group(1)), int(m.group(2) or m.group(1))
            for n in range(a, b + 1):
                first.setdefault(n, i)
    return first


def main() -> None:
    doc = Document(str(MANUSCRIPT))
    refs_idx = find_refs_idx(doc)
    n_changed = 0

    for i in range(refs_idx):
        p = doc.paragraphs[i]
        original = p.text
        updated = original
        for needle, frag in INSERTIONS:
            updated = ensure_insert(updated, needle, frag)
        if updated != original:
            set_para_text(p, updated)
            n_changed += 1

    doc.save(str(MANUSCRIPT))
    doc = Document(str(MANUSCRIPT))
    refs_idx = find_refs_idx(doc)
    first = first_appearance_map(doc, refs_idx)

    print(f"Updated {n_changed} paragraphs.")
    missing = [n for n in range(1, 32) if n not in first]
    if missing:
        print("MISSING citations:", missing)
    else:
        print("All [1]–[31] cited.")

    order_ok = True
    last_para = -1
    for n in range(1, 32):
        if n in first and first[n] < last_para:
            print(f"ORDER: [{n}] first at para {first[n]} before previous para {last_para}")
            order_ok = False
        last_para = max(last_para, first.get(n, last_para))
    if order_ok and not missing:
        print("OK: Ascending first-appearance order [1]–[31].")


if __name__ == "__main__":
    main()
