#!/usr/bin/env python3
"""
Renumber supplementary items 1–32 by first citation order in NewManuscript.docx.
In-text citations, end inventory, and supplementary files use the same global numbers.
"""

from __future__ import annotations

import re
from copy import deepcopy
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph

TUMOR = Path(__file__).resolve().parent
MANUSCRIPT = TUMOR / "NewManuscript.docx"
SUPP_DIR = TUMOR / "supplementary"

METHOD_TITLES = {
    1: "Scope and relationship to the main manuscript",
    2: "Primary analytic cohort (discovery)",
    3: "Synergy and enrichment metrics (multiplicative, additive, and protective scores)",
    4: "Stage-stratified Cox and multivariable extensions",
    5: "DX2COLLECTION_YEAR × mutation-combination interaction models and sensitivity filters",
    6: "Internal validation (stratified split and cross-validation)",
    7: "External validation in public PDAC cohorts (cBioPortal)",
    8: "Overview of supplementary figure content",
    9: "Software and computational environment",
}

FIGURE_TITLES = {
    1: "Triple-mutation additive deviations from independence (top positive and negative bars by stage)",
    2: "Univariate Cox forest plot for top triple-mutation combinations by stage",
    3: "Diagnosis-to-collection timing × pairwise combination interaction volcano (DX2COLLECTION_YEAR ≥ 0)",
    4: "Diagnosis-to-collection timing × triple-mutation interaction volcano (DX2COLLECTION_YEAR ≥ 0)",
    5: "TP53/KRAS patterns in short overall-survival subsets (exploratory sensitivity)",
    6: "Diabetes-stratified exploratory gene analysis",
    7: "Triple-mutation timing interaction volcano under DX2COLLECTION_YEAR 0–5 years sensitivity filter",
    8: "Multiple-comparison correction diagnostics for combination screens",
    9: "Internal discovery–validation cohort diagnostics",
    10: "Pathway-oriented functional validation summary",
    11: "Deceased-versus-living comprehensive mutation enrichment overview",
    12: "Multivariable Cox internal validation summary by stage",
}

TABLE_TITLES = {
    1: "Timing-interaction sensitivity comparison for DX2COLLECTION_YEAR × mutation combinations (adjusted Cox; ALL, DX≥0, and 0–5 year filters)",
    2: "External replication of TP53+KRAS co-mutation prevalence and univariate overall survival in public PDAC cohorts (TCGA, Pan-Cancer Atlas, CPTAC, MSK, and discovery cohort)",
    3: "Stage-stratified mutation frequencies, pairwise and triple additive co-occurrence, and univariate Cox overall survival models",
    4: "Diagnosis-to-collection timing × combination interaction Cox models (primary cohort)",
    5: "Timing-interaction Cox outputs under dx_yr sensitivity filters (ALL, DX≥0, DX 0–5 years)",
    6: "Multivariable Cox coefficients and internal validation metrics by stage",
    7: "Three-group stage landscape, mutation rates, clinical associations, and synergy summaries",
    8: "Top-ranked synergistic dual-mutation combinations by stage",
    9: "Internal discovery–validation cohort comparison and cross-validation metrics",
    10: "Dual-mutation combination significance, FDR/Bonferroni summaries, and multiplicity correction",
    11: "Deceased-versus-living enrichment: cohort summary, mutation frequencies, synergy, and clinical factors",
}

UNCITED_ORDER = [
    ("method", 1),
    ("method", 4),
    ("figure", 5),
    ("figure", 6),
    ("table", 4),
    ("table", 6),
    ("table", 9),
]


def item_label(typ: str, num: int, title: str) -> str:
    kind = {"method": "Method", "figure": "Figure", "table": "Table"}[typ]
    return f"Supplementary {kind} {num}. {title}"


def title_for_key(key: tuple[str, int]) -> str:
    typ, old = key
    if typ == "method":
        return METHOD_TITLES[old]
    if typ == "figure":
        return FIGURE_TITLES[old]
    return TABLE_TITLES[old]


def delete_paragraph(paragraph) -> None:
    el = paragraph._element
    el.getparent().remove(el)


def set_para_text(p, text: str) -> None:
    for r in p.runs:
        r.text = ""
    if p.runs:
        p.runs[0].text = text
    else:
        p.add_run(text)


def insert_paragraph_after(paragraph, text: str) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._element.addnext(new_p)
    np = Paragraph(new_p, paragraph._parent)
    np.add_run(text)
    return np


def normalize_s_prefixes(text: str) -> str:
    if not text.strip():
        return text
    placeholders: dict[str, str] = {}

    def protect(m: re.Match) -> str:
        key = f"__MAIN_{len(placeholders)}__"
        placeholders[key] = m.group(0)
        return key

    text = re.sub(r"\b(?:Table|Figure) ([1-7])\.", protect, text)
    text = re.sub(r"Supplementary Tables S(\d+)", r"Supplementary Tables \1", text)
    text = re.sub(r"Supplementary Table S(\d+)", r"Supplementary Table \1", text)
    text = re.sub(r"Supplementary Figures S(\d+)", r"Supplementary Figures \1", text)
    text = re.sub(r"Supplementary Figure S(\d+)", r"Supplementary Figure \1", text)
    text = re.sub(r"Figure S(\d+)", r"Supplementary Figure \1", text)
    text = re.sub(r"Table S(\d+)", r"Supplementary Table \1", text)
    text = re.sub(r"(Supplementary (?:Figures|Tables) \d+)–S(\d+)", r"\1–\2", text)
    text = re.sub(r"(Supplementary (?:Figures|Tables) [\d–-]+) and S(\d+)", r"\1 and \2", text)
    for key, val in placeholders.items():
        text = text.replace(key, val)
    return text


def extract_refs_in_order(text: str) -> list[tuple[str, int]]:
    hits: list[tuple[int, str, int]] = []

    def add(typ: str, n: str, pos: int) -> None:
        hits.append((pos, typ, int(n)))

    for m in re.finditer(r"Supplementary Method (\d+)", text):
        add("method", m.group(1), m.start())
    for m in re.finditer(r"Supplementary Figure (\d+)", text):
        add("figure", m.group(1), m.start())
    for m in re.finditer(r"Supplementary Table (\d+)", text):
        add("table", m.group(1), m.start())
    for m in re.finditer(r"Supplementary Figures (\d+)[–-](\d+)", text):
        add("figure", m.group(1), m.start())
        add("figure", m.group(2), m.start() + 1)
    for m in re.finditer(r"Supplementary Figures [\d–-]+ and (\d+)", text):
        add("figure", m.group(1), m.start())
    for m in re.finditer(r"Supplementary Tables (\d+)[–-](\d+)", text):
        add("table", m.group(1), m.start())
        add("table", m.group(2), m.start() + 1)
    for m in re.finditer(r"Supplementary Tables (\d+), (\d+), and (\d+)", text):
        add("table", m.group(1), m.start())
        add("table", m.group(2), m.start() + 1)
        add("table", m.group(3), m.start() + 2)
    for m in re.finditer(r"Supplementary Tables (\d+) and (\d+)", text):
        add("table", m.group(1), m.start())
        add("table", m.group(2), m.start() + 1)
    for m in re.finditer(r"Supplementary Tables (\d+), (\d+)", text):
        add("table", m.group(1), m.start())
        add("table", m.group(2), m.start() + 1)
    for m in re.finditer(r"Supplementary Tables (\d+)-(\d+)", text):
        lo, hi = int(m.group(1)), int(m.group(2))
        for n in range(lo, hi + 1):
            add("table", str(n), m.start())

    hits.sort(key=lambda x: x[0])
    order: list[tuple[str, int]] = []
    seen: set[tuple[str, int]] = set()
    for _, typ, n in hits:
        key = (typ, n)
        if key not in seen:
            seen.add(key)
            order.append(key)
    return order


def build_global_order(doc: Document, body_end: int) -> list[tuple[str, int]]:
    order: list[tuple[str, int]] = []
    seen: set[tuple[str, int]] = set()
    for i in range(body_end):
        text = normalize_s_prefixes(doc.paragraphs[i].text)
        for key in extract_refs_in_order(text):
            if key not in seen:
                seen.add(key)
                order.append(key)
    for key in UNCITED_ORDER:
        if key not in seen:
            seen.add(key)
            order.append(key)
    return order


def build_mapping(order: list[tuple[str, int]]) -> dict[tuple[str, int], int]:
    return {key: idx + 1 for idx, key in enumerate(order)}


def token(typ: str, old: int) -> str:
    return f"__SUP_{typ.upper()}_{old:02d}__"


def text_to_tokens(text: str) -> str:
    text = normalize_s_prefixes(text)

    text = re.sub(
        r"Supplementary Figures (\d+)[–-](\d+) and (\d+)",
        lambda m: (
            f"Supplementary Figures {token('figure', int(m.group(1)))}–"
            f"{token('figure', int(m.group(2)))} and {token('figure', int(m.group(3)))}"
        ),
        text,
    )
    text = re.sub(
        r"Supplementary Figures (\d+)[–-](\d+)",
        lambda m: f"{token('figure', int(m.group(1)))}–{token('figure', int(m.group(2)))}",
        text,
    )
    text = re.sub(
        r"Supplementary Tables (\d+), (\d+), and (\d+)",
        lambda m: (
            f"{token('table', int(m.group(1)))}, {token('table', int(m.group(2)))}, "
            f"and {token('table', int(m.group(3)))}"
        ),
        text,
    )
    text = re.sub(
        r"Supplementary Tables (\d+) and (\d+)",
        lambda m: f"{token('table', int(m.group(1)))} and {token('table', int(m.group(2)))}",
        text,
    )
    text = re.sub(
        r"Supplementary Tables (\d+), (\d+)",
        lambda m: f"{token('table', int(m.group(1)))}, {token('table', int(m.group(2)))}",
        text,
    )
    text = re.sub(
        r"Supplementary Tables (\d+)[–-](\d+)",
        lambda m: f"{token('table', int(m.group(1)))}–{token('table', int(m.group(2)))}",
        text,
    )

    def repl_table_range(m: re.Match) -> str:
        lo, hi = int(m.group(1)), int(m.group(2))
        return f"{token('table', lo)}–{token('table', hi)}"

    text = re.sub(r"Supplementary Tables (\d+)-(\d+)", repl_table_range, text)
    text = re.sub(r"Supplementary Method (\d+)", lambda m: token("method", int(m.group(1))), text)
    text = re.sub(r"Supplementary Figure (\d+)", lambda m: token("figure", int(m.group(1))), text)
    text = re.sub(r"Supplementary Table (\d+)", lambda m: token("table", int(m.group(1))), text)
    return text


def tokens_to_new(text: str, mapping: dict[tuple[str, int], int]) -> str:
    def tok_to_label(m: re.Match) -> str:
        typ = m.group(1).lower()
        old = int(m.group(2))
        new = mapping[(typ, old)]
        kind = {"method": "Method", "figure": "Figure", "table": "Table"}[typ]
        return f"Supplementary {kind} {new}"

    text = re.sub(r"__SUP_(METHOD|FIGURE|TABLE)_(\d+)__", tok_to_label, text)

    # Collapse adjacent table/figure spans to "Supplementary Tables 6–14" style
    def collapse_range(m: re.Match) -> str:
        kind = m.group(1)
        plural = "Tables" if kind == "Table" else "Figures"
        return f"Supplementary {plural} {m.group(2)}–{m.group(3)}"

    text = re.sub(
        r"Supplementary Table (\d+)[–-]Supplementary Table (\d+)",
        collapse_range,
        text,
    )
    text = re.sub(
        r"Supplementary Figure (\d+)[–-]Supplementary Figure (\d+)",
        lambda m: f"Supplementary Figures {m.group(1)}–{m.group(2)}",
        text,
    )
    text = re.sub(
        r"Supplementary Tables? (\d+), Supplementary Tables? (\d+), and Supplementary Tables? (\d+)",
        lambda m: f"Supplementary Tables {m.group(1)}, {m.group(2)}, and {m.group(3)}",
        text,
    )
    text = re.sub(
        r"Supplementary Tables? (\d+) and Supplementary Tables? (\d+)",
        lambda m: f"Supplementary Tables {m.group(1)} and {m.group(2)}",
        text,
    )
    text = re.sub(
        r"Supplementary Tables? (\d+), Supplementary Tables? (\d+)",
        lambda m: f"Supplementary Tables {m.group(1)} and {m.group(2)}",
        text,
    )
    text = re.sub(
        r"Supplementary Figures? (\d+)[–-]Supplementary Figures? (\d+) and Supplementary Figures? (\d+)",
        lambda m: f"Supplementary Figures {m.group(1)}–{m.group(2)} and {m.group(3)}",
        text,
    )
    return text


def apply_renumber_text(text: str, mapping: dict[tuple[str, int], int]) -> str:
    return tokens_to_new(text_to_tokens(text), mapping)


def find_section_start(doc: Document, heading: str) -> int:
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip() == heading:
            return i
    raise RuntimeError(f"Section not found: {heading}")


def rebuild_manuscript_inventory(doc: Document, intro_idx: int, order: list[tuple[str, int]], mapping: dict) -> None:
    ack_idx = find_section_start(doc, "Acknowledgments")
    while intro_idx + 1 < ack_idx:
        delete_paragraph(doc.paragraphs[intro_idx + 1])
        ack_idx = find_section_start(doc, "Acknowledgments")

    intro_p = doc.paragraphs[intro_idx]
    set_para_text(
        intro_p,
        "The following supplementary items accompany the main manuscript and are numbered sequentially "
        "(items 1–32) in order of first citation in the main text. Extended content is provided in separate files "
        "(Supplementary Methods.docx; Supplementary Figures.docx; Supplementary Tables.xlsx).",
    )

    lines = [f"{g}. {item_label(k[0], mapping[k], title_for_key(k))}" for g, k in enumerate(order, 1)]
    anchor = intro_p
    for line in lines:
        anchor = insert_paragraph_after(anchor, line)


def update_manuscript(mapping: dict[tuple[str, int], int], order: list[tuple[str, int]]) -> None:
    doc = Document(str(MANUSCRIPT))
    supp_start = find_section_start(doc, "Supplementary Materials")
    ack_idx = find_section_start(doc, "Acknowledgments")
    intro_idx = supp_start
    for i in range(supp_start, ack_idx):
        if "supplementary items accompany" in doc.paragraphs[i].text.lower():
            intro_idx = i
            break

    n_changed = 0
    for i, p in enumerate(doc.paragraphs):
        if i >= ack_idx:
            break
        if supp_start <= i < ack_idx and i != intro_idx:
            continue
        new = apply_renumber_text(p.text, mapping)
        if new != p.text:
            set_para_text(p, new)
            n_changed += 1

    rebuild_manuscript_inventory(doc, intro_idx, order, mapping)
    doc.save(str(MANUSCRIPT))
    print(f"Manuscript: {n_changed} body paragraphs updated; inventory rebuilt (1–32).")


def update_supplementary_methods(mapping: dict[tuple[str, int], int]) -> None:
    """Rebuild Methods.docx from generator source, then apply global numbers."""
    import tempfile

    from build_supplementary_materials import build_supplementary_methods

    path = SUPP_DIR / "Supplementary Methods.docx"
    tmp = Path(tempfile.gettempdir()) / "supp_methods_base.docx"
    build_supplementary_methods(TUMOR, tmp)

    doc = Document(str(tmp))
    section_re = re.compile(r"^(\d+)\.\s+(.+)$")
    sections: list[tuple[int, str, list[str]]] = []
    current_old: int | None = None
    current_title = ""
    current_body: list[str] = []

    for p in doc.paragraphs:
        m = section_re.match(p.text.strip())
        if m and p.style and "Heading" in (p.style.name or ""):
            if current_old is not None:
                sections.append((current_old, current_title, current_body))
            current_old = int(m.group(1))
            current_title = m.group(2)
            current_body = []
        elif current_old is not None and p.text.strip():
            current_body.append(p.text)
    if current_old is not None:
        sections.append((current_old, current_title, current_body))

    doc2 = Document()
    title = doc2.add_paragraph("Supplementary Methods")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.bold = True
    doc2.add_paragraph()

    sections.sort(key=lambda x: mapping[("method", x[0])])
    for old, _title, body in sections:
        new = mapping[("method", old)]
        h = doc2.add_paragraph(f"Supplementary Method {new}. {METHOD_TITLES[old]}")
        for run in h.runs:
            run.bold = True
        for line in body:
            doc2.add_paragraph(apply_renumber_text(line, mapping))

    doc2.save(str(path))
    tmp.unlink(missing_ok=True)
    print(f"Updated {path.name}")


def update_supplementary_figures(mapping: dict[tuple[str, int], int]) -> None:
    path = SUPP_DIR / "Supplementary Figures.docx"
    doc = Document(str(path))
    fig_heading_re = re.compile(r"^Supplementary Figure S?(\d+)\.")

    body = doc.element.body
    children = list(body)
    title_children: list = []
    blocks: list[tuple[int, int, list]] = []  # global_num, old_num, elements
    current_block: list = []
    current_old: int | None = None

    for child in children:
        tag = child.tag.split("}")[-1]
        if tag == "p":
            p = Paragraph(child, doc)
            m = fig_heading_re.match(p.text.strip())
            if m:
                if current_old is not None:
                    blocks.append((mapping[("figure", current_old)], current_old, current_block))
                current_old = int(m.group(1))
                current_block = [child]
            elif current_block:
                current_block.append(child)
            else:
                title_children.append(child)
        elif current_block:
            current_block.append(child)
        else:
            title_children.append(child)

    if current_old is not None:
        blocks.append((mapping[("figure", current_old)], current_old, current_block))

    for child in list(body):
        body.remove(child)

    for child in title_children:
        body.append(child)

    for global_num, old, block in sorted(blocks, key=lambda x: x[0]):
        # Update heading text in first paragraph of block
        first = Paragraph(block[0], doc)
        set_para_text(first, f"Supplementary Figure {global_num}. {FIGURE_TITLES[old]}")
        for child in block:
            body.append(child)

    doc.save(str(path))
    print(f"Updated {path.name}")


def update_supplementary_tables(mapping: dict[tuple[str, int], int]) -> None:
    path = SUPP_DIR / "Supplementary Tables.xlsx"
    xl = pd.ExcelFile(path)
    rows = []
    for old in range(1, 12):
        sheets = [s for s in xl.sheet_names if re.match(rf"^S{old}_", s)]
        rows.append(
            {
                "Global_Number": mapping[("table", old)],
                "Supplementary_Table": f"Supplementary Table {mapping[('table', old)]}",
                "Legacy_Sheet_Prefix": f"S{old}_",
                "Sheets": ", ".join(sheets),
                "Description": TABLE_TITLES[old],
            }
        )
    idx_df = pd.DataFrame(rows).sort_values("Global_Number")

    # Rewrite workbook preserving sheets, replace Table_Index only
    from openpyxl import load_workbook

    wb = load_workbook(path)
    if "Table_Index" in wb.sheetnames:
        del wb["Table_Index"]
    ws = wb.create_sheet("Table_Index", 0)
    for r_idx, row in enumerate([idx_df.columns.tolist()] + idx_df.values.tolist(), 1):
        for c_idx, val in enumerate(row, 1):
            ws.cell(row=r_idx, column=c_idx, value=val)
    wb.save(path)
    print(f"Updated Table_Index in {path.name}")


def canonical_mapping() -> tuple[list[tuple[str, int]], dict[tuple[str, int], int]]:
    """Fixed order from first full pass (do not re-derive from already-renumbered text)."""
    order = [
        ("method", 6),
        ("method", 7),
        ("table", 2),
        ("method", 8),
        ("figure", 10),
        ("table", 1),
        ("table", 3),
        ("table", 4),
        ("table", 5),
        ("table", 6),
        ("table", 7),
        ("table", 8),
        ("table", 9),
        ("table", 10),
        ("method", 9),
        ("method", 5),
        ("figure", 3),
        ("figure", 7),
        ("figure", 4),
        ("figure", 8),
        ("method", 2),
        ("figure", 11),
        ("table", 11),
        ("method", 3),
        ("figure", 12),
        ("figure", 1),
        ("figure", 2),
        ("figure", 9),
        ("method", 1),
        ("method", 4),
        ("figure", 5),
        ("figure", 6),
    ]
    return order, build_mapping(order)


def polish_supplementary_citation_text(text: str) -> str:
    """Fix broken ranges and lingering legacy S-prefix / old table numbers in body text."""
    text = normalize_s_prefixes(text)
    text = re.sub(
        r"Supplementary Methods\s*§\s*(\d+)",
        lambda m: f"Supplementary Method {canonical_mapping()[1][('method', int(m.group(1)))]}",
        text,
    )

    # Duplicate 'Supplementary Figures Supplementary Figure' artifacts
    text = re.sub(r"Supplementary Figures Supplementary Figure ", "Supplementary Figure ", text)
    text = re.sub(
        r"Supplementary Figure (\d+)[–-]Supplementary Figure (\d+)",
        r"Supplementary Figures \1–\2",
        text,
    )
    text = re.sub(
        r"Supplementary Figures Supplementary Figure (\d+)[–-]Supplementary Figure (\d+) and Supplementary Figure (\d+)",
        r"Supplementary Figures \1–\2 and \3",
        text,
    )

    # Targeted legacy S-table / old-number fixes (from inject_supplementary_citations intent)
    replacements = [
        ("Supplementary Tables 3, S7, and S11", "Supplementary Tables 7, 11, and 23"),
        ("Supplementary Tables 5;", "Supplementary Table 9;"),
        ("Supplementary Tables 5)", "Supplementary Table 9)"),
        ("Supplementary Tables 10)", "Supplementary Table 14)"),
        ("Supplementary Tables 10.", "Supplementary Table 14."),
        ("Supplementary Tables 14–23", "Supplementary Tables 14 and 23"),
        ("Supplementary Tables 6;", "Supplementary Table 10;"),
        ("Supplementary Tables 6)", "Supplementary Table 10)"),
        ("Supplementary Tables 9)", "Supplementary Table 13)"),
        ("Supplementary Figures 17–Supplementary Figure 19 and Supplementary Figure 18", "Supplementary Figures 17–19 and 18"),
        ("Supplementary Figure 26–Supplementary Figure 27", "Supplementary Figures 26–27"),
    ]
    for old, new in replacements:
        text = text.replace(old, new)
    return text


def fix_manuscript_citation_polish() -> None:
    doc = Document(str(MANUSCRIPT))
    ack_idx = find_section_start(doc, "Acknowledgments")
    supp_start = find_section_start(doc, "Supplementary Materials")
    n = 0
    for i, p in enumerate(doc.paragraphs):
        if i >= ack_idx or (supp_start <= i < ack_idx):
            continue
        new = polish_supplementary_citation_text(p.text)
        if new != p.text:
            set_para_text(p, new)
            n += 1
    doc.save(str(MANUSCRIPT))
    print(f"Polished {n} manuscript paragraphs (citation cleanup).")


def fix_manuscript_ranges_and_inventory() -> None:
    order, mapping = canonical_mapping()
    doc = Document(str(MANUSCRIPT))
    ack_idx = find_section_start(doc, "Acknowledgments")
    supp_start = find_section_start(doc, "Supplementary Materials")
    intro_idx = supp_start
    for i in range(supp_start, ack_idx):
        if "supplementary items accompany" in doc.paragraphs[i].text.lower():
            intro_idx = i
            break

    for i, p in enumerate(doc.paragraphs):
        if i >= ack_idx:
            break
        if supp_start <= i < ack_idx and i != intro_idx:
            continue
        t = p.text
        t = re.sub(
            r"Supplementary Table (\d+)[–-]Supplementary Table (\d+)",
            lambda m: f"Supplementary Tables {m.group(1)}–{m.group(2)}",
            t,
        )
        if t != p.text:
            set_para_text(p, t)

    rebuild_manuscript_inventory(doc, intro_idx, order, mapping)
    doc.save(str(MANUSCRIPT))
    print("Fixed table ranges and rebuilt inventory list.")


def main() -> None:
    order, mapping = canonical_mapping()

    print("Global numbering (list position -> Supplementary Type N):")
    for pos, key in enumerate(order, 1):
        typ, old = key
        print(f"  List {pos}: {typ} (was {typ} {old}) -> Supplementary {typ.title()} {mapping[key]}")

    update_manuscript(mapping, order)
    update_supplementary_methods(mapping)
    update_supplementary_figures(mapping)
    update_supplementary_tables(mapping)


if __name__ == "__main__":
    import sys

    if len(sys.argv) > 1 and sys.argv[1] == "polish":
        fix_manuscript_citation_polish()
        fix_manuscript_ranges_and_inventory()
    elif len(sys.argv) > 1 and sys.argv[1] == "fix-inventory":
        fix_manuscript_ranges_and_inventory()
    else:
        main()
