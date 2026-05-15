#!/usr/bin/env python3
"""
Re-embed main-text figures in NewManuscript.docx: caption → image → legend.

Figures were removed when citation scripts called set_para_text() on image paragraphs.
This script splits merged caption+legend blocks, then inserts PNGs above each legend.
"""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches

TUMOR = Path(__file__).resolve().parent
NEW2 = TUMOR.parent
MANUSCRIPT = TUMOR / "NewManuscript.docx"

FIGURE_PNG: dict[int, Path] = {
    1: NEW2 / "01_Dead_Alive_Comprehensive_Analysis.png",
    2: TUMOR / "Stage_3Group_Comprehensive_Analysis.png",
    3: TUMOR / "Stage_Figure2_AB_ForestPlots_VSTACK.png",
    4: TUMOR / "Stage_06_OS_Months_Methodology_Analysis.png",
    5: TUMOR / "Stage_05_TP53_KRAS_Detailed_Analysis.png",
    6: TUMOR / "Stage_02_TP53_KRAS_Focused_Analysis.png",
    7: TUMOR / "Stage_Figure6_AB_AdditiveBars_CoxPairsForest.png",
}

FIGURE_WIDTH_IN: dict[int, float] = {7: 6.2}
DEFAULT_WIDTH = 6.5


def find_supp_start(doc: Document) -> int:
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip() == "Supplementary Materials" and i > 50:
            return i
    return len(doc.paragraphs)


def para_has_image(p) -> bool:
    return "pic:pic" in p._element.xml or "w:drawing" in p._element.xml


def set_para_text(p, text: str) -> None:
    for r in p.runs:
        r.text = ""
    if p.runs:
        p.runs[0].text = text
    else:
        p.add_run(text)


def split_caption_legend(text: str, fig_num: int) -> tuple[str, str] | None:
    pat = rf"^(Figure\s+{fig_num}\.\s*[^.]+\.)\s+(Panel\s+[A-Z].*)$"
    m = re.match(pat, text.strip(), re.IGNORECASE | re.DOTALL)
    if m:
        return m.group(1).strip(), m.group(2).strip()
    return None


def remove_image_only_paragraphs(doc: Document, end: int) -> int:
    removed = 0
    for i in range(end - 1, -1, -1):
        p = doc.paragraphs[i]
        if para_has_image(p) and not p.text.strip():
            p._element.getparent().remove(p._element)
            removed += 1
    return removed


def split_all_merged_captions(doc: Document) -> None:
    """Split Figure N. Title. Panel A... into caption + legend paragraphs."""
    end = find_supp_start(doc)
    to_split: list[tuple[int, str, str]] = []
    for i in range(end):
        t = doc.paragraphs[i].text.strip()
        m = re.match(r"^Figure\s+(\d+)\.", t, re.I)
        if not m:
            continue
        n = int(m.group(1))
        if n not in FIGURE_PNG:
            continue
        sp = split_caption_legend(t, n)
        if not sp:
            continue
        nxt = i + 1
        while nxt < end and not doc.paragraphs[nxt].text.strip():
            nxt += 1
        if nxt < end and doc.paragraphs[nxt].text.strip().startswith("Panel "):
            continue  # already split
        to_split.append((i, sp[0], sp[1]))

    for cap_i, cap_text, leg_text in sorted(to_split, key=lambda x: -x[0]):
        p = doc.paragraphs[cap_i]
        set_para_text(p, cap_text)
        anchor = cap_i + 1
        if anchor < len(doc.paragraphs):
            doc.paragraphs[anchor].insert_paragraph_before(leg_text)
        else:
            doc.add_paragraph(leg_text)


def discover_blocks(doc: Document, end: int) -> list[dict]:
    seen: set[int] = set()
    blocks: list[dict] = []
    for i in range(end):
        t = doc.paragraphs[i].text.strip()
        m = re.match(r"^Figure\s+(\d+)\.", t, re.I)
        if not m:
            continue
        n = int(m.group(1))
        if n in seen or n not in FIGURE_PNG:
            continue
        seen.add(n)

        j = i + 1
        while j < end and not doc.paragraphs[j].text.strip():
            j += 1
        if j >= end:
            continue
        if re.match(r"^Figure\s+\d+\.", doc.paragraphs[j].text.strip(), re.I):
            continue

        blocks.append({"num": n, "cap_idx": i, "leg_idx": j})
    return blocks


def insert_image_before_legend(doc: Document, leg_idx: int, png: Path, width_in: float) -> bool:
    if leg_idx <= 0:
        return False
    prev = doc.paragraphs[leg_idx - 1]
    if para_has_image(prev):
        return False
    if not png.is_file():
        print(f"  ✗ missing: {png}")
        return False

    img_p = doc.paragraphs[leg_idx].insert_paragraph_before()
    img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    img_p.add_run().add_picture(str(png), width=Inches(width_in))
    return True


def main() -> int:
    doc = Document(str(MANUSCRIPT))
    end = find_supp_start(doc)

    n_removed = remove_image_only_paragraphs(doc, end)
    print(f"Removed {n_removed} misplaced image paragraph(s)")

    split_all_merged_captions(doc)
    end = find_supp_start(doc)

    blocks = discover_blocks(doc, end)
    blocks.sort(key=lambda b: b["num"])
    print(f"Found {len(blocks)} figure block(s)")

    # Insert high legend index first
    for b in sorted(blocks, key=lambda x: -x["leg_idx"]):
        n = b["num"]
        w = FIGURE_WIDTH_IN.get(n, DEFAULT_WIDTH)
        ok = insert_image_before_legend(doc, b["leg_idx"], FIGURE_PNG[n], w)
        print(
            f"  Figure {n}: caption @{b['cap_idx']}, legend @{b['leg_idx']} "
            f"→ {'OK' if ok else 'skip'}"
        )

    doc.save(str(MANUSCRIPT))

    # Verify layout
    doc2 = Document(str(MANUSCRIPT))
    end2 = find_supp_start(doc2)
    ok_count = 0
    for b in discover_blocks(doc2, end2):
        c, leg = b["cap_idx"], b["leg_idx"]
        good = leg == c + 2 and para_has_image(doc2.paragraphs[c + 1])
        if good:
            ok_count += 1
        else:
            print(f"  ⚠ Figure {b['num']}: cap={c}, img@+1={para_has_image(doc2.paragraphs[c+1]) if c+1<end2 else False}, leg={leg}")
    print(f"\nSaved {MANUSCRIPT} — {ok_count}/{len(blocks)} figures in caption→image→legend order")
    return 0 if ok_count == len(blocks) else 1


if __name__ == "__main__":
    raise SystemExit(main())
