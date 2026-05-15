#!/usr/bin/env python3
"""Audit supplementary citations vs end inventory; report mismatches."""

from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document

MANUSCRIPT = Path(__file__).resolve().parent / "NewManuscript.docx"

METHOD_TITLES = {
    1: "Scope and relationship to the main manuscript",
    2: "Primary analytic cohort (discovery)",
    3: "Synergy and enrichment metrics (multiplicative, additive, and protective scores)",
    4: "Stage-stratified Cox and multivariable extensions",
    5: "DX2COLLECTION_YEAR × mutation-combination interaction models and sensitivity filters",
    6: "Internal validation (stratified split and cross-validation)",
    7: "External validation in public PDAC cohorts (cBioPortal)",
    8: "Overview of supplementary figure content",
    9: "Software and computational environment",
}

FIGURE_TITLES = {
    1: "Triple-mutation additive deviations from independence",
    2: "Univariate Cox forest plot for top triple-mutation combinations",
    3: "Diagnosis-to-collection timing × pairwise combination interaction volcano",
    4: "Diagnosis-to-collection timing × triple-mutation interaction volcano",
    5: "TP53/KRAS patterns in short overall-survival subsets",
    6: "Diabetes-stratified exploratory gene analysis",
    7: "Triple-mutation timing interaction volcano under DX2COLLECTION_YEAR 0–5 years",
    8: "Multiple-comparison correction diagnostics",
    9: "Internal discovery–validation cohort diagnostics",
    10: "Pathway-oriented functional validation summary",
    11: "Deceased-versus-living comprehensive mutation enrichment overview",
    12: "Multivariable Cox internal validation summary by stage",
}

TABLE_TITLES = {
    1: "Timing-interaction sensitivity comparison",
    2: "External replication of TP53+KRAS",
    3: "Stage-stratified mutation frequencies",
    4: "Diagnosis-to-collection timing × combination interaction Cox models",
    5: "Timing-interaction Cox outputs under dx_yr sensitivity filters",
    6: "Multivariable Cox coefficients and internal validation metrics",
    7: "Three-group stage landscape",
    8: "Top-ranked synergistic dual-mutation combinations",
    9: "Internal discovery–validation cohort comparison",
    10: "Dual-mutation combination significance, FDR/Bonferroni",
    11: "Deceased-versus-living enrichment",
}

# paragraph needle -> required refs (semantic check)
SEMANTIC_CHECKS = [
    ("Internal validation of descriptive synergy", {"method": {6}, "figure": {9}, "table": {9}}),
    ("Using the public cBioPortal API", {"method": {7}, "table": {2}}),
    ("The master analytic cohort comprised 2330", {"method": {2}, "table": {3, 7, 11}}),
    ("Statistical significance was assessed using chi-square", {"method": {1}}),
    ("Figure 3 summarizes stage-stratified Cox", {"method": {4}, "table": {6}, "figure": {12}}),
    ("DX2COLLECTION YEAR measures", {"method": {5}, "table": {1, 5}, "figure": {3, 4, 7}}),
    ("Table 1. Timing interaction highlights", {"table": {1}}),
    ("triple-level panels:", {"figure": {1, 2}}),
]


def extract_refs(text: str, kind: str) -> set[int]:
    nums: set[int] = set()
    k = kind.capitalize()
    for m in re.finditer(rf"Supplementary {k} (\d+)", text):
        nums.add(int(m.group(1)))
    for m in re.finditer(rf"Supplementary {k}s (\d+)[–-](\d+)", text):
        nums.add(int(m.group(1)))
        nums.add(int(m.group(2)))
    for m in re.finditer(rf"Supplementary {k}s [\d–-]+ and Supplementary {k} (\d+)", text):
        nums.add(int(m.group(1)))
    for m in re.finditer(rf"Supplementary {k}s (\d+) and (\d+)", text):
        nums.add(int(m.group(1)))
        nums.add(int(m.group(2)))
    for m in re.finditer(rf"Supplementary {k}s (\d+), (\d+), and (\d+)", text):
        nums.update(int(x) for x in m.groups())
    return nums


def parse_end_inventory(doc: Document, supp: int, ack: int) -> dict[str, dict[int, str]]:
    items: dict[str, dict[int, str]] = {"method": {}, "figure": {}, "table": {}}
    section: str | None = None
    for i in range(supp, ack):
        t = doc.paragraphs[i].text.strip()
        if t == "Supplementary Methods":
            section = "method"
        elif t == "Supplementary Figures":
            section = "figure"
        elif t == "Supplementary Tables":
            section = "table"
        elif section:
            m = re.match(r"Supplementary (Method|Figure|Table) (\d+)\.\s*(.+)", t)
            if m:
                items[m.group(1).lower()][int(m.group(2))] = m.group(3)
    return items


def title_ok(actual: str, expected_prefix: str) -> bool:
    return expected_prefix.lower()[:40] in actual.lower()[:80]


def main() -> int:
    doc = Document(str(MANUSCRIPT))
    supp = next(i for i, p in enumerate(doc.paragraphs) if p.text.strip() == "Supplementary Materials")
    ack = next(i for i, p in enumerate(doc.paragraphs) if p.text.strip() == "Acknowledgments")

    body_m, body_f, body_t = set(), set(), set()
    for i in range(supp):
        t = doc.paragraphs[i].text
        body_m |= extract_refs(t, "method")
        body_f |= extract_refs(t, "figure")
        body_t |= extract_refs(t, "table")

    end = parse_end_inventory(doc, supp, ack)
    errors: list[str] = []
    warnings: list[str] = []

    for typ, maxn, titles in [
        ("method", 9, METHOD_TITLES),
        ("figure", 12, FIGURE_TITLES),
        ("table", 11, TABLE_TITLES),
    ]:
        body_set = {"method": body_m, "figure": body_f, "table": body_t}[typ]
        end_set = set(end[typ])
        expected = set(range(1, maxn + 1))

        if end_set != expected:
            errors.append(f"End list {typ}: expected 1–{maxn}, got {sorted(end_set)}")

        missing_body = expected - body_set
        if missing_body:
            warnings.append(f"Never cited in body — {typ} {sorted(missing_body)}")

        uncited_in_end = body_set - end_set
        if uncited_in_end:
            errors.append(f"Cited in body but missing from end list — {typ} {sorted(uncited_in_end)}")

        for n in range(1, maxn + 1):
            if n in end[typ] and not title_ok(end[typ][n], titles[n]):
                errors.append(f"End list title mismatch {typ} {n}: got '{end[typ][n][:50]}...'")

        order = sorted(end[typ])
        if order != list(range(1, maxn + 1)):
            errors.append(f"End list {typ} not sequential: {order}")

    for needle, required in SEMANTIC_CHECKS:
        for i in range(supp):
            t = doc.paragraphs[i].text
            if needle not in t:
                continue
            for typ, nums in required.items():
                found = extract_refs(t, typ)
                if not nums.issubset(found):
                    errors.append(
                        f"Para {i} [{needle[:40]}...]: expected {typ} {sorted(nums)}, found {sorted(found)}"
                    )
            break
        else:
            warnings.append(f"Semantic check anchor not found: {needle[:50]}")

    # Duplicate / legacy patterns
    for i in range(supp):
        t = doc.paragraphs[i].text
        if re.search(r"Supplementary Methods\s*§", t):
            errors.append(f"Para {i}: legacy § notation")
        if re.search(r"Supplementary (?:Figure|Table) S\d+", t):
            errors.append(f"Para {i}: legacy S-prefix")
        if ".." in t:
            warnings.append(f"Para {i}: double period")

    print("=== SUPPLEMENTARY REFERENCE AUDIT ===\n")
    print(f"Body cited: Methods {sorted(body_m)}, Figures {sorted(body_f)}, Tables {sorted(body_t)}")
    print(f"End list:   Methods {sorted(end['method'])}, Figures {sorted(end['figure'])}, Tables {sorted(end['table'])}")

    if errors:
        print(f"\nERRORS ({len(errors)}):")
        for e in errors:
            print("  ✗", e)
    else:
        print("\n✓ No structural errors.")

    if warnings:
        print(f"\nWARNINGS ({len(warnings)}):")
        for w in warnings:
            print("  !", w)

    return 1 if errors else 0


if __name__ == "__main__":
    raise SystemExit(main())
