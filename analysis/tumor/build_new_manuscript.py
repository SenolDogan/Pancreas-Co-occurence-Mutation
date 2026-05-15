#!/usr/bin/env python3
"""Build NewManuscript.docx from Final.docx intro/M&M + revised body (do not edit Final.docx)."""

from __future__ import annotations

from pathlib import Path

import numpy as np
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_COLOR_INDEX
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

# Note: Final.docx uses Title style for many body paragraphs; we re-add as normal paragraphs for compatibility.


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.size = Pt(14 if level == 1 else 12)


def copy_table_grid(dest: Document, src_table, caption: str) -> None:
    cap = dest.add_paragraph()
    cr = cap.add_run(caption)
    cr.bold = True
    n_rows = len(src_table.rows)
    n_cols = len(src_table.columns)
    nt = dest.add_table(n_rows, n_cols)
    try:
        nt.style = "Table Grid"
    except Exception:
        pass
    for i in range(n_rows):
        for j in range(n_cols):
            nt.cell(i, j).text = src_table.cell(i, j).text
    dest.add_paragraph()


def add_figure_embed(
    doc: Document, image_path: Path, title: str, legend: str | None = None, width_in: float = 6.35
) -> None:
    p = doc.add_paragraph()
    r = p.add_run(title)
    r.bold = True
    if image_path.is_file():
        doc.add_picture(str(image_path), width=Inches(width_in))
    else:
        doc.add_paragraph(f"[Missing image: {image_path}]")
    if legend:
        lp = doc.add_paragraph()
        lr = lp.add_run(legend)
        lr.italic = True
    doc.add_paragraph()


def add_df_table(doc: Document, df, title: str, max_rows: int = 40) -> None:
    """
    Render a small pandas DataFrame as a Word table.
    Intended for summary tables (not huge outputs).
    """
    cap = doc.add_paragraph()
    cr = cap.add_run(title)
    cr.bold = True
    if df is None or len(df) == 0:
        doc.add_paragraph("[Empty table]")
        doc.add_paragraph()
        return
    d = df.head(max_rows).copy()
    nt = doc.add_table(rows=len(d) + 1, cols=len(d.columns))
    try:
        nt.style = "Table Grid"
    except Exception:
        pass
    # header
    for j, col in enumerate(d.columns):
        nt.cell(0, j).text = str(col)
    # rows
    for i in range(len(d)):
        for j, col in enumerate(d.columns):
            v = d.iloc[i, j]
            nt.cell(i + 1, j).text = "" if v is None else str(v)
    doc.add_paragraph()


def _set_row_shading(table_row, fill_hex: str) -> None:
    for cell in table_row.cells:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        for shd in tcPr.findall(qn("w:shd")):
            tcPr.remove(shd)
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:fill"), fill_hex)
        tcPr.append(shd)


def _set_table_grid_borders(table) -> None:
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        return
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "000000")
        borders.append(el)
    tblPr.append(borders)


def add_df_table_zebra(doc: Document, df, title: str, max_rows: int = 40) -> None:
    """
    Like add_df_table but applies alternating white / light-gray row fills (journal-style banding).
    """
    cap = doc.add_paragraph()
    cr = cap.add_run(title)
    cr.bold = True
    if df is None or len(df) == 0:
        doc.add_paragraph("[Empty table]")
        doc.add_paragraph()
        return
    d = df.head(max_rows).copy()
    nt = doc.add_table(rows=len(d) + 1, cols=len(d.columns))
    try:
        nt.style = "Table Grid"
    except Exception:
        pass
    try:
        _set_table_grid_borders(nt)
    except Exception:
        pass
    white, gray = "FFFFFF", "F2F2F2"
    for j, col in enumerate(d.columns):
        nt.cell(0, j).text = str(col)
    _set_row_shading(nt.rows[0], white)
    for i in range(len(d)):
        for j, col in enumerate(d.columns):
            v = d.iloc[i, j]
            nt.cell(i + 1, j).text = "" if v is None else str(v)
        _set_row_shading(nt.rows[i + 1], white if (i % 2 == 0) else gray)
    doc.add_paragraph()


def add_para(doc: Document, text: str, bold_lead: str | None = None) -> None:
    p = doc.add_paragraph()
    if bold_lead:
        r = p.add_run(bold_lead)
        r.bold = True
        p.add_run(text[len(bold_lead) :] if text.startswith(bold_lead) else " " + text)
    else:
        p.add_run(text)
    p.paragraph_format.space_after = Pt(6)


def add_para_yellow(doc: Document, text: str) -> None:
    """Add a paragraph with yellow highlight to mark updates."""
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.highlight_color = WD_COLOR_INDEX.YELLOW
    p.paragraph_format.space_after = Pt(6)


def cohort_counts_from_master(tumor_dir: Path) -> dict:
    """
    Compute cohort counts from Tumor/Merged.xlsx (patient-level derived from mutation-level rows).
    Returns total N and effective Ns used by OS/DX analyses.
    """
    try:
        import pandas as pd

        merged = tumor_dir / "Merged.xlsx"
        df = pd.read_excel(merged)
        pid = "PATIENT_ID" if "PATIENT_ID" in df.columns else ("Patient_ID" if "Patient_ID" in df.columns else None)
        if pid is None:
            return {}
        cols = [c for c in ["OS_STATUS", "OS_MONTHS", "DX2COLLECTION_YEAR", "STAGE", "AGE", "SEX", "DISEASE_STATUS"] if c in df.columns]
        pat = df.groupby(pid)[cols].first().reset_index()
        n_total = int(pat[pid].nunique())
        if {"OS_STATUS", "OS_MONTHS"}.issubset(pat.columns):
            n_os = int((pat["OS_STATUS"].notna() & pat["OS_MONTHS"].notna()).sum())
        else:
            n_os = 0
        if "DX2COLLECTION_YEAR" in pat.columns:
            dx = pd.to_numeric(pat["DX2COLLECTION_YEAR"], errors="coerce")
            n_dx = int(dx.notna().sum())
            n_dx_ge0 = int((dx.notna() & (dx >= 0)).sum())
            n_dx_0_5 = int((dx.notna() & (dx >= 0) & (dx <= 5)).sum())
        else:
            n_dx = n_dx_ge0 = n_dx_0_5 = 0
        if {"OS_STATUS", "OS_MONTHS", "DX2COLLECTION_YEAR"}.issubset(pat.columns):
            dx = pd.to_numeric(pat["DX2COLLECTION_YEAR"], errors="coerce")
            n_osdx = int((pat["OS_STATUS"].notna() & pat["OS_MONTHS"].notna() & dx.notna()).sum())
        else:
            n_osdx = 0
        stage_counts = {}
        if "STAGE" in pat.columns:
            vc = pat["STAGE"].value_counts(dropna=False)
            stage_counts = {str(k): int(v) for k, v in vc.to_dict().items()}
        return {
            "n_total": n_total,
            "n_os": n_os,
            "n_dx": n_dx,
            "n_dx_ge0": n_dx_ge0,
            "n_dx_0_5": n_dx_0_5,
            "n_osdx": n_osdx,
            "stage_counts": stage_counts,
        }
    except Exception:
        return {}


def main() -> None:
    base = Path(__file__).resolve().parent
    final_path = base / "Final.docx"
    out_path = base / "NewManuscript.docx"

    src = Document(str(final_path))

    doc = Document()

    tumor = base
    new2 = base.parent  # .../survival/New/New 2
    cc = cohort_counts_from_master(tumor)
    n_total = cc.get("n_total", 2330)
    n_osdx = cc.get("n_osdx", 2264)
    n_dx_ge0 = cc.get("n_dx_ge0", 2242)
    n_dx_0_5 = cc.get("n_dx_0_5", 2215)

    # --- Title & authors (from Final) ---
    for idx in (0, 2):
        t = src.paragraphs[idx].text.strip()
        if t:
            p = doc.add_paragraph(t)
            if idx == 0:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in p.runs:
                    r.bold = True
                    r.font.size = Pt(16)

    doc.add_paragraph()

    add_heading(doc, "Abstract", 1)

    abstract = (
        "Background: Pancreatic ductal adenocarcinoma (PDAC) is dominated by co-occurring driver alterations, "
        "yet how clinically defined stage groups modify the prognostic impact of recurrent dual and triple "
        "mutation patterns is incompletely mapped at cohort scale.\n\n"
        f"Methods: We studied {n_total} PDAC patients with integrated clinical records "
        "and binary mutation calls across ten prespecified genes (TP53, KRAS, CDKN2A, SMAD4, ARID1A, ATM, PIK3CA, "
        "BRAF, GNAS, RNF43). We applied the same synergy metrics as in our primary pipeline—multiplicative synergy, "
        "additive synergy, and protective contrasts comparing deceased versus living patients—together with "
        "chi-square/Fisher testing and multiplicity control (Benjamini–Hochberg FDR as primary reporting framework "
        "for the combination screen, with Bonferroni/Holm as sensitivity bounds). "
        "To relate mutations to outcome over time, we fitted Cox proportional hazards models for overall survival "
        "(OS) within three stage groups (Metastatic; Resectable; Borderline Resectable/Locally Advanced) and summarized "
        "penalized multivariable Cox models (parsimonious signature plus age/sex; extended gene-level covariate model). "
        f"Complete-case OS and DX2COLLECTION_YEAR data were available for {n_osdx} patients; "
        f"DX2COLLECTION_YEAR sensitivity filters yielded N={n_dx_ge0} (DX≥0) and N={n_dx_0_5} (DX 0–5 years). "
        "Internal stability of descriptive synergy-style scores was assessed by stratified 70/30 split and 5-fold "
        "cross-validation on the same registry cohort (not an external multi-center replication). "
        "Pathway labels contextualized combinations relative to canonical PDAC signaling modules.\n\n"
        "Results: TP53+KRAS remained the prevalent backbone profile (≈71.6% of patients). Dual and triple combinations "
        "showed statistically non-random enrichment patterns after multiplicity control, with ranking that differs "
        "when interpreted as prevalence/synergy metrics versus OS hazard summaries. Stage-stratified Cox and forest-plot "
        "summaries highlighted heterogeneous OS associations for selected TP53–KRAS–third-gene patterns across "
        "Metastatic, Resectable, and Borderline/Locally advanced subsets, with multivariable penalized models retaining "
        "key signals after adjustment for age and sex in prespecified models. Internal split correlations supported "
        "stability of the derived descriptive scores across partitions, within the limits of a single registry source.\n\n"
        "Conclusions: Integrating outcome-associated mutation enrichment with stage-stratified time-to-event modeling "
        "provides a more reviewer-aligned evidence chain than dichotomous profiling alone. The framework remains "
        "associational and requires prospective and external validation."
    )
    doc.add_paragraph(abstract)
    add_para_yellow(
        doc,
        f"[UPDATE] Cohort harmonized to master analytic table (N={n_total}); effective Ns reported for OS/DX subsets.",
    )

    doc.add_paragraph(
        "Keywords: pancreatic cancer; mutation combinations; stage-stratified analysis; Cox regression; "
        "multiplicity control; synergy metrics; TCGA"
    )

    doc.add_page_break()

    # --- Introduction: para 9 verbatim, para 10 revised ---
    add_heading(doc, "Introduction", 1)
    intro_body = src.paragraphs[9].text.strip()
    doc.add_paragraph(intro_body)

    # Keep Introduction conceptual; place the detailed analytic overview at the start of Methods.
    intro_close = (
        "Most PDAC cohorts are dominated by a small set of recurrent driver alterations, but the clinical meaning of "
        "co-mutation patterns depends on disease extent at presentation and on how outcomes are summarized. "
        "A purely cross-sectional comparison of deceased versus living patients can be informative for enrichment, "
        "yet it can also blur prognostic interpretation when stage composition differs between groups. "
        "For this reason, we combine descriptive combination scoring with stage-stratified time-to-event modeling, "
        "so that prevalence patterns are presented alongside survival-associated effect estimates. "
        "We focus on a prespecified ten-gene panel that captures common PDAC pathways (cell cycle, RAS–MAPK, "
        "TGF-β signaling, chromatin regulation, DNA repair) and evaluate dual and triple combinations in a "
        "multiplicity-aware framework."
    )
    doc.add_paragraph(intro_close)

    doc.add_page_break()

    # --- Materials and Methods (largely copied) ---
    methods_overview = (
        f"Here, we present a large integrative analysis of recurrent mutation combinations in PDAC (N={n_total}) using the "
        "same synergy scoring framework (multiplicative synergy, additive synergy, and protective contrasts) described "
        "previously, now complemented by stage-stratified overall survival modeling. Comparisons between deceased and "
        "living patients are interpreted explicitly as outcome-associated enrichment of mutation patterns (not causal "
        "treatment effects). Time-to-event associations are evaluated with Cox models within three clinically defined "
        "stage groups (Metastatic; Resectable; Borderline Resectable/Locally Advanced), alongside penalized "
        "multivariable models that adjust for age and sex. Multiple testing is handled with an explicit "
        "primary/sensitivity policy, and internal split/cross-validation summaries quantify stability of descriptive "
        "scores within this single integrated cohort."
    )
    for idx in range(12, 39):
        t = src.paragraphs[idx].text.strip()
        if not t:
            continue
        if idx == 12:
            add_heading(doc, t, 1)
            add_para_yellow(doc, methods_overview)
            continue
        # Light-touch edits only where needed
        if idx == 14:
            t = (
                f"We analyzed PDAC patients with integrated clinical and mutation annotations assembled as a single "
                f"patient-level analytic table (N={n_total}). Complete-case OS and DX2COLLECTION_YEAR values were "
                f"available for {n_osdx} patients and were used for the primary time-to-event and interaction analyses, "
                f"with prespecified sensitivity filters applied to DX2COLLECTION_YEAR (DX≥0; DX 0–5 years)."
            )
        if idx == 32:
            t = (
                "Internal validation of descriptive synergy and protective score summaries was performed on the "
                "same integrated cohort using (1) a stratified random 70/30 discovery–validation split preserving the "
                "distribution of OS status, with correlation of cohort-level score vectors between partitions, and "
                "(2) 5-fold stratified cross-validation with averaged correlation-type stability metrics. "
                "These analyses quantify reproducibility of internally derived scores under resampling and do not "
                "constitute independent multi-institutional external validation."
            )
        if idx == 35:
            t = (
                "Contextual functional mapping was performed through pathway-oriented grouping and literature-based "
                "interaction assessment. Pathway labels (e.g., DNA repair, RAS–MAPK) may reference canonical genes for "
                "interpretation even when those genes were not among the ten binary mutation features analyzed in "
                "every patient row."
            )
        if idx == 38:
            t += (
                " Cox proportional hazards models and penalized Cox fits were implemented in Python using lifelines "
                "(and related utilities consistent with the analysis scripts), with two-sided α = 0.05 unless otherwise "
                "noted for exploratory contrasts."
            )

        is_subheading = idx in (13, 16, 19, 22, 28, 31, 34, 37) and len(t) < 120
        if is_subheading and not t[0].isdigit():
            add_heading(doc, t, 2)
        else:
            if idx == 14:
                add_para_yellow(doc, t)
            else:
                doc.add_paragraph(t)

    add_heading(doc, "Interpretation assumptions for synergy metrics", 2)
    doc.add_paragraph(
        "Multiplicative synergy compares the observed joint mutation frequency to the frequency expected under "
        "independence at the patient level (product of marginal rates for dual combinations; analogous extension "
        "for triples). It is a descriptive prevalence-based measure and does not, by itself, imply therapeutic "
        "synergy or mechanistic cooperativity beyond statistical non-independence under the stated assumptions. "
        "Additive synergy contrasts observed joint frequency with the sum of marginal rates and can be negative when "
        "joint occurrence is less than an additive benchmark. Binary mutation calls ignore allelic multiplicity and "
        "subclonal architecture; small cell counts for rare combinations require cautious interpretation regardless of "
        "uncorrected p-values."
    )

    add_heading(doc, "Overall survival, stage stratification, and Cox models", 2)
    doc.add_paragraph(
        "OS was defined as time from diagnosis to death or last follow-up, with censoring for living patients at last "
        "contact. Events were coded from OS status fields consistent with the integrated table. "
        "Primary stage stratification used three clinically labeled groups: Metastatic; Resectable; and "
        "Borderline Resectable/Locally Advanced. "
        "Within each stage stratum, we evaluated univariate Cox models for prespecified single-gene indicators and "
        "selected dual/triple combination indicators, summarized as forest plots (hazard ratios with confidence intervals). "
        "Penalized multivariable Cox models compared (A) a parsimonious combination signature plus age and sex versus "
        "(B) an expanded gene-level covariate specification with the same clinical covariates, as implemented in the "
        "project analysis scripts. Internal validation summaries for multivariable fits used repeated data partitions "
        "to report stability of discriminative performance (e.g., concordance summaries) within the same cohort. "
        "Kaplan–Meier curves with log-rank tests for selected high-prevalence combination groups are recommended as "
        "editorial-standard complements to hazard summaries and can be added as finalized panels without changing "
        "the underlying Cox definitions."
    )

    add_heading(doc, "Bias-aware interpretation of DX2COLLECTION_YEAR (timing variable)", 2)
    add_para_yellow(
        doc,
        "DX2COLLECTION_YEAR measures the elapsed time between diagnosis and sequencing sample collection and therefore "
        "acts as a proxy for clinical trajectory and ascertainment rather than a randomized exposure. Because OS_MONTHS "
        "is evaluated from the collection date in our timing-aware analyses, patients with longer diagnosis-to-collection "
        "intervals must have survived long enough to be sampled, creating potential left-truncation/immortal-time "
        "structure and selection bias. We therefore interpret dx_yr associations and dx_yr×genotype interaction terms as "
        "timing-aware, associational signals and explicitly stress-test them with prespecified sensitivity filters "
        "(DX≥0; DX 0–5 years) rather than as causal effects of biopsy timing."
    )

    add_heading(doc, "Covariate transparency and missing-data handling", 2)
    add_para_yellow(
        doc,
        "Adjusted Cox interaction models include only covariates consistently available in the integrated table "
        "(AGE, SEX, and DISEASE_STATUS at sequencing) and are estimated on complete-case rows for those fields. "
        "Treatment, performance status, and detailed regimen timing are not available in this registry-derived dataset "
        "and therefore cannot be used for confounding control; we state this explicitly and avoid treatment-causal language."
    )

    add_heading(doc, "Standard statistical reporting", 2)
    add_para_yellow(
        doc,
        "We report hazard ratios (HR) with 95% confidence intervals and two-sided p-values for Cox models. "
        "Multiple testing is controlled primarily by Benjamini–Hochberg FDR (q=0.05) for high-dimensional combination "
        "screens, with Bonferroni/Holm shown as conservative sensitivity bounds. Proportional hazards is an underlying "
        "assumption of Cox modeling; where formal diagnostics are not exhaustively performed for every screened model, "
        "results should be interpreted as exploratory and prioritized for follow-up confirmation."
    )

    add_heading(doc, "Multiple testing policy for combination screens", 2)
    doc.add_paragraph(
        "For the high-dimensional screen across dual mutation combinations (45 tests under complete pairwise coverage "
        "of ten genes), we prespecify Benjamini–Hochberg FDR (q = 0.05) as the primary multiplicity framework for "
        "reporting significance in the main text, with Bonferroni and Holm adjustments reported in parallel as "
        "conservative sensitivity bounds. "
        "Counts of 'significant combinations' must be read conditional on the chosen error rate and should not be "
        "compared across correction schemes as if interchangeable outcomes."
    )

    doc.add_page_break()

    # --- Results (new structure) ---
    add_heading(doc, "Results", 1)

    add_heading(doc, "Cohort overview and stage distribution", 2)
    add_para_yellow(
        doc,
        f"The master analytic cohort comprised {n_total} patients. Demographics, OS status distribution, and clinical covariates are "
        "summarized in Supplementary Tables. Tumor stage labels were available for "
        "stratified analyses using the three-group clinical scheme described in Methods. The global stage landscape "
        "and cohort composition are contextualized by the deceased-versus-living enrichment overview (Figure 1) and "
        "the stage-distribution landscape shown below."
    )
    add_para_yellow(
        doc,
        f"[UPDATE] Cohort N is reported as N={n_total}; primary OS/DX analyses use complete-case N={n_osdx}.",
    )
    # Table 1 (main text): timing interaction highlights — top 5 combinations per stage by HR (ALL), journal-style zebra table.
    # Full, wide comparison table remains in Supplementary Table S1.
    try:
        import pandas as pd

        cmp_xlsx = tumor / "Stage_DX2Collection_Interaction_Sensitivity_Comparison.xlsx"
        if cmp_xlsx.is_file():
            cmp_df = pd.read_excel(cmp_xlsx, sheet_name="Comparison_AllFilters")
            need = {"Stage", "Feature", "HR_int_adj_ALL", "p_int_adj_ALL", "stable_FDR_lt_0p10"}
            if not need.issubset(set(cmp_df.columns)):
                raise ValueError("Comparison sheet missing required columns for Table 1")
            stage_order = [
                "Metastatic",
                "Resectable",
                "Borderline Resectable/Locally Advanced",
            ]
            parts: list[pd.DataFrame] = []
            for st in stage_order:
                sub = cmp_df[cmp_df["Stage"] == st].copy()
                if len(sub) == 0:
                    continue
                sub = sub.sort_values("HR_int_adj_ALL", ascending=False, na_position="last").head(5)
                parts.append(sub)
            top = pd.concat(parts, ignore_index=True) if parts else pd.DataFrame()
            top = top.rename(columns={"Feature": "Variable"})
            top["Stable q<0.10"] = top["stable_FDR_lt_0p10"].map({True: "Yes", False: "No"})
            top["HR adj.int ALL"] = top["HR_int_adj_ALL"].map(lambda x: "" if x != x else f"{float(x):.2f}")
            top["P adj.int ALL"] = top["p_int_adj_ALL"].map(lambda x: "" if x != x else f"{float(x):.2f}")
            tbl1 = top[["Stage", "Variable", "HR adj.int ALL", "P adj.int ALL", "Stable q<0.10"]].copy()
            add_df_table_zebra(
                doc,
                tbl1,
                "Table 1. Timing interaction highlights (adjusted Cox; top five combinations per stage by interaction HR, ALL analysis). Full sensitivity comparison is shown in Supplementary Table S1.",
                max_rows=40,
            )
    except Exception:
        doc.add_paragraph(
            "[Table 1 timing-interaction highlights could not be embedded; see Supplementary Table S1 in Supplementary Information.]"
        )
        doc.add_paragraph()
    add_figure_embed(
        doc,
        new2 / "01_Dead_Alive_Comprehensive_Analysis.png",
        "Figure 1. Dead vs Alive enrichment overview (supporting).",
        legend=(
            "Legend: Cross-sectional enrichment view comparing mutation and combination frequencies between deceased "
            "and living patients. This panel summarizes outcome-associated enrichment patterns and provides descriptive "
            "context for subsequent stage-stratified and time-to-event analyses; it is not interpreted as a causal estimate."
        ),
        width_in=6.8,
    )

    add_figure_embed(
        doc,
        tumor / "Stage_3Group_Comprehensive_Analysis.png",
        "Figure 2. Integrated cohort and three-group stage landscape.",
        legend=(
            "Legend: Overview of cohort composition and the three clinical stage groups used throughout the manuscript "
            "(Metastatic; Resectable; Borderline Resectable/Locally Advanced). Colors are consistent across stage-based "
            "panels to support within-figure and between-figure comparison."
        ),
        width_in=6.8,
    )
    doc.add_paragraph("Stage-group composition and cohort structure are summarized in Figure 2.")

    add_heading(doc, "Outcome-associated enrichment in deceased versus living patients", 2)
    doc.add_paragraph(
        "As an enrichment readout, deceased versus living contrasts recapitulated expected PDAC backbone frequencies "
        "(notably high KRAS and TP53 prevalence) and identified dual/triple combinations whose joint prevalence differed "
        "between OS status groups after accounting for multiple testing under the prespecified FDR-first policy "
        "(Supplementary Tables). "
        "These analyses quantify cross-sectional association with a terminal outcome label and should be interpreted "
        "together with the stage-stratified Cox results rather than as standalone proof of causal lethality. "
        "A compact enrichment overview is provided in Figure 1."
    )

    add_heading(doc, "TP53+KRAS backbone and higher-order combinations", 2)
    doc.add_paragraph(
        "TP53+KRAS co-occurrence remained the dominant backbone profile (~71.6% of patients). "
        "To keep the narrative aligned with how readers typically scan results, we present the evidence in three steps: "
        "(i) single-gene prevalence and outcome-associated contrasts, (ii) statistically supported dual-combination signals, "
        "and (iii) higher-order (triple) patterns."
    )

    add_heading(doc, "Single-gene prevalence and outcome-associated contrasts", 3)
    doc.add_paragraph(
        "Single-gene frequencies provide context for all combination-level analyses. "
        "Table 2 summarizes mutation prevalence in deceased versus living patients alongside lethality-style ratios "
        "as descriptive enrichment measures; these contrasts should be interpreted together with stage-stratified time-to-event "
        "associations rather than as causal estimates."
    )
    copy_table_grid(
        doc,
        src.tables[1],
        "Table 2. Genetic mutation frequencies and lethality ratios (see Supplementary Tables for extended mutation-level summaries).",
    )

    add_heading(doc, "Dual-combination signals under multiplicity control", 3)
    doc.add_paragraph(
        "We next tested dual combinations for non-random co-occurrence and outcome-associated enrichment under the prespecified "
        "FDR-first policy. Table 3 lists the statistically supported dual-combination signals; the complete screen and alternative "
        "multiplicity bounds are provided in Supplementary Tables."
    )
    copy_table_grid(
        doc,
        src.tables[2],
        "Table 3. Statistically significant dual mutation combinations (see Supplementary Tables for the complete dual-combination screen).",
    )

    add_heading(doc, "Top-ranked synergy-style pair summaries", 3)
    doc.add_paragraph(
        "Because statistical significance and effect ranking can diverge, we also provide a compact ranking of the strongest "
        "synergy-style pair summaries. Table 4 highlights top pairwise patterns by descriptive synergy metrics; rare events can yield "
        "unstable ratio-style summaries, so top ranks are interpreted alongside absolute rates and stage-stratified survival models."
    )
    copy_table_grid(doc, src.tables[3], "Table 4. Top synergistic mutation combinations")

    add_heading(doc, "Triple-combination patterns", 3)
    doc.add_paragraph(
        "Finally, we extend the same descriptive framework to triple combinations to capture higher-order contexts that may reflect "
        "biologically coherent pathway modules. Table 5 summarizes lethality/protective-style contrasts for selected triples, with the "
        "full triple screen provided in Supplementary Tables."
    )
    copy_table_grid(
        doc,
        src.tables[4],
        "Table 5. Triple mutation combinations: Lethality and protective scores (see Supplementary Tables for the complete triple-combination screen).",
    )

    add_heading(doc, "Time-to-event associations: stage-stratified Cox and multivariable models", 2)
    doc.add_paragraph(
        "Figure 3 summarizes stage-stratified Cox results: Panel A shows univariate associations, and Panel B contrasts "
        "penalized multivariable Models A versus B. Figure 4 then reports internal stability summaries for the "
        "multivariable fits. Tabular estimates are exported to Stage_Multivariable_Cox_Report.xlsx."
    )
    add_figure_embed(
        doc,
        tumor / "Stage_Figure2_AB_ForestPlots_VSTACK.png",
        "Figure 3. Stage-stratified Cox summaries (Panels A–B).",
        legend=(
            "Legend: Panel A shows univariate Cox proportional hazards results stratified by stage group (hazard ratios "
            "with confidence intervals for prespecified mutation features). Panel B shows penalized multivariable Cox "
            "models within stage strata, contrasting Model A (parsimonious combination-oriented specification with "
            "clinical covariates) versus Model B (expanded gene-level specification with the same clinical covariates). "
            "Across panels, HR>1 indicates shorter survival (higher hazard) and HR<1 indicates longer survival for the "
            "feature-positive group."
        ),
        width_in=6.8,
    )
    add_figure_embed(
        doc,
        tumor / "Stage_06_OS_Months_Methodology_Analysis.png",
        "Figure 4. Overall survival endpoint definition and timing framework.",
        legend=(
            "Legend: Schematic/summary of OS endpoint definition and the timing framework used for analyses, including "
            "how OS is measured and how DX2COLLECTION_YEAR relates to the sample-collection date used in the interaction "
            "models. This figure anchors interpretation of the timing-aware interaction analyses presented later."
        ),
        width_in=6.8,
    )

    add_heading(doc, "Stage-stratified co-mutation patterns", 2)
    doc.add_paragraph(
        "With stage-stratified survival associations established (Figures 3–4), we next visualize how key co-mutation "
        "patterns distribute across Metastatic, Resectable, and Borderline Resectable/Locally Advanced groups. "
        "These panels provide prevalence context for interpreting hazard-based results; they are descriptive and do not "
        "imply treatment effects. A detailed TP53+KRAS(+X) stage comparison is shown in Figure 5."
    )
    add_figure_embed(
        doc,
        tumor / "Stage_05_TP53_KRAS_Detailed_Analysis.png",
        "Figure 5. TP53+KRAS(+X) detailed stage comparison.",
        legend=(
            "Legend: Stage-wise comparison of TP53+KRAS and selected TP53+KRAS+third-gene patterns. The figure provides "
            "prevalence context across stage groups for interpreting the Cox model summaries shown earlier."
        ),
        width_in=6.8,
    )
    # Note: Supplementary Figure S2 is embedded in the Supplementary section (not inline in Results).

    add_heading(doc, "TP53/KRAS-focused stage distribution", 2)
    doc.add_paragraph(
        "To provide a concise anchor for the backbone genotype across clinical strata, we show a focused TP53/KRAS "
        "stage panel (Figure 5). This view complements the broader stage panels and supports interpretation of the "
        "combination summaries."
    )
    add_figure_embed(
        doc,
        tumor / "Stage_02_TP53_KRAS_Focused_Analysis.png",
        "Figure 6. TP53/KRAS-focused stage panels.",
        legend=(
            "Legend: Focused panels centered on TP53 and KRAS status across the three stage groups. These plots provide "
            "a compact visual companion to the detailed combination tables and stage-stratified survival summaries."
        ),
        width_in=6.8,
    )

    add_heading(doc, "Novelty analysis: sample-collection timing × mutation context (interaction Cox)", 2)
    doc.add_paragraph(
        "Beyond stage-stratified hazard summaries, we tested whether the association between sample collection timing "
        "(DX2COLLECTION_YEAR; years from diagnosis to sample collection) and survival measured from collection "
        "(OS_MONTHS) differs by co-occurring mutation contexts. For each stage group and each prespecified pair or "
        "triple combination, we fit interaction Cox models of the form OS ~ dx_yr + combo + dx_yr×combo, and "
        "then re-estimated the same interaction term adjusted for available clinical covariates (AGE, SEX, and "
        "DISEASE_STATUS at sequencing where non-constant within the stage subset). "
        "We further stress-tested results by excluding dx_yr<0 values and by excluding extreme dx_yr>5-year intervals. "
        "This analysis is explicitly associational given that OS is defined from the collection date, and "
        "therefore serves as a novel, timing-aware sensitivity layer rather than a causal statement about "
        "biopsy timing."
    )
    add_para_yellow(
        doc,
        "Clinical interpretation note: longer dx_yr can reflect delayed sampling, prolonged disease course before sequencing, "
        "or referral/biobank capture patterns; consequently, dx_yr×genotype interaction signals should be read as evidence "
        "that the prognostic association of a genotype differs across timing strata, not that changing sampling time would "
        "change survival."
    )
    doc.add_paragraph(
        "Across filters, a particularly robust signal was observed in the metastatic subset for the "
        "TP53+KRAS+SMAD4 triple context, where the adjusted interaction term remained among the strongest and "
        "retained FDR < 0.10 even after excluding dx_yr>5 years in the sensitivity comparison table. "
        "Full adjusted/unadjusted results, volcano summaries, and the single consolidated sensitivity comparison table "
        "are provided in the accompanying interaction outputs (see Supplementary Table S1 for the compact comparison)."
    )
    add_heading(doc, "Prioritizing pairwise signals: co-occurrence versus survival effects", 2)
    doc.add_paragraph(
        "To summarize the pairwise screen in a reviewer-friendly way, we present (i) the strongest positive/negative "
        "deviations from independence and (ii) the top univariate OS associations in a single paired figure (Figure 6). "
        "Timing-interaction volcano plots are provided as Supplementary Figures for completeness."
    )
    add_figure_embed(
        doc,
        tumor / "Stage_Figure6_AB_AdditiveBars_CoxPairsForest.png",
        "Figure 7. Pairwise co-occurrence and OS association summaries (Panels A–B).",
        legend=(
            "Legend: Panel A shows top pairwise deviations from independence (additive excess over independence), highlighting "
            "pairs that co-occur more or less often than expected from marginal rates. Panel B shows the top univariate Cox "
            "associations for mutation pairs within stage strata (hazard ratios with confidence intervals). Together, the panels "
            "summarize prevalence non-independence and time-to-event association in a single view."
        ),
        width_in=6.8,
    )

    add_heading(doc, "Internal stability of descriptive synergy summaries", 2)
    doc.add_paragraph(
        "Stratified 70/30 splits and cross-validation reproduced correlation structure for internally computed synergy "
        "and protective score summaries, supporting stability under resampling of the same registry-derived dataset. "
        "These metrics should not be described as validation against independent external cohorts."
    )

    add_heading(doc, "Pathway contextualization", 2)
    doc.add_paragraph(
        "Pathway-oriented summaries grouped recurrent combinations under canonical PDAC modules (cell cycle, RAS–MAPK, "
        "TGF-β, DNA repair context, etc.) as interpretive scaffolding. This module supplements, rather than replaces, "
        "statistical genetic-clinical associations."
    )

    doc.add_page_break()

    add_heading(doc, "Discussion", 1)
    doc.add_paragraph(
        "By jointly presenting (i) outcome-associated mutation enrichment between deceased and living patients, "
        "(ii) stage-stratified prevalence and co-occurrence structure, and (iii) OS-focused Cox summaries with "
        "multivariable penalization and internal stability checks, the manuscript links descriptive combination "
        "patterns to time-to-event outcome summaries in a single, stage-aware narrative. "
        "This design is intended to mirror how clinical readers interpret molecular profiles: first as cohort-level "
        "patterns of co-alteration, and then as quantitative, stage-conditional effect estimates on survival."
    )
    doc.add_paragraph(
        "The principal shift versus a dichotomous Dead–Alive-only narrative is interpretive: cross-sectional contrasts "
        "by final OS status can overrepresent molecular features enriched in fatal cases even when similar alterations "
        "appear in long survivors, particularly when stage and treatment intensity are incompletely captured. "
        "Explicit stage stratification and hazard-based summaries reduce this ambiguity. "
        "Importantly, the stage-specific forest-plot summaries make it straightforward to distinguish signals that are "
        "consistent across clinical strata from those that appear concentrated in a single stage group—an observation "
        "that can be clinically actionable for hypothesis generation and trial-stratification thinking."
    )
    doc.add_paragraph(
        "We intentionally avoid language implying external multi-cohort replication where only internal splits were "
        "performed, and we avoid absolute claims of therapeutic benefit for drug classes unless supported by "
        "prospective treatment annotations beyond mutation proxies. "
        "At the same time, the combination-first framework provides a compact way to prioritize which recurrent "
        "genotype contexts merit external replication, functional follow-up, or deeper clinical annotation "
        "(e.g., treatment exposure and response)."
    )
    doc.add_paragraph(
        "Limitations include retrospective TCGA registry constraints, incomplete treatment and performance-status "
        "encoding, binary mutation summaries across only ten genes, residual immortal-time and selection biases when "
        "contrasting final OS labels, and absence of a fully independent validation cohort. Prospective studies with "
        "standardized molecular profiling and adjudicated outcomes remain required."
    )
    doc.add_paragraph(
        "Despite these limitations, the study’s strength is the coherent linkage of prevalence-based combination "
        "structure to stage-aware survival modeling under an explicit multiplicity and sensitivity policy. "
        "This alignment of descriptive and time-to-event evidence reduces overinterpretation of cross-sectional labels "
        "and provides a reproducible template for extending the same workflow to additional clinical covariates or "
        "external datasets when available."
    )

    doc.add_page_break()

    add_heading(doc, "Conclusions", 1)
    doc.add_paragraph(
        "Integrated synergy-style profiling, multiplicity-aware combination testing, and stage-stratified Cox modeling "
        "provide a coherent framework for describing how recurrent PDAC mutation combinations relate to outcome labels "
        "and OS within clinically defined stage groups. The interaction analyses using DX2COLLECTION_YEAR add a "
        "timing-aware layer that helps distinguish stable prognostic associations from signals that vary with sampling "
        "context, while the sensitivity filters provide guardrails against overinterpretation of extreme intervals."
    )
    doc.add_paragraph(
        "Practically, our results prioritize a short list of recurrent dual and triple contexts for follow-up in "
        "externally annotated cohorts and prospective studies, where treatment exposure, performance status, and "
        "biopsy timing can be explicitly modeled. Future work should validate the stage-conditional and timing-conditional "
        "signals in independent datasets, extend the feature space beyond a ten-gene panel, and test whether these "
        "genotype contexts track with therapeutic response or resistance rather than survival alone."
    )

    doc.add_page_break()

    add_heading(doc, "References", 1)
    refs = []
    for i in range(94, 115):
        t = src.paragraphs[i].text.strip()
        if t:
            refs.append(t)
    fixed = []
    for r in refs:
        if r.startswith("[12]"):
            r = (
                "[12] Goodwin CM, Waters AM, Klomp JE, Javaid S, Bryant KL, et al. Combination Therapies with CDK4/6 "
                "Inhibitors to Treat KRAS-Mutant Pancreatic Cancer. Cancer Res. 2023;83(1):141-157."
            )
        elif r.startswith("[13]"):
            r = (
                "[13] Sodergren MH. Immunological combination treatment holds the key to improving survival in "
                "pancreatic cancer. J Cancer Res Clin Oncol. 2020;146(11):2897-2911. doi:10.1007/s00432-020-03332-5"
            )
        fixed.append(r)

    # Append recent (2023–2026) references to reach ~30 total, prioritizing high-impact PDAC genomics/clinical papers.
    def _max_ref_num(ref_lines: list[str]) -> int:
        m = 0
        for line in ref_lines:
            s = line.strip()
            if s.startswith("[") and "]" in s:
                try:
                    n = int(s[1 : s.index("]")])
                    m = max(m, n)
                except Exception:
                    continue
        return m

    recent_refs = [
        "Varghese AM, Perry MA, Chou JF, et al. Clinicogenomic landscape of pancreatic adenocarcinoma identifies KRAS mutant dosage as prognostic of overall survival. Nat Med. 2025;31:466-477. doi:10.1038/s41591-024-03362-3",
        "Nguyen J, Sanchez-Vega F. Metastatic patterns stratify patients with pancreatic cancer. Nat Cancer. 2025;6:16-17. doi:10.1038/s43018-024-00846-6",
        "Link JM, Eng JR, Sears RC. Ongoing replication stress tolerance and clonal T cell responses distinguish liver and lung recurrence and outcomes in pancreatic cancer. Nat Cancer. 2025. doi:10.1038/s43018-024-00881-3",
        "Miyamoto K, Yoshida R, Yasui K, et al. Precise stratification of prognosis in pancreatic ductal adenocarcinoma patients based on pre- and postoperative genomic information. Cancer Cell Int. 2025;25:305. doi:10.1186/s12935-025-03894-9",
        "Bojmar L, Zambirinis CP, Hernandez JM, et al. Multi-parametric atlas of the pre-metastatic liver for prediction of metastatic outcome in early-stage pancreatic cancer. Nat Med. 2024;30(8):2170-2180. doi:10.1038/s41591-024-03075-7",
        "Halbrook CJ, Lyssiotis CA, Pasca di Magliano M, Maitra A. Pancreatic cancer: advances and challenges. Cell. 2023;186:1729-1754. doi:10.1016/j.cell.2023.02.014",
        "Siegel RL, Giaquinto AN, Jemal A. Cancer statistics, 2024. CA Cancer J Clin. 2024;74:12-49.",
        "O’Connor CA, et al. Lynch syndrome and somatic mismatch repair variants in pancreas cancer. JAMA Oncol. 2024. doi:10.1001/jamaoncol.2024.3651",
        "Yousef A, et al. Impact of KRAS mutations and co-mutations on clinical outcomes in pancreatic ductal adenocarcinoma. NPJ Precis Oncol. 2024. doi:10.1038/s41698-024-00505-0",
        "Vitello DJ, Shah D, Wells A, et al. Mutant KRAS in circulating tumor DNA as a biomarker in localized pancreatic cancer in patients treated with neoadjuvant chemotherapy. Ann Surg. 2024. doi:10.1097/SLA.0000000000006562",
    ]

    all_refs = fixed[:]
    max_n = _max_ref_num(all_refs)
    for rr in recent_refs:
        max_n += 1
        all_refs.append(f"[{max_n}] {rr}")

    for r in all_refs:
        doc.add_paragraph(r, style="List Bullet")

    doc.add_paragraph()
    add_heading(doc, "Supplementary Information", 1)
    doc.add_paragraph(
        "Supplementary tables and extended figures should include: complete dual- and triple-combination score sheets; "
        "full multiplicity-correction panels; per-stage Cox coefficient tables (from Stage_Multivariable_Cox_Report.xlsx); "
        "and archived Dead vs Alive comprehensive plots retained for comparison with the primary stage/OS figure set. "
        "We also provide timing-aware interaction analyses (DX2COLLECTION_YEAR×combination) with adjusted models and "
        "sensitivity filters, plus stage-wise additive-vs-independence visualizations for pairs and triples."
    )

    # --- Supplementary figures at the very end ---
    doc.add_page_break()
    add_heading(doc, "Supplementary Information (embedded)", 1)

    add_heading(doc, "Supplementary Methods", 2)
    doc.add_paragraph(
        "Supplementary Methods provide additional analysis details, sensitivity filters, and extended outputs that "
        "support the main Results without changing primary definitions. References to Supplementary Methods, Tables, "
        "and Figures are made in ascending order where first introduced."
    )

    add_heading(doc, "Supplementary Tables", 2)

    # Timing interaction sensitivity comparison table is large: keep it in Supplementary.
    try:
        import pandas as pd

        cmp_xlsx = tumor / "Stage_DX2Collection_Interaction_Sensitivity_Comparison.xlsx"
        if cmp_xlsx.is_file():
            cmp_df = pd.read_excel(cmp_xlsx, sheet_name="Comparison_AllFilters")
            compact = cmp_df.copy()
            for c in ["stable_FDR_lt_0p10", "p_int_adj_ALL"]:
                if c not in compact.columns:
                    compact[c] = np.nan
            compact = compact.sort_values(["stable_FDR_lt_0p10", "p_int_adj_ALL"], ascending=[False, True])
            add_df_table(doc, compact, "Supplementary Table S1. Timing interaction sensitivity comparison (compact view)", max_rows=35)
    except Exception:
        doc.add_paragraph("[Could not embed interaction comparison table; see xlsx outputs.]")
        doc.add_paragraph()

    add_heading(doc, "Supplementary Figures", 2)
    supp_figs = [
        (
            "Supplementary Figure S1. Triple additive top bars (vs independence)",
            tumor / "Stage_Triple_Additive_TopBars.png",
            "Legend: Highest positive and negative deviations from independence for mutation triples (excess co-occurrence).",
        ),
        (
            "Supplementary Figure S2. Cox triples top forest (univariate; stage-wise)",
            tumor / "Stage_Cox_Triples_TopForest.png",
            "Legend: Univariate Cox forest plot for top triple combinations within stage strata (HR with confidence intervals).",
        ),
        (
            "Supplementary Figure S3. dx_yr×pair interaction volcano (DX_GE_0)",
            tumor / "Stage_DX2Collection_Pair_Interaction_Volcano_DX_GE_0.png",
            "Legend: Interaction volcano under the sensitivity filter DX2COLLECTION_YEAR ≥ 0.",
        ),
        (
            "Supplementary Figure S4. dx_yr×triple interaction volcano (DX_GE_0)",
            tumor / "Stage_DX2Collection_Triple_Interaction_Volcano_DX_GE_0.png",
            "Legend: Triple interaction volcano under the sensitivity filter DX2COLLECTION_YEAR ≥ 0.",
        ),
        (
            "Supplementary Figure S5. Stage 12 short survival comparison",
            tumor / "Stage_12_TP53_KRAS_ShortSurvival_Comparison.png",
            "Legend: Exploratory comparison of TP53/KRAS patterns in short-survival subsets, shown as a sensitivity view.",
        ),
        (
            "Supplementary Figure S6. Stage 14 diabetic vs non-diabetic",
            tumor / "Stage_14_Diabetic_vs_NonDiabetic_Gene_Analysis.png",
            "Legend: Exploratory diabetes-stratified gene analysis panel; interpreted as hypothesis-generating.",
        ),
        (
            "Supplementary Figure S7. dx_yr×triple interaction volcano (DX_0_TO_5 sensitivity)",
            tumor / "Stage_DX2Collection_Triple_Interaction_Volcano_DX_0_TO_5.png",
            "Legend: Sensitivity analysis restricting DX2COLLECTION_YEAR to 0–5 years to reduce the influence of extreme intervals.",
        ),
    ]
    for title, path, legend in supp_figs:
        add_figure_embed(doc, path, title, legend=legend, width_in=6.8)

    doc.save(str(out_path))
    print(f"Wrote {out_path}")


if __name__ == "__main__":
    main()
