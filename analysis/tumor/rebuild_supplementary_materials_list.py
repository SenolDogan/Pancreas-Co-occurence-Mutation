#!/usr/bin/env python3
"""Replace end-of-manuscript supplementary block with journal-style numbered inventory only."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


def delete_paragraph(paragraph) -> None:
    el = paragraph._element
    el.getparent().remove(el)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.size = Pt(14 if level == 1 else 12)


def add_list_item(doc: Document, text: str) -> None:
    doc.add_paragraph(text, style="List Number")


def main() -> None:
    path = Path(__file__).resolve().parent / "NewManuscript.docx"
    doc = Document(str(path))

    start_idx = None
    for i, p in enumerate(doc.paragraphs):
        t = p.text.strip()
        if t == "Supplementary Information" or t.startswith("Supplementary Information"):
            if start_idx is None and i > 100:  # end-of-manuscript block only
                start_idx = i
                break

    if start_idx is None:
        for i, p in enumerate(doc.paragraphs):
            if p.text.strip() == "Supplementary Materials":
                start_idx = i
                break

    if start_idx is None:
        raise SystemExit("Could not find supplementary section start index")

    # Delete from start_idx through end (reverse order)
    while len(doc.paragraphs) > start_idx:
        delete_paragraph(doc.paragraphs[start_idx])

    doc.add_page_break()
    title = doc.add_paragraph("Supplementary Materials")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.bold = True
        run.font.size = Pt(16)
    doc.add_paragraph()

    intro = doc.add_paragraph(
        "The following supplementary items accompany the main manuscript. "
        "Extended methods, figures, and tabular data are provided as separate "
        "supplementary files (Supplementary Methods; Supplementary Figures; Supplementary Tables)."
    )
    intro.paragraph_format.space_after = Pt(8)

    add_heading(doc, "Supplementary Methods", 2)
    methods = [
        "Scope and relationship to the main manuscript",
        "Primary analytic cohort (discovery)",
        "Synergy and enrichment metrics (multiplicative, additive, and protective scores)",
        "Stage-stratified Cox and multivariable extensions",
        "DX2COLLECTION_YEAR × mutation-combination interaction models and sensitivity filters",
        "Internal validation (stratified split and cross-validation)",
        "External validation in public PDAC cohorts (cBioPortal)",
        "Overview of supplementary figure content",
        "Software and computational environment",
    ]
    for i, m in enumerate(methods, start=1):
        add_list_item(doc, f"{i}. {m}")

    doc.add_paragraph()
    add_heading(doc, "Supplementary Figures", 2)
    figures = [
        "Triple-mutation additive deviations from independence (top positive and negative bars by stage)",
        "Univariate Cox forest plot for top triple-mutation combinations by stage",
        "Diagnosis-to-collection timing × pairwise combination interaction volcano (DX2COLLECTION_YEAR ≥ 0)",
        "Diagnosis-to-collection timing × triple-mutation interaction volcano (DX2COLLECTION_YEAR ≥ 0)",
        "TP53/KRAS patterns in short overall-survival subsets (exploratory sensitivity)",
        "Diabetes-stratified exploratory gene analysis",
        "Triple-mutation timing interaction volcano under DX2COLLECTION_YEAR 0–5 years sensitivity filter",
        "Multiple-comparison correction diagnostics for combination screens",
        "Internal discovery–validation cohort diagnostics",
        "Pathway-oriented functional validation summary",
        "Deceased-versus-living comprehensive mutation enrichment overview",
        "Multivariable Cox internal validation summary by stage",
    ]
    for i, cap in enumerate(figures, start=1):
        p = doc.add_paragraph(style="List Number")
        r1 = p.add_run(f"Figure S{i}. ")
        r1.bold = True
        p.add_run(cap)

    doc.add_paragraph()
    add_heading(doc, "Supplementary Tables", 2)
    # Numbering S1–S11 aligns with in-text citations and Supplementary Tables.xlsx workbook groups.
    tables = [
        (
            "Timing-interaction sensitivity comparison for DX2COLLECTION_YEAR × mutation combinations "
            "(adjusted Cox; ALL, DX≥0, and 0–5 year filters)"
        ),
        (
            "External replication of TP53+KRAS co-mutation prevalence and univariate overall survival "
            "in public PDAC cohorts (TCGA, Pan-Cancer Atlas, CPTAC, MSK, and discovery cohort)"
        ),
        (
            "Stage-stratified mutation frequencies, pairwise and triple additive co-occurrence, "
            "and univariate Cox overall survival models"
        ),
        (
            "Diagnosis-to-collection timing × combination interaction Cox models (primary cohort)"
        ),
        (
            "Timing-interaction Cox outputs under dx_yr sensitivity filters (ALL, DX≥0, DX 0–5 years)"
        ),
        (
            "Multivariable Cox coefficients and internal validation metrics by stage"
        ),
        (
            "Three-group stage landscape, mutation rates, clinical associations, and synergy summaries"
        ),
        (
            "Top-ranked synergistic dual-mutation combinations by stage"
        ),
        (
            "Internal discovery–validation cohort comparison and cross-validation metrics"
        ),
        (
            "Dual-mutation combination significance, FDR/Bonferroni summaries, and multiplicity correction"
        ),
        (
            "Deceased-versus-living enrichment: cohort summary, mutation frequencies, synergy, and clinical factors"
        ),
    ]
    for i, cap in enumerate(tables, start=1):
        p = doc.add_paragraph(style="List Number")
        r1 = p.add_run(f"Table S{i}. ")
        r1.bold = True
        p.add_run(cap)

    doc.save(str(path))
    print(f"Rebuilt supplementary inventory from paragraph {start_idx}: {path}")


if __name__ == "__main__":
    main()
