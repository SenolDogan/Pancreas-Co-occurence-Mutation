#!/usr/bin/env python3
"""
Restore type-specific supplementary numbering:
  Supplementary Methods 1–9, Figures 1–12, Tables 1–11
Updates NewManuscript.docx (body + end inventory) and supplementary/ files only.
"""

from __future__ import annotations

import re
import tempfile
from copy import deepcopy
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.shared import Pt
from docx.text.paragraph import Paragraph

TUMOR = Path(__file__).resolve().parent
MANUSCRIPT = TUMOR / "NewManuscript.docx"
SUPP_DIR = TUMOR / "supplementary"

METHOD_TITLES = {
    1: "Scope and relationship to the main manuscript",
    2: "Primary analytic cohort (discovery)",
    3: "Synergy and enrichment metrics (multiplicative, additive, and protective scores)",
    4: "Stage-stratified Cox and multivariable extensions",
    5: "DX2COLLECTION_YEAR × mutation-combination interaction models and sensitivity filters",
    6: "Internal validation (stratified split and cross-validation)",
    7: "External validation in public PDAC cohorts (cBioPortal)",
    8: "Overview of supplementary figure content",
    9: "Software and computational environment",
}

FIGURE_TITLES = {
    1: "Triple-mutation additive deviations from independence (top positive and negative bars by stage)",
    2: "Univariate Cox forest plot for top triple-mutation combinations by stage",
    3: "Diagnosis-to-collection timing × pairwise combination interaction volcano (DX2COLLECTION_YEAR ≥ 0)",
    4: "Diagnosis-to-collection timing × triple-mutation interaction volcano (DX2COLLECTION_YEAR ≥ 0)",
    5: "TP53/KRAS patterns in short overall-survival subsets (exploratory sensitivity)",
    6: "Diabetes-stratified exploratory gene analysis",
    7: "Triple-mutation timing interaction volcano under DX2COLLECTION_YEAR 0–5 years sensitivity filter",
    8: "Multiple-comparison correction diagnostics for combination screens",
    9: "Internal discovery–validation cohort diagnostics",
    10: "Pathway-oriented functional validation summary",
    11: "Deceased-versus-living comprehensive mutation enrichment overview",
    12: "Multivariable Cox internal validation summary by stage",
}

TABLE_TITLES = {
    1: "Timing-interaction sensitivity comparison for DX2COLLECTION_YEAR × mutation combinations (adjusted Cox; ALL, DX≥0, and 0–5 year filters)",
    2: "External replication of TP53+KRAS co-mutation prevalence and univariate overall survival in public PDAC cohorts (TCGA, Pan-Cancer Atlas, CPTAC, MSK, and discovery cohort)",
    3: "Stage-stratified mutation frequencies, pairwise and triple additive co-occurrence, and univariate Cox overall survival models",
    4: "Diagnosis-to-collection timing × combination interaction Cox models (primary cohort)",
    5: "Timing-interaction Cox outputs under dx_yr sensitivity filters (ALL, DX≥0, DX 0–5 years)",
    6: "Multivariable Cox coefficients and internal validation metrics by stage",
    7: "Three-group stage landscape, mutation rates, clinical associations, and synergy summaries",
    8: "Top-ranked synergistic dual-mutation combinations by stage",
    9: "Internal discovery–validation cohort comparison and cross-validation metrics",
    10: "Dual-mutation combination significance, FDR/Bonferroni summaries, and multiplicity correction",
    11: "Deceased-versus-living enrichment: cohort summary, mutation frequencies, synergy, and clinical factors",
}

# global (appearance) number -> type-specific number (from prior renumber_supplementary_by_appearance.py)
GLOBAL_TO_METHOD = {1: 6, 2: 7, 4: 8, 15: 9, 16: 5, 21: 2, 24: 3, 29: 1, 30: 4}
GLOBAL_TO_TABLE = {3: 2, 6: 1, 7: 3, 8: 4, 9: 5, 10: 6, 11: 7, 12: 8, 13: 9, 14: 10, 23: 11}
GLOBAL_TO_FIGURE = {5: 10, 17: 3, 18: 7, 19: 4, 20: 8, 22: 11, 25: 12, 26: 1, 27: 2, 28: 9, 31: 5, 32: 6}


def delete_paragraph(paragraph) -> None:
    el = paragraph._element
    el.getparent().remove(el)


def set_para_text(p, text: str) -> None:
    for r in p.runs:
        r.text = ""
    if p.runs:
        p.runs[0].text = text
    else:
        p.add_run(text)


def find_section_start(doc: Document, heading: str) -> int:
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip() == heading:
            return i
    raise RuntimeError(f"Section not found: {heading}")


def global_to_type_text(text: str) -> str:
    """Map global supplementary numbers back to type-specific 1..N."""

    def g2m(m: re.Match) -> str:
        g = int(m.group(1))
        return f"Supplementary Method {GLOBAL_TO_METHOD.get(g, g)}"

    def g2f(m: re.Match) -> str:
        g = int(m.group(1))
        return f"Supplementary Figure {GLOBAL_TO_FIGURE.get(g, g)}"

    def g2t(m: re.Match) -> str:
        g = int(m.group(1))
        return f"Supplementary Table {GLOBAL_TO_TABLE.get(g, g)}"

    # Ranges/lists (high numbers first within each replacement pass)
    def map_nums(s: str, mapping: dict[int, int]) -> str:
        parts = re.split(r"(\d+)", s)
        out = []
        for part in parts:
            if part.isdigit():
                n = int(part)
                out.append(str(mapping.get(n, n)))
            else:
                out.append(part)
        return "".join(out)

    text = re.sub(
        r"Supplementary Figures (\d+)[–-](\d+) and (\d+)",
        lambda m: (
            f"Supplementary Figures "
            f"{GLOBAL_TO_FIGURE.get(int(m.group(1)), int(m.group(1)))}–"
            f"{GLOBAL_TO_FIGURE.get(int(m.group(2)), int(m.group(2)))} and "
            f"{GLOBAL_TO_FIGURE.get(int(m.group(3)), int(m.group(3)))}"
        ),
        text,
    )
    text = re.sub(
        r"Supplementary Figures (\d+)[–-](\d+)",
        lambda m: (
            f"Supplementary Figures "
            f"{GLOBAL_TO_FIGURE.get(int(m.group(1)), int(m.group(1)))}–"
            f"{GLOBAL_TO_FIGURE.get(int(m.group(2)), int(m.group(2)))}"
        ),
        text,
    )
    text = re.sub(
        r"Supplementary Tables (\d+)[–-](\d+)",
        lambda m: (
            f"Supplementary Tables "
            f"{GLOBAL_TO_TABLE.get(int(m.group(1)), int(m.group(1)))}–"
            f"{GLOBAL_TO_TABLE.get(int(m.group(2)), int(m.group(2)))}"
        ),
        text,
    )
    text = re.sub(
        r"Supplementary Tables (\d+), (\d+), and (\d+)",
        lambda m: (
            f"Supplementary Tables "
            f"{GLOBAL_TO_TABLE.get(int(m.group(1)), int(m.group(1)))}, "
            f"{GLOBAL_TO_TABLE.get(int(m.group(2)), int(m.group(2)))}, and "
            f"{GLOBAL_TO_TABLE.get(int(m.group(3)), int(m.group(3)))}"
        ),
        text,
    )
    text = re.sub(
        r"Supplementary Tables (\d+) and (\d+)",
        lambda m: (
            f"Supplementary Tables "
            f"{GLOBAL_TO_TABLE.get(int(m.group(1)), int(m.group(1)))} and "
            f"{GLOBAL_TO_TABLE.get(int(m.group(2)), int(m.group(2)))}"
        ),
        text,
    )
    text = re.sub(
        r"Supplementary Table (\d+) and Supplementary Table (\d+)",
        lambda m: (
            f"Supplementary Table {GLOBAL_TO_TABLE.get(int(m.group(1)), int(m.group(1)))} "
            f"and Supplementary Table {GLOBAL_TO_TABLE.get(int(m.group(2)), int(m.group(2)))}"
        ),
        text,
    )

    # Single refs — replace high global IDs first to avoid partial overlaps
    for g in sorted(GLOBAL_TO_METHOD, reverse=True):
        text = re.sub(rf"\bSupplementary Method {g}\b", f"Supplementary Method {GLOBAL_TO_METHOD[g]}", text)
    for g in sorted(GLOBAL_TO_FIGURE, reverse=True):
        text = re.sub(rf"\bSupplementary Figure {g}\b", f"Supplementary Figure {GLOBAL_TO_FIGURE[g]}", text)
    for g in sorted(GLOBAL_TO_TABLE, reverse=True):
        text = re.sub(rf"\bSupplementary Table {g}\b", f"Supplementary Table {GLOBAL_TO_TABLE[g]}", text)

    return text


def insert_paragraph_after(paragraph, text: str, style: str | None = None) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._element.addnext(new_p)
    np = Paragraph(new_p, paragraph._parent)
    if style:
        np.style = style
    np.add_run(text)
    return np


def insert_heading_after(paragraph, text: str, level: int = 2) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._element.addnext(new_p)
    np = Paragraph(new_p, paragraph._parent)
    np.style = f"Heading {level}"
    run = np.add_run(text)
    run.bold = True
    run.font.size = Pt(12)
    return np


def remove_duplicate_tail_inventory(doc: Document) -> None:
    """Delete a second Methods/Figures/Tables block after back matter (Ethics, etc.)."""
    first_methods = None
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip() == "Supplementary Methods":
            if first_methods is None:
                first_methods = i
                continue
            # Second block: delete from this heading through end
            while i < len(doc.paragraphs):
                delete_paragraph(doc.paragraphs[i])
            return


def rebuild_end_inventory(doc: Document) -> None:
    start = find_section_start(doc, "Supplementary Materials")
    remove_duplicate_tail_inventory(doc)

    # Remove misplaced intro paragraphs not directly after title
    while True:
        removed = False
        for i, p in enumerate(doc.paragraphs):
            if p.text.strip().startswith(
                "The following supplementary items accompany the main manuscript"
            ) and i != start + 1:
                delete_paragraph(p)
                removed = True
                break
        if not removed:
            break

    ack = find_section_start(doc, "Acknowledgments")
    while start + 1 < ack:
        delete_paragraph(doc.paragraphs[start + 1])
        ack = find_section_start(doc, "Acknowledgments")

    anchor = doc.paragraphs[start]
    anchor = insert_paragraph_after(
        anchor,
        "The following supplementary items accompany the main manuscript. "
        "Extended methods, figures, and tabular data are provided as separate "
        "supplementary files (Supplementary Methods.docx; Supplementary Figures.docx; "
        "Supplementary Tables.xlsx).",
    )

    anchor = insert_heading_after(anchor, "Supplementary Methods", 2)
    for n, title in METHOD_TITLES.items():
        anchor = insert_paragraph_after(anchor, f"Supplementary Method {n}. {title}", style="List Number")

    anchor = insert_paragraph_after(anchor, "")
    anchor = insert_heading_after(anchor, "Supplementary Figures", 2)
    for n, title in FIGURE_TITLES.items():
        anchor = insert_paragraph_after(anchor, f"Supplementary Figure {n}. {title}", style="List Number")

    anchor = insert_paragraph_after(anchor, "")
    anchor = insert_heading_after(anchor, "Supplementary Tables", 2)
    for n, title in TABLE_TITLES.items():
        anchor = insert_paragraph_after(anchor, f"Supplementary Table {n}. {title}", style="List Number")


def fix_inventory_placement_only() -> None:
    doc = Document(str(MANUSCRIPT))
    rebuild_end_inventory(doc)
    doc.save(str(MANUSCRIPT))
    print("Rebuilt Supplementary Materials inventory in correct location.")


def update_manuscript() -> None:
    doc = Document(str(MANUSCRIPT))
    supp_start = find_section_start(doc, "Supplementary Materials")
    ack = find_section_start(doc, "Acknowledgments")
    n = 0
    for i, p in enumerate(doc.paragraphs):
        if i >= ack:
            break
        if supp_start <= i < ack:
            continue
        new = global_to_type_text(p.text)
        if new != p.text:
            set_para_text(p, new)
            n += 1
    rebuild_end_inventory(doc)
    doc.save(str(MANUSCRIPT))
    print(f"Manuscript: {n} body paragraphs updated; end inventory rebuilt (Methods/Figures/Tables 1..N).")


def update_supplementary_methods() -> None:
    from build_supplementary_materials import build_supplementary_methods

    path = SUPP_DIR / "Supplementary Methods.docx"
    tmp = Path(tempfile.gettempdir()) / "supp_methods_base.docx"
    build_supplementary_methods(TUMOR, tmp)
    doc = Document(str(tmp))
    section_re = re.compile(r"^(\d+)\.\s+")
    doc2 = Document()
    title = doc2.add_paragraph("Supplementary Methods")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.bold = True
    doc2.add_paragraph()

    current_n: int | None = None
    body: list[str] = []
    for p in doc.paragraphs:
        m = section_re.match(p.text.strip())
        if m and p.style and "Heading" in (p.style.name or ""):
            if current_n is not None:
                h = doc2.add_paragraph(f"Supplementary Method {current_n}. {METHOD_TITLES[current_n]}")
                for run in h.runs:
                    run.bold = True
                for line in body:
                    doc2.add_paragraph(line)
                doc2.add_paragraph()
            current_n = int(m.group(1))
            body = []
        elif current_n is not None and p.text.strip():
            body.append(p.text)
    if current_n is not None:
        h = doc2.add_paragraph(f"Supplementary Method {current_n}. {METHOD_TITLES[current_n]}")
        for run in h.runs:
            run.bold = True
        for line in body:
            doc2.add_paragraph(line)

    doc2.save(str(path))
    tmp.unlink(missing_ok=True)
    print(f"Updated {path.name}")


def update_supplementary_figures() -> None:
    path = SUPP_DIR / "Supplementary Figures.docx"
    doc = Document(str(path))
    fig_re = re.compile(r"^Supplementary Figure (\d+)\.")
    body = doc.element.body
    blocks: list[tuple[int, list]] = []
    title_children: list = []
    current_old: int | None = None
    current_block: list = []

    for child in list(body):
        tag = child.tag.split("}")[-1]
        if tag == "p":
            p = Paragraph(child, doc)
            m = fig_re.match(p.text.strip())
            if m:
                if current_old is not None:
                    blocks.append((current_old, current_block))
                current_old = int(m.group(1))
                current_block = [child]
            elif current_block:
                current_block.append(child)
            else:
                title_children.append(child)
        elif current_block:
            current_block.append(child)
        else:
            title_children.append(child)
    if current_old is not None:
        blocks.append((current_old, current_block))

    # Map global figure number -> type number if needed
    def to_type(n: int) -> int:
        return GLOBAL_TO_FIGURE.get(n, n)

    for child in list(body):
        body.remove(child)
    for child in title_children:
        body.append(child)

    for old_g, block in sorted(blocks, key=lambda x: to_type(x[0])):
        n = to_type(old_g)
        first = Paragraph(block[0], doc)
        set_para_text(first, f"Supplementary Figure {n}. {FIGURE_TITLES[n]}")
        for child in block:
            body.append(child)

    doc.save(str(path))
    print(f"Updated {path.name}")


def update_supplementary_tables() -> None:
    path = SUPP_DIR / "Supplementary Tables.xlsx"
    xl = pd.ExcelFile(path)
    rows = []
    for old in range(1, 12):
        sheets = [s for s in xl.sheet_names if re.match(rf"^S{old}_", s)]
        rows.append(
            {
                "Supplementary_Table": f"Supplementary Table {old}",
                "Legacy_Sheet_Prefix": f"S{old}_",
                "Sheets": ", ".join(sheets),
                "Description": TABLE_TITLES[old],
            }
        )
    idx_df = pd.DataFrame(rows)
    from openpyxl import load_workbook

    wb = load_workbook(path)
    if "Table_Index" in wb.sheetnames:
        del wb["Table_Index"]
    ws = wb.create_sheet("Table_Index", 0)
    for r_idx, row in enumerate([idx_df.columns.tolist()] + idx_df.values.tolist(), 1):
        for c_idx, val in enumerate(row, 1):
            ws.cell(row=r_idx, column=c_idx, value=val)
    wb.save(path)
    print(f"Updated Table_Index in {path.name}")


def main() -> None:
    update_manuscript()
    update_supplementary_methods()
    update_supplementary_figures()
    update_supplementary_tables()


if __name__ == "__main__":
    import sys

    if len(sys.argv) > 1 and sys.argv[1] == "fix-inventory":
        fix_inventory_placement_only()
    else:
        main()
