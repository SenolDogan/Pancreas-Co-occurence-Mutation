#!/usr/bin/env python3
"""
Restore and complete in-text supplementary citations in NewManuscript.docx.
Uses type-specific numbering: Methods 1–9, Figures 1–12, Tables 1–11.
Only modifies citation text in the body (not main scientific prose).
"""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document

MANUSCRIPT = Path(__file__).resolve().parent / "NewManuscript.docx"

# (unique substring in paragraph, citation fragment to ensure at end)
# Fragments use final Supplementary Method/Figure/Table numbering.
CITATION_RULES: list[tuple[str, str]] = [
    # Methods 1–4 (previously missing in manuscript)
    (
        "Statistical significance was assessed using chi-square",
        " (Supplementary Method 1).",
    ),
    (
        "The master analytic cohort comprised 2330 patients",
        " (Supplementary Method 2; Supplementary Tables 3, 7, and 11).",
    ),
    (
        "Figure 3 summarizes stage-stratified Cox results",
        " (Supplementary Method 4; Supplementary Table 6; Supplementary Figure 12).",
    ),
    (
        "Figure 2. Integrated cohort and three-group stage landscape",
        " (Supplementary Table 7).",
    ),
    # Methods 5–9, validation, software
    (
        "Internal validation of descriptive synergy",
        " (Supplementary Method 6; Supplementary Figure 9; Supplementary Table 9).",
    ),
    (
        "Using the public cBioPortal API",
        " (Supplementary Method 7; Supplementary Table 2).",
    ),
    (
        "All analyses were performed using Python 3.x",
        " (Supplementary Method 9; Supplementary Tables 1–10).",
    ),
    (
        "DX2COLLECTION YEAR measures the elapsed time",
        " (Supplementary Method 5; Supplementary Table 1; Supplementary Table 5; "
        "Supplementary Figures 3–4 and Supplementary Figure 7).",
    ),
    (
        "For the high-dimensional screen across dual mutation",
        " (Supplementary Figure 8; Supplementary Table 10).",
    ),
    # Results — tables
    (
        "Table 1. Timing interaction highlights",
        " (Supplementary Table 1).",
    ),
    (
        "Stable q<0.10 = interaction FDR",
        " (Supplementary Table 1).",
    ),
    (
        "Figure 1 summarizes a cross-sectional dead-versus-alive",
        " (Supplementary Figure 11; Supplementary Table 11).",
    ),
    (
        "As an enrichment readout, deceased versus living",
        " (Supplementary Tables 10 and 11; Supplementary Figure 11; Supplementary Method 3).",
    ),
    (
        "Across paad_tcga_gdc, paad_tcga_pan_can_atlas",
        " (Supplementary Table 2; Supplementary Method 7).",
    ),
    (
        "Table 2. Genetic mutation frequencies",
        " (Supplementary Tables 3 and 11).",
    ),
    (
        "Table 3. Statistically significant dual mutation",
        " (Supplementary Table 10).",
    ),
    (
        "Table 4. Top synergistic mutation",
        " (Supplementary Tables 7 and 8).",
    ),
    (
        "Table 5. Triple mutation combinations",
        " (Supplementary Tables 3 and 8).",
    ),
  # Figures 1–7 pairwise / timing
    (
        "Beyond stage-stratified hazard summaries, we tested whether the association",
        " (Supplementary Method 5; Supplementary Table 1 and Supplementary Table 5; "
        "Supplementary Table 4; Supplementary Figures 3–4 and Supplementary Figure 7).",
    ),
    (
        "particularly robust signal was observed in the metastatic subset",
        " (see Supplementary Table 1 for the compact comparison).",
    ),
    (
        "To summarize the pairwise screen in a reviewer-friendly way",
        " (triple-level panels: Supplementary Figures 1–2; timing volcanoes: "
        "Supplementary Figures 3–4 and Supplementary Figure 7).",
    ),
    (
        "detailed TP53+KRAS(+X) stage comparison is shown in Figure 5",
        " (exploratory short-OS sensitivity: Supplementary Figure 5).",
    ),
    (
        "focused TP53/KRAS stage panel (Figure 5)",
        " (Supplementary Figure 5).",
    ),
    (
        "Contextual functional mapping was performed through pathway",
        " (Supplementary Method 8; Supplementary Figure 10; exploratory panels in Supplementary Figures 5 and 6).",
    ),
    (
        "Stratified 70/30 splits and cross-validation",
        " (Supplementary Figure 9; Supplementary Table 9).",
    ),
    (
        "Pathway-oriented summaries grouped recurrent",
        " (Supplementary Figure 10).",
    ),
    (
        "External validity and cohort independence (cBioPortal)",
        " (Supplementary Table 2).",
    ),
    # Abstract — external validation
    (
        "External replication in four public PDAC cohorts",
        " (Supplementary Method 7; Supplementary Table 2).",
    ),
]

# Patterns to strip before re-appending (remove broken/duplicate citation tails)
STRIP_PATTERNS = [
    r"\s*\(Supplementary Method[^)]*\)",
    r"\s*\(Supplementary Table[^)]*\)",
    r"\s*\(Supplementary Figure[^)]*\)",
    r"\s*\(Supplementary Tables[^)]*\)",
    r"\s*\(Supplementary Figures[^)]*\)",
    r"\s*\(triple-level panels:[^)]*\)",
    r"\s*\(archived overview[^)]*\)",
    r"\s*\(see Supplementary Table \d+ for the compact comparison\)",
    r"\s*\(exploratory short-OS sensitivity:[^)]*\)",
    r"\s*\(exploratory panels in Supplementary Figures[^)]*\)",
    r"\s*External validation: Supplementary Method[^.]*\.",
    r"\s*\(numbered as Supplementary Table[^)]*\)",
]


def set_para_text(p, text: str) -> None:
    for r in p.runs:
        r.text = ""
    if p.runs:
        p.runs[0].text = text
    else:
        p.add_run(text)


def strip_supplementary_citations(text: str) -> str:
    out = text
    for _ in range(4):
        prev = out
        for pat in STRIP_PATTERNS:
            out = re.sub(pat, "", out, flags=re.IGNORECASE)
        if out == prev:
            break
    out = re.sub(r"\s{2,}", " ", out).strip()
    return out


def citation_key(fragment: str) -> str:
    return fragment.strip().lstrip("(").split(";")[0][:50]


def ensure_citation(text: str, fragment: str) -> str:
    key = citation_key(fragment)
    if key in text:
        return text
    frag = fragment if fragment.startswith(" ") or fragment.startswith("(") else f" {fragment}"
    if not frag.startswith("(") and not text.endswith("."):
        return text.rstrip() + frag
    return text.rstrip() + frag


def find_supplementary_materials_idx(doc: Document) -> int:
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip() == "Supplementary Materials" and i > 50:
            return i
    raise RuntimeError("Supplementary Materials section not found")


def main() -> None:
    doc = Document(str(MANUSCRIPT))
    supp_idx = find_supplementary_materials_idx(doc)
    n_strip = 0
    n_add = 0

    for i in range(supp_idx):
        p = doc.paragraphs[i]
        original = p.text
        cleaned = strip_supplementary_citations(original)
        if cleaned != original:
            n_strip += 1

        updated = cleaned
        for needle, frag in CITATION_RULES:
            if needle in updated:
                before = updated
                updated = ensure_citation(updated, frag)
                if updated != before:
                    n_add += 1

        # Remove redundant duplicate parenthetical when already named in prose
        updated = re.sub(
            r"(Supplementary Table 1)\.\s*\(Supplementary Table 1\)",
            r"\1",
            updated,
        )
        updated = re.sub(
            r"(Supplementary Figure 5)\)\.\s*\(Supplementary Figure 5\)",
            r"\1).",
            updated,
        )
        updated = re.sub(r"\.{2,}", ".", updated)
        updated = re.sub(r"\s+\.", ".", updated)
        if updated != original:
            set_para_text(p, updated)

    doc.save(str(MANUSCRIPT))
    print(f"Updated citations in {MANUSCRIPT.name}: stripped {n_strip} paragraphs, added/confirmed {n_add} fragments.")

    # Report coverage
    methods: set[int] = set()
    figures: set[int] = set()
    tables: set[int] = set()
    for i in range(supp_idx):
        t = doc.paragraphs[i].text
        methods.update(int(x) for x in re.findall(r"Supplementary Method (\d+)", t))
        figures.update(int(x) for x in re.findall(r"Supplementary Figure (\d+)", t))
        tables.update(int(x) for x in re.findall(r"Supplementary Table (\d+)", t))

    print("Cited Methods:", sorted(methods), f"({len(methods)}/9)")
    print("Cited Figures:", sorted(figures), f"({len(figures)}/12)")
    print("Cited Tables:", sorted(tables), f"({len(tables)}/11)")
    missing_m = set(range(1, 10)) - methods
    missing_f = set(range(1, 13)) - figures
    missing_t = set(range(1, 12)) - tables
    if missing_m:
        print("  Missing Methods:", sorted(missing_m))
    if missing_f:
        print("  Missing Figures:", sorted(missing_f))
    if missing_t:
        print("  Missing Tables:", sorted(missing_t))


if __name__ == "__main__":
    main()
