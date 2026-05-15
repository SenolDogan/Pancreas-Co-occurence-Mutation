#!/usr/bin/env python3
"""Fix [18]–[31] placement so first-appearance order is ascending."""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document

MANUSCRIPT = Path(__file__).resolve().parent / "NewManuscript.docx"


def set_para_text(p, text: str) -> None:
    for r in p.runs:
        r.text = ""
    if p.runs:
        p.runs[0].text = text
    else:
        p.add_run(text)


def strip_ref(text: str, n: int) -> str:
    text = re.sub(rf"\s*\[{n}\]", "", text)
    text = re.sub(rf"\[{n},\s*", "[", text)
    text = re.sub(rf",\s*{n}\]", "]", text)
    return re.sub(r"\s{2,}", " ", text).strip()


def main() -> None:
    doc = Document(str(MANUSCRIPT))

  # Remove misplaced [20],[21] from early Methods headings
    t52 = doc.paragraphs[52].text.replace(" [20]", "").strip()
    set_para_text(doc.paragraphs[52], t52)

    t53 = doc.paragraphs[53].text
    t53 = re.sub(r"\s*\[21\]", "", t53)
    set_para_text(doc.paragraphs[53], t53.strip())

    # Multiplicity para: [18] [19] -> [19]; append [20,21]
    t61 = doc.paragraphs[61].text
    t61 = strip_ref(t61, 18)
    if "[20]" not in t61:
        t61 = re.sub(
            r"(\[19\])\s*(\(Supplementary)",
            r"\1 [20,21] \2",
            t61,
        )
        if "[20,21]" not in t61:
            t61 = t61.replace("[19]", "[19,20,21]", 1)
    set_para_text(doc.paragraphs[61], t61)

    # Statistical Analysis: add [18]
    t41 = doc.paragraphs[41].text
    if "[18]" not in t41:
        t41 = t41.rstrip()
        if t41.endswith("(Supplementary Method 1)."):
            t41 = t41.replace(
                "(Supplementary Method 1).",
                "[18] (Supplementary Method 1).",
            )
        else:
            t41 += " [18]"
    set_para_text(doc.paragraphs[41], t41)

    # External discussion: [25] after MSK, [26] after CPTAC (reading order)
    t128 = doc.paragraphs[128].text
    t128 = re.sub(r"\s*\[25\]", "", t128)
    if "[25]" not in t128:
        t128 = t128.replace(
            "despite platform and annotation differences.",
            "despite platform and annotation differences. [25]",
        )
    if "[26]" not in t128:
        t128 = t128.replace(
            "under the simplified external Cox specification.",
            "under the simplified external Cox specification. [26]",
        )
    set_para_text(doc.paragraphs[128], t128)

    # Discussion order: [27] before [28,29]
    t129 = doc.paragraphs[129].text
    t130 = doc.paragraphs[130].text
    t129 = re.sub(r"\s*\[28\]\s*\[29\]", "", t129)
    t130 = re.sub(r"\s*\[27\]", "", t130)
    if "[27]" not in t129:
        t129 = t129.rstrip() + " [27]"
    if "[28,29]" not in t130 and "[28]" not in t130:
        t130 = t130.rstrip() + " [28,29]"
    set_para_text(doc.paragraphs[129], t129)
    set_para_text(doc.paragraphs[130], t130)

    doc.save(str(MANUSCRIPT))

    # Verify
    doc = Document(str(MANUSCRIPT))
    refs_idx = next(i for i, p in enumerate(doc.paragraphs) if p.text.strip() == "References")
    first: dict[int, int] = {}
    for i in range(refs_idx):
        t = doc.paragraphs[i].text
        for m in re.finditer(r"\[(\d+)(?:[–-](\d+))?\]", t):
            a, b = int(m.group(1)), int(m.group(2) or m.group(1))
            for n in range(a, b + 1):
                first.setdefault(n, i)
        for m in re.finditer(r"\[(\d+),\s*(\d+)\]", t):
            for n in range(int(m.group(1)), int(m.group(2)) + 1):
                first.setdefault(n, i)

    print("First appearance:")
    for n in range(1, 32):
        print(f"  [{n:2d}] para {first.get(n, '—')}")

    ok = True
    prev_p, prev_n = -1, 0
    for n in range(1, 32):
        if n not in first:
            print(f"MISSING [{n}]")
            ok = False
            continue
        if first[n] < prev_p or (first[n] == prev_p and n < prev_n):
            pass  # same para multiple refs OK if ascending n
        if n > 1 and first[n] < prev_p:
            print(f"ORDER FAIL: [{n}] para {first[n]} before [{n-1}] para {prev_p}")
            ok = False
        prev_p, prev_n = first[n], n
    if ok:
        print("OK: ascending order")


if __name__ == "__main__":
    main()
