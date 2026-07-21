#!/usr/bin/env python3
"""
Remove main-text Figure 8 (schematic) and Figure 9 (KM) from Nat Commun boost,
and embed them after Supplementary Figure 12 as Supplementary Figures 13–14.

Requirements:
  - Tumor/Figure_Decision_Schematic_ThreeLayers.png
  - Tumor/KM_TP53_KRAS_SuppFigure14_AB.png (from nat_commun_boost_analyses.py)

Recommended order:
  MPLBACKEND=Agg python3 nat_commun_boost_analyses.py
  python3 move_nat_commun_figures_to_supplementary.py
"""

from __future__ import annotations

import argparse
import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.shared import Inches
from docx.text.paragraph import Paragraph


TUMOR = Path(__file__).resolve().parent
MANUSCRIPT = TUMOR / "NewManuscript.docx"
SCHEMATIC = TUMOR / "Figure_Decision_Schematic_ThreeLayers.png"
KM_COMBINED = TUMOR / "KM_TP53_KRAS_SuppFigure14_AB.png"


def delete_paragraph(paragraph: Paragraph) -> None:
    paragraph._element.getparent().remove(paragraph._element)


def paragraph_index(doc: Document, predicate) -> int | None:
    for i, p in enumerate(doc.paragraphs):
        if predicate(p):
            return i
    return None


def add_paragraph_after(anchor: Paragraph, text: str = "") -> Paragraph:
    new_p = OxmlElement("w:p")
    anchor._element.addnext(new_p)
    np = Paragraph(new_p, anchor._parent)
    if text:
        np.add_run(text)
    return np


def insert_centered_picture_after(anchor: Paragraph, png: Path, width_in: float) -> Paragraph:
    new_p = OxmlElement("w:p")
    anchor._element.addnext(new_p)
    np = Paragraph(new_p, anchor._parent)
    np.alignment = WD_ALIGN_PARAGRAPH.CENTER
    np.add_run().add_picture(str(png), width=Inches(width_in))
    return np


def para_has_picture(p: Paragraph) -> bool:
    return "pic:pic" in p._element.xml or "w:drawing" in p._element.xml


def collapse_extra_blanks_between_indices(doc: Document, kw_idx: int, intro_idx: int) -> None:
    """Keep at most one empty paragraph between Keywords (kw_idx) and Introduction (intro_idx)."""
    blank_indices: list[int] = []
    for i in range(kw_idx + 1, intro_idx):
        p = doc.paragraphs[i]
        if not p.text.strip() and not para_has_picture(p):
            blank_indices.append(i)
    for bi in reversed(blank_indices[1:]):
        delete_paragraph(doc.paragraphs[bi])


def remove_main_boost_fig_blocks(doc: Document) -> None:
    i8 = paragraph_index(
        doc, lambda p: p.text.strip().startswith("Figure 8. Analytical decision schematic")
    )
    intro_i = paragraph_index(doc, lambda p: p.text.strip().startswith("Introduction"))
    if i8 is not None and intro_i is not None and i8 < intro_i:
        for j in range(intro_i - 1, i8 - 1, -1):
            delete_paragraph(doc.paragraphs[j])
        kw_i = paragraph_index(doc, lambda p: p.text.strip().startswith("Keywords:"))
        intro_i2 = paragraph_index(doc, lambda p: p.text.strip().startswith("Introduction"))
        if kw_i is not None and intro_i2 is not None and kw_i < intro_i2:
            collapse_extra_blanks_between_indices(doc, kw_i, intro_i2)

    i9 = paragraph_index(doc, lambda p: p.text.strip().startswith("Figure 9. Kaplan"))
    i_ext = paragraph_index(
        doc, lambda p: p.text.strip().startswith("External replication in public cohorts")
    )
    if i9 is not None and i_ext is not None and i9 < i_ext:
        for j in range(i_ext - 1, i9 - 1, -1):
            delete_paragraph(doc.paragraphs[j])


def logrank_pa_pb_from_feasibility() -> tuple[str, str]:
    """Parse NatCommun_Feasibility_Report.md for panel (A)/(B) log-rank p-values."""
    md = TUMOR / "NatCommun_Feasibility_Report.md"
    if not md.is_file():
        return "", ""
    text = md.read_text()
    pa = re.search(r"Panel \*\*\(A\)\*\*:\s*log-rank p=([\d.eE+-]+)", text)
    pb = re.search(r"Panel \*\*\(B\)\*\*:\s*log-rank p=([\d.eE+-]+)", text)
    return (pa.group(1) if pa else "", pb.group(1) if pb else "")


def delete_supplementary_figures_13_14(doc: Document) -> None:
    """Remove SF13 caption, schematic image, SF14 caption, KM image."""
    ix = paragraph_index(doc, lambda p: p.text.strip().startswith("Supplementary Figure 13."))
    if ix is None:
        return
    while ix < len(doc.paragraphs):
        t = doc.paragraphs[ix].text.strip()
        if t.startswith("Supplementary Tables"):
            break
        delete_paragraph(doc.paragraphs[ix])


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument(
        "--overwrite",
        action="store_true",
        help="Remove existing Supplementary Figures 13–14 block and recreate",
    )
    args = parser.parse_args()

    if not MANUSCRIPT.is_file():
        print(f"Missing {MANUSCRIPT}")
        return 1
    if not SCHEMATIC.is_file():
        print(f"Missing {SCHEMATIC}")
        return 1
    if not KM_COMBINED.is_file():
        print(f"Missing {KM_COMBINED}; run: MPLBACKEND=Agg python3 nat_commun_boost_analyses.py")
        return 1

    doc = Document(str(MANUSCRIPT))
    has_sf13 = any("Supplementary Figure 13." in p.text for p in doc.paragraphs)

    if has_sf13 and args.overwrite:
        delete_supplementary_figures_13_14(doc)
        has_sf13 = False

    if has_sf13 and not args.overwrite:
        print("Supplementary Figure 13 already present; skipping. Use --overwrite to recreate.")
        return 0

    if any("Figure 8. Analytical decision schematic" in p.text for p in doc.paragraphs):
        remove_main_boost_fig_blocks(doc)

    sf12_i = paragraph_index(
        doc, lambda p: p.text.strip().startswith("Supplementary Figure 12.")
    )
    if sf12_i is None:
        print("Could not find Supplementary Figure 12.")
        return 1
    anchor = doc.paragraphs[sf12_i]

    p13 = add_paragraph_after(
        anchor,
        "Supplementary Figure 13. Three-layer analytical decision schematic for the discovery cohort. "
        "The workflow links (i) cross-sectional enrichment screens (deceased versus living), "
        "(ii) stage-stratified Cox overall survival models, and (iii) diagnosis-to-biosample timing "
        "interactions (DX2COLLECTION YEAR). Public cBioPortal cohorts were used for univariate "
        "replication of prespecified dual-mutation pairs only; external timing matching is not supported.",
    )
    img13 = insert_centered_picture_after(p13, SCHEMATIC, 6.4)

    pa, pb = logrank_pa_pb_from_feasibility()
    caption14 = (
        "Supplementary Figure 14. Kaplan–Meier overall survival for TP53+KRAS versus all other patients "
        "in the discovery cohort (OS_MONTHS>0). Panel (A) all prespecified stage groups; panel (B) metastatic stratum. "
        "Log-rank p-values are printed on each panel (uncorrected two-group test; interpret together with "
        "stage-stratified Cox models in the main text)."
    )
    if pa and pb:
        caption14 += f" Numeric log-rank p-values (see NatCommun_Feasibility_Report.md): (A) p={pa}; (B) p={pb}."

    p14 = add_paragraph_after(img13, caption14)
    insert_centered_picture_after(p14, KM_COMBINED, 6.8)

    doc.save(str(MANUSCRIPT))
    print(f"Updated {MANUSCRIPT}: Supplementary Figures 13–14 added; main-text Fig 8/9 removed if present.")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
