#!/usr/bin/env python3
"""
Unify supplementary numbering in NewManuscript.docx only:
  - Supplementary Method 1–9
  - Supplementary Figure 1–12 (was Figure S1…)
  - Supplementary Table 1–11 (was Table S1…)
  - End inventory: single sequential list 1–32 under Supplementary Materials
"""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


METHODS = [
    "Supplementary Method 1. Scope and relationship to the main manuscript",
    "Supplementary Method 2. Primary analytic cohort (discovery)",
    "Supplementary Method 3. Synergy and enrichment metrics (multiplicative, additive, and protective scores)",
    "Supplementary Method 4. Stage-stratified Cox and multivariable extensions",
    "Supplementary Method 5. DX2COLLECTION_YEAR × mutation-combination interaction models and sensitivity filters",
    "Supplementary Method 6. Internal validation (stratified split and cross-validation)",
    "Supplementary Method 7. External validation in public PDAC cohorts (cBioPortal)",
    "Supplementary Method 8. Overview of supplementary figure content",
    "Supplementary Method 9. Software and computational environment",
]

FIGURES = [
    "Supplementary Figure 1. Triple-mutation additive deviations from independence (top positive and negative bars by stage)",
    "Supplementary Figure 2. Univariate Cox forest plot for top triple-mutation combinations by stage",
    "Supplementary Figure 3. Diagnosis-to-collection timing × pairwise combination interaction volcano (DX2COLLECTION_YEAR ≥ 0)",
    "Supplementary Figure 4. Diagnosis-to-collection timing × triple-mutation interaction volcano (DX2COLLECTION_YEAR ≥ 0)",
    "Supplementary Figure 5. TP53/KRAS patterns in short overall-survival subsets (exploratory sensitivity)",
    "Supplementary Figure 6. Diabetes-stratified exploratory gene analysis",
    "Supplementary Figure 7. Triple-mutation timing interaction volcano under DX2COLLECTION_YEAR 0–5 years sensitivity filter",
    "Supplementary Figure 8. Multiple-comparison correction diagnostics for combination screens",
    "Supplementary Figure 9. Internal discovery–validation cohort diagnostics",
    "Supplementary Figure 10. Pathway-oriented functional validation summary",
    "Supplementary Figure 11. Deceased-versus-living comprehensive mutation enrichment overview",
    "Supplementary Figure 12. Multivariable Cox internal validation summary by stage",
]

TABLES = [
    "Supplementary Table 1. Timing-interaction sensitivity comparison for DX2COLLECTION_YEAR × mutation combinations (adjusted Cox; ALL, DX≥0, and 0–5 year filters)",
    "Supplementary Table 2. External replication of TP53+KRAS co-mutation prevalence and univariate overall survival in public PDAC cohorts (TCGA, Pan-Cancer Atlas, CPTAC, MSK, and discovery cohort)",
    "Supplementary Table 3. Stage-stratified mutation frequencies, pairwise and triple additive co-occurrence, and univariate Cox overall survival models",
    "Supplementary Table 4. Diagnosis-to-collection timing × combination interaction Cox models (primary cohort)",
    "Supplementary Table 5. Timing-interaction Cox outputs under dx_yr sensitivity filters (ALL, DX≥0, DX 0–5 years)",
    "Supplementary Table 6. Multivariable Cox coefficients and internal validation metrics by stage",
    "Supplementary Table 7. Three-group stage landscape, mutation rates, clinical associations, and synergy summaries",
    "Supplementary Table 8. Top-ranked synergistic dual-mutation combinations by stage",
    "Supplementary Table 9. Internal discovery–validation cohort comparison and cross-validation metrics",
    "Supplementary Table 10. Dual-mutation combination significance, FDR/Bonferroni summaries, and multiplicity correction",
    "Supplementary Table 11. Deceased-versus-living enrichment: cohort summary, mutation frequencies, synergy, and clinical factors",
]


def delete_paragraph(paragraph) -> None:
    el = paragraph._element
    el.getparent().remove(el)


def set_para_text(p, text: str) -> None:
    for r in p.runs:
        r.text = ""
    if p.runs:
        p.runs[0].text = text
    else:
        p.add_run(text)


def normalize_supplementary_citations(text: str) -> str:
    """Standardize S-suffix supplementary IDs; do not alter main Table 1–5 / Figure 1–7 captions."""
    if not text.strip():
        return text

    # Protect main-text table/figure captions (e.g. "Table 1. Timing", "Figure 3.")
    placeholders: dict[str, str] = {}

    def protect(m: re.Match) -> str:
        key = f"__MAIN_{len(placeholders)}__"
        placeholders[key] = m.group(0)
        return key

    text = re.sub(
        r"\b(?:Table|Figure) ([1-7])\.",
        protect,
        text,
    )

    text = re.sub(r"Supplementary Tables S(\d+)", r"Supplementary Tables \1", text)
    text = re.sub(r"Supplementary Table S(\d+)", r"Supplementary Table \1", text)
    text = re.sub(r"Supplementary Figures S(\d+)", r"Supplementary Figures \1", text)
    text = re.sub(r"Supplementary Figure S(\d+)", r"Supplementary Figure \1", text)
    text = re.sub(r"Figure S(\d+)", r"Supplementary Figure \1", text)
    text = re.sub(r"Table S(\d+)", r"Supplementary Table \1", text)

    text = re.sub(
        r"(Supplementary (?:Figures|Tables) \d+)–S(\d+)",
        r"\1–\2",
        text,
    )
    text = re.sub(
        r"(Supplementary (?:Figures|Tables) [\d–]+) and S(\d+)",
        r"\1 and \2",
        text,
    )
    text = re.sub(
        r"Supplementary Tables (\d+)–(\d+)",
        lambda m: f"Supplementary Tables {m.group(1)}–{m.group(2)}",
        text,
    )

    for key, val in placeholders.items():
        text = text.replace(key, val)

    return text


def find_supplementary_inventory_start(doc: Document) -> int:
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip() == "Supplementary Methods" and i > 100:
            return i
    for i, p in enumerate(doc.paragraphs):
        t = p.text.strip()
        if t.startswith("Supplementary Method 1.") and i > 100:
            return i
    raise RuntimeError("Could not find supplementary inventory start")


def find_acknowledgments_start(doc: Document) -> int:
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip() == "Acknowledgments":
            return i
    raise RuntimeError("Could not find Acknowledgments section")


def insert_paragraph_after(paragraph, text: str):
    from docx.oxml import OxmlElement
    from docx.text.paragraph import Paragraph

    new_p = OxmlElement("w:p")
    paragraph._element.addnext(new_p)
    np = Paragraph(new_p, paragraph._parent)
    np.add_run(text)
    return np


def rebuild_inventory(doc: Document, intro_idx: int) -> None:
    start = find_supplementary_inventory_start(doc)
    # Delete from old subheading through paragraph before Acknowledgments
    while True:
        paras = doc.paragraphs
        if start >= len(paras):
            break
        t = paras[start].text.strip()
        if t == "Acknowledgments":
            break
        delete_paragraph(paras[start])

    items: list[str] = []
    for n, line in enumerate(METHODS, start=1):
        items.append(f"{n}. {line}")
    for n, line in enumerate(FIGURES, start=10):
        items.append(f"{n}. {line}")
    for n, line in enumerate(TABLES, start=22):
        items.append(f"{n}. {line}")

    intro_p = doc.paragraphs[intro_idx]
    for item in reversed(items):
        insert_paragraph_after(intro_p, item)


def main() -> None:
    path = Path(__file__).resolve().parent / "NewManuscript.docx"
    doc = Document(str(path))

    intro_idx = None
    for i, p in enumerate(doc.paragraphs):
        if "supplementary items accompany" in p.text.lower():
            intro_idx = i
            break

    ack_idx = find_acknowledgments_start(doc)

    # Normalize citations in body only (before supplementary inventory / acknowledgments)
    end_body = ack_idx
    n_norm = 0
    for i, p in enumerate(doc.paragraphs):
        if i >= end_body:
            break
        new = normalize_supplementary_citations(p.text)
        if new != p.text:
            set_para_text(p, new)
            n_norm += 1

    # Update intro
    if intro_idx is not None:
        set_para_text(
            doc.paragraphs[intro_idx],
            "The following supplementary items accompany the main manuscript and are numbered sequentially "
            "(items 1–32): Supplementary Methods (items 1–9), Supplementary Figures (items 10–21), and "
            "Supplementary Tables (items 22–32). Extended content is provided in separate files "
            "(Supplementary Methods.docx; Supplementary Figures.docx; Supplementary Tables.xlsx).",
        )

    rebuild_inventory(doc, intro_idx if intro_idx is not None else find_supplementary_inventory_start(doc) - 2)

    # Normalize data availability line if present
    for p in doc.paragraphs:
        if "Supplementary Methods, Supplementary Figures" in p.text:
            t = normalize_supplementary_citations(p.text)
            set_para_text(p, t)

    doc.save(str(path))
    print(f"Normalized {n_norm} body paragraphs; rebuilt inventory 1–32 in {path}")


if __name__ == "__main__":
    main()
