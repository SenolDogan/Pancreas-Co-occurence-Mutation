#!/usr/bin/env python3
"""Apply pre-submission citation/format fixes to NewManuscript.docx."""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph

MANUSCRIPT = Path(__file__).resolve().parent / "NewManuscript.docx"


def set_para_text(p, text: str) -> None:
    for r in p.runs:
        r.text = ""
    if p.runs:
        p.runs[0].text = text
    else:
        p.add_run(text)


def insert_paragraph_after(paragraph, text: str) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._element.addnext(new_p)
    np = Paragraph(new_p, paragraph._parent)
    np.add_run(text)
    return np


def find_para(doc: Document, startswith: str) -> int:
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip().startswith(startswith):
            return i
    raise ValueError(f"Paragraph not found: {startswith[:50]}")


def main() -> None:
    doc = Document(str(MANUSCRIPT))

    # 1. Table 1 caption — remove duplicate supplementary cite
    i65 = find_para(doc, "Table 1. Timing interaction")
    t65 = doc.paragraphs[i65].text
    t65 = re.sub(r"\s*\(Supplementary Table 1\)\.\s*$", ".", t65)
    set_para_text(doc.paragraphs[i65], t65.strip())

    # 2. Table 1 in-text before table (timing section)
    i112 = find_para(doc, "Beyond stage-stratified hazard summaries")
    t112 = doc.paragraphs[i112].text
    if "Table 1" not in t112.split("Supplementary")[0]:
        t112 = t112.replace(
            "Beyond stage-stratified hazard summaries, we tested whether",
            "Table 1 summarizes the leading timing-interaction signals by stage (top combinations by interaction HR). "
            "Beyond stage-stratified hazard summaries, we tested whether",
        )
        set_para_text(doc.paragraphs[i112], t112)

    # 3. Figure 5 caption period
    i106 = find_para(doc, "Figure 5 TP53")
    t106 = doc.paragraphs[i106].text.replace("Figure 5 TP53", "Figure 5. TP53", 1)
    set_para_text(doc.paragraphs[i106], t106)

    # 4. Figure 1 — formal caption + legend paragraph
    i68 = find_para(doc, "Figure 1 summarizes")
    old68 = doc.paragraphs[i68].text
    # Strip trailing supplementary from legend body for split
    supp_tail = ""
    m = re.search(r"\s*(\(Supplementary Figure 11.*)$", old68)
    if m:
        supp_tail = " " + m.group(1)
        body = old68[: m.start()].strip()
    else:
        body = old68

    body = body.replace(
        "Figure 1 summarizes a cross-sectional dead-versus-alive comparison of mutation patterns in the same integrated cohort:",
        "",
    ).strip()
    if body.startswith(":"):
        body = body[1:].strip()

    caption = (
        "Figure 1. Cross-sectional deceased-versus-living mutation enrichment overview."
    )
    set_para_text(doc.paragraphs[i68], caption)

    legend = body
    if not legend.endswith("."):
        legend += "."
    legend += supp_tail
    # Insert legend after caption if next para is not already legend
    nxt = doc.paragraphs[i68 + 1].text.strip() if i68 + 1 < len(doc.paragraphs) else ""
    if not nxt.startswith("(A) patient counts"):
        insert_paragraph_after(doc.paragraphs[i68], legend)

    # 5. Reference format [19–21]
    i61 = find_para(doc, "For the high-dimensional screen across dual mutation")
    t61 = doc.paragraphs[i61].text
    t61 = t61.replace("[19] [20,21]", "[19–21]")
    set_para_text(doc.paragraphs[i61], t61)

    # 6. Figure 7 in-text (pairwise → Figure 7, not Figure 6)
    i116 = find_para(doc, "To summarize the pairwise screen")
    t116 = doc.paragraphs[i116].text
    t116 = t116.replace(
        "in a single paired figure (Figure 6)",
        "in a single paired figure (Figure 7)",
    )
    if "Figure 7" not in t116.split("Figure 6")[0]:
        t116 = t116.replace(
            "Timing-interaction volcano plots",
            "Complementary TP53/KRAS-focused stage panels are shown in Figure 6. Timing-interaction volcano plots",
        )
    set_para_text(doc.paragraphs[i116], t116)

    # 7. Move [23] from Figure 5 para to stage/backbone para 77
    i104 = find_para(doc, "With stage-stratified survival associations established")
    t104 = doc.paragraphs[i104].text
    t104 = re.sub(r"\s*\[23\]\s*", " ", t104)
    t104 = re.sub(r"\s{2,}", " ", t104).strip()
    set_para_text(doc.paragraphs[i104], t104)

    i77 = find_para(doc, "TP53+KRAS co-occurrence remained")
    t77 = doc.paragraphs[i77].text
    if "[23]" not in t77:
        t77 = t77.replace(
            "higher-order (triple) patterns.",
            "higher-order (triple) patterns, consistent with heterogeneous metastatic presentation across strata [23].",
        )
        set_para_text(doc.paragraphs[i77], t77)

    doc.save(str(MANUSCRIPT))
    print(f"Saved fixes to {MANUSCRIPT}")


if __name__ == "__main__":
    main()
