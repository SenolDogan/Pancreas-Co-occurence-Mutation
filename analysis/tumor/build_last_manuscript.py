#!/usr/bin/env python3
"""
Build Last.docx by selecting the strongest, most novel results from
Final.docx and NewManuscript.docx, with a single coherent scientific narrative.

Notes:
- Does not modify Final.docx or NewManuscript.docx.
- Embeds key figures/tables inline; places extended validation/diagnostics in Supplementary.
"""

from __future__ import annotations

from pathlib import Path

import numpy as np
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.size = Pt(14 if level == 1 else 12)


def add_para(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.add_run(text)
    p.paragraph_format.space_after = Pt(6)


def add_figure_embed(doc: Document, image_path: Path, title: str, legend: str | None = None, width_in: float = 6.8) -> None:
    p = doc.add_paragraph()
    r = p.add_run(title)
    r.bold = True
    if image_path.is_file():
        doc.add_picture(str(image_path), width=Inches(width_in))
    else:
        add_para(doc, f"[Missing image: {image_path}]")
    if legend:
        lp = doc.add_paragraph()
        lr = lp.add_run(legend)
        lr.italic = True
    doc.add_paragraph()


def add_df_table(doc: Document, df, title: str, max_rows: int = 25) -> None:
    cap = doc.add_paragraph()
    cr = cap.add_run(title)
    cr.bold = True
    if df is None or len(df) == 0:
        add_para(doc, "[Empty table]")
        return
    d = df.head(max_rows).copy()
    nt = doc.add_table(rows=len(d) + 1, cols=len(d.columns))
    try:
        nt.style = "Table Grid"
    except Exception:
        pass
    for j, col in enumerate(d.columns):
        nt.cell(0, j).text = str(col)
    for i in range(len(d)):
        for j, col in enumerate(d.columns):
            v = d.iloc[i, j]
            nt.cell(i + 1, j).text = "" if v is None else str(v)
    doc.add_paragraph()


def copy_table_grid(dest: Document, src_table, caption: str) -> None:
    cap = dest.add_paragraph()
    cr = cap.add_run(caption)
    cr.bold = True
    n_rows = len(src_table.rows)
    n_cols = len(src_table.columns)
    nt = dest.add_table(n_rows, n_cols)
    try:
        nt.style = "Table Grid"
    except Exception:
        pass
    for i in range(n_rows):
        for j in range(n_cols):
            nt.cell(i, j).text = src_table.cell(i, j).text
    dest.add_paragraph()


def main() -> None:
    tumor = Path(__file__).resolve().parent
    base = tumor.parent  # .../survival/New/New 2

    final_path = tumor / "Final.docx"
    new_path = tumor / "NewManuscript.docx"
    out_path = tumor / "Last.docx"

    src_final = Document(str(final_path))
    src_new = Document(str(new_path))

    doc = Document()

    # Title & authors: prefer Final formatting
    for idx in (0, 2):
        t = src_final.paragraphs[idx].text.strip()
        if not t:
            continue
        p = doc.add_paragraph(t)
        if idx == 0:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(16)
    doc.add_paragraph()

    # Abstract: reuse New (more stage/timing aware), lightly tightened
    add_heading(doc, "Abstract", 1)
    add_para(
        doc,
        "Background: Pancreatic ductal adenocarcinoma (PDAC) is dominated by co-occurring driver alterations, yet how "
        "clinically defined stage groups and sampling-time heterogeneity modify the prognostic meaning of recurrent "
        "dual and triple mutation patterns remains incompletely mapped.\n\n"
        "Methods: We studied 2,264 PDAC patients with integrated clinical records and binary mutation calls across ten "
        "prespecified genes (TP53, KRAS, CDKN2A, SMAD4, ARID1A, ATM, PIK3CA, BRAF, GNAS, RNF43). We quantified outcome-associated "
        "enrichment (deceased vs living) and multiplicity-controlled combination screens, and we evaluated stage-stratified "
        "overall survival (OS) associations using Cox models. We further tested a novelty layer: whether the relationship between "
        "DX2COLLECTION_YEAR (diagnosis-to-collection interval) and OS differs across mutation-combination contexts using interaction Cox models.\n\n"
        "Results: Co-mutation patterns were common and non-independent, but their survival associations varied by stage stratum. "
        "A compact timing-interaction summary highlighted combinations in which sampling-time heterogeneity materially modified OS association "
        "after adjustment and sensitivity filtering.\n\n"
        "Conclusions: Integrating enrichment-style combination profiling with stage-stratified survival modeling and timing-aware interaction tests "
        "yields a coherent evidence chain for PDAC co-mutation interpretation and motivates external validation."
    )
    add_para(doc, "Keywords: pancreatic cancer; co-mutations; stage stratification; Cox regression; interaction models; multiplicity control")

    doc.add_page_break()

    # Introduction: keep from New (already humanized); add one bridging paragraph from Final where helpful
    add_heading(doc, "Introduction", 1)
    intro_1 = src_new.paragraphs[0].text.strip()  # title line in New; skip if empty
    # New stores actual intro text later; safer to search for "Introduction" heading and take next 2 paras
    # but keep simple: reuse Final's intro paragraph + New closing framing.
    add_para(doc, src_final.paragraphs[9].text.strip())
    add_para(
        doc,
        "We address this gap by linking descriptive co-occurrence and multiplicity-aware combination screens to stage-stratified time-to-event models, "
        "and by introducing a timing-aware sensitivity layer that asks whether diagnosis-to-collection interval modifies survival associations within specific "
        "mutation-combination contexts."
    )
    doc.add_page_break()

    # Methods: keep from New (already structured)
    add_heading(doc, "Materials and Methods", 1)
    # Copy Methods block from New between its "Materials and Methods" and "Results" markers
    # (paragraph indices are stable in the generator; but we keep robust by copying a fixed window used before)
    for p in src_new.paragraphs:
        if p.text.strip() == "Materials and Methods":
            break
    # NewManuscript.docx is generated, so direct paragraph scanning is noisy; keep minimal methods paragraph set:
    add_para(
        doc,
        "We analyzed an integrated PDAC cohort (n=2,264) with clinical annotations and binary mutation calls for ten prespecified genes. "
        "Outcome-associated enrichment contrasts (deceased vs living) were treated as descriptive. For OS, we applied Cox proportional hazards models within "
        "three stage groups (Metastatic; Resectable; Borderline Resectable/Locally Advanced). Multiple testing was handled with an explicit FDR-first policy, "
        "with conservative bounds reported as sensitivity. Timing-interaction models evaluated OS ~ dx_yr + combo + dx_yr×combo, with adjusted and sensitivity-filtered fits."
    )
    doc.add_page_break()

    # Results: choose strongest narrative from New; add concise validation/diagnostics from Final as Supplementary
    add_heading(doc, "Results", 1)

    add_heading(doc, "Cohort overview and outcome-associated enrichment", 2)
    add_para(
        doc,
        "We first summarize outcome-associated enrichment comparing deceased versus living patients as a descriptive overview (Figure 1). "
        "We then contextualize the cohort by the three-group stage landscape used for stratified analyses (Figure 2)."
    )

    # Table 1: timing interaction highlights (novelty)
    try:
        import pandas as pd

        cmp_xlsx = tumor / "Stage_DX2Collection_Interaction_Sensitivity_Comparison.xlsx"
        if cmp_xlsx.is_file():
            cmp_df = pd.read_excel(cmp_xlsx, sheet_name="Comparison_AllFilters")
            keep = [c for c in ["Stage", "Feature", "HR_int_adj_ALL", "p_int_adj_ALL", "q_int_adj_ALL", "stable_FDR_lt_0p10"] if c in cmp_df.columns]
            top = cmp_df.copy()
            if "stable_FDR_lt_0p10" in top.columns and "p_int_adj_ALL" in top.columns:
                top = top.sort_values(["stable_FDR_lt_0p10", "p_int_adj_ALL"], ascending=[False, True])
            top = top[keep].head(12)
            for c in list(top.columns):
                if c.startswith("HR") or c.startswith("p_") or c.startswith("q_") or c.endswith("_ALL"):
                    try:
                        top[c] = top[c].astype(float).map(lambda x: "" if x != x else f"{x:.2f}")
                    except Exception:
                        pass
            add_df_table(
                doc,
                top,
                "Table 1. Timing interaction highlights (adjusted Cox; top signals across stages). Full sensitivity comparison appears in Supplementary Table S1.",
                max_rows=20,
            )
    except Exception:
        add_para(doc, "[Table 1 could not be embedded; see Supplementary Table S1.]")

    add_figure_embed(
        doc,
        base / "01_Dead_Alive_Comprehensive_Analysis.png",
        "Figure 1. Dead vs Alive enrichment overview (supporting).",
        legend=(
            "Legend: Multi-panel overview summarizing deceased-versus-living enrichment, combination scores, and selected clinical summaries. "
            "This figure is descriptive context and is interpreted alongside stage-stratified time-to-event modeling."
        ),
    )
    add_figure_embed(
        doc,
        tumor / "Stage_3Group_Comprehensive_Analysis.png",
        "Figure 2. Integrated cohort and three-group stage landscape.",
        legend="Legend: Cohort composition across Metastatic, Resectable, and Borderline Resectable/Locally Advanced stage groups.",
    )

    add_heading(doc, "Stage-stratified survival associations and backbone context", 2)
    add_para(
        doc,
        "Stage-stratified Cox summaries are shown in Figure 3, and the OS/timing framework is summarized in Figure 4. "
        "Backbone context is provided by TP53+KRAS(+X) and TP53/KRAS-focused stage panels (Figures 5–6)."
    )
    add_figure_embed(
        doc,
        tumor / "Stage_Figure2_AB_ForestPlots_VSTACK.png",
        "Figure 3. Stage-stratified Cox summaries (Panels A–B).",
        legend=(
            "Legend: Panel A shows univariate Cox associations within stage strata; Panel B contrasts penalized multivariable models. "
            "HR>1 indicates shorter survival for feature-positive groups."
        ),
    )
    add_figure_embed(
        doc,
        tumor / "Stage_06_OS_Months_Methodology_Analysis.png",
        "Figure 4. Overall survival endpoint definition and timing framework.",
        legend="Legend: Schematic summary of OS definition and timing variables used for interaction modeling.",
    )
    add_figure_embed(
        doc,
        tumor / "Stage_05_TP53_KRAS_Detailed_Analysis.png",
        "Figure 5. TP53+KRAS(+X) detailed stage comparison.",
        legend="Legend: Stage-wise prevalence context for TP53+KRAS and selected TP53+KRAS+third-gene patterns.",
    )
    add_figure_embed(
        doc,
        tumor / "Stage_02_TP53_KRAS_Focused_Analysis.png",
        "Figure 6. TP53/KRAS-focused stage panels.",
        legend="Legend: Focused TP53 and KRAS status across stage groups.",
    )

    add_heading(doc, "Pairwise prioritization: co-occurrence versus OS association", 2)
    add_para(doc, "A paired summary of non-independence and OS association for mutation pairs is shown in Figure 7.")
    add_figure_embed(
        doc,
        tumor / "Stage_Figure6_AB_AdditiveBars_CoxPairsForest.png",
        "Figure 7. Pairwise co-occurrence and OS association summaries (Panels A–B).",
        legend=(
            "Legend: Panel A shows top deviations from independence for mutation pairs; Panel B shows top univariate Cox associations "
            "for mutation pairs within stage strata."
        ),
    )

    doc.add_page_break()

    add_heading(doc, "Discussion", 1)
    add_para(
        doc,
        "This Last manuscript prioritizes the most publishable evidence chain across both drafts: a descriptive enrichment overview, "
        "stage-stratified time-to-event associations, and a timing-aware interaction layer that addresses a clinically relevant source of heterogeneity "
        "(diagnosis-to-collection interval). The central contribution is not a new driver list, but a stage- and timing-aware interpretation framework "
        "for recurrent co-mutation patterns."
    )

    doc.add_page_break()

    add_heading(doc, "References", 1)
    # Use the same references block that NewManuscript already expanded to ~30
    # (copy from NewManuscript reference paragraphs as bullets)
    in_refs = False
    for p in src_new.paragraphs:
        t = p.text.strip()
        if t == "References":
            in_refs = True
            continue
        if in_refs:
            if t == "Supplementary Information":
                break
            if t:
                doc.add_paragraph(t, style="List Bullet")

    doc.add_page_break()
    add_heading(doc, "Supplementary Information", 1)
    add_para(
        doc,
        "Supplementary materials include the full timing-interaction sensitivity table (Supplementary Table S1) and extended figures "
        "for triples, interaction volcano diagnostics, and exploratory sensitivity panels. We also retain the original validation/correction and "
        "functional-context panels from the earlier draft as extended context."
    )

    add_heading(doc, "Supplementary Tables", 2)
    try:
        import pandas as pd

        cmp_xlsx = tumor / "Stage_DX2Collection_Interaction_Sensitivity_Comparison.xlsx"
        if cmp_xlsx.is_file():
            cmp_df = pd.read_excel(cmp_xlsx, sheet_name="Comparison_AllFilters")
            add_df_table(doc, cmp_df, "Supplementary Table S1. Timing interaction sensitivity comparison (compact view)", max_rows=35)
    except Exception:
        add_para(doc, "[Could not embed Supplementary Table S1; see xlsx outputs.]")

    add_heading(doc, "Supplementary Figures", 2)
    supp = [
        ("Supplementary Figure S1. Triple additive top bars (vs independence)", tumor / "Stage_Triple_Additive_TopBars.png"),
        ("Supplementary Figure S2. Cox triples top forest (univariate; stage-wise)", tumor / "Stage_Cox_Triples_TopForest.png"),
        ("Supplementary Figure S3. dx_yr×pair interaction volcano (DX_GE_0)", tumor / "Stage_DX2Collection_Pair_Interaction_Volcano_DX_GE_0.png"),
        ("Supplementary Figure S4. dx_yr×triple interaction volcano (DX_GE_0)", tumor / "Stage_DX2Collection_Triple_Interaction_Volcano_DX_GE_0.png"),
        ("Supplementary Figure S5. Stage 12 short survival comparison", tumor / "Stage_12_TP53_KRAS_ShortSurvival_Comparison.png"),
        ("Supplementary Figure S6. Stage 14 diabetic vs non-diabetic", tumor / "Stage_14_Diabetic_vs_NonDiabetic_Gene_Analysis.png"),
        ("Supplementary Figure S7. dx_yr×triple interaction volcano (DX_0_TO_5 sensitivity)", tumor / "Stage_DX2Collection_Triple_Interaction_Volcano_DX_0_TO_5.png"),
        # From Final-style extended context
        ("Supplementary Figure S8. Multiple comparison correction analysis", tumor / "Stage_08_Multiple_Comparison_Correction_Analysis.png"),
        ("Supplementary Figure S9. Validation cohort analysis", tumor / "Stage_09_Validation_Cohort_Analysis.png"),
        ("Supplementary Figure S10. Functional validation analysis", tumor / "Stage_10_Functional_Validation_Analysis.png"),
    ]
    for title, path in supp:
        add_figure_embed(doc, path, title, legend="Legend: See panel title.", width_in=6.8)

    doc.save(str(out_path))
    print(f"Wrote {out_path}")


if __name__ == "__main__":
    main()

