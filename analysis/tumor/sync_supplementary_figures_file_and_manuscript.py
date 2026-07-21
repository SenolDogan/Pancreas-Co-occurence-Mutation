#!/usr/bin/env python3
"""
1) Build or overwrite Tumor/supplementary/Supplementary Figures.docx with figures 1–14 (embedded PNGs).
2) Strip embedded images after Supplementary Figure 13–14 entries in NewManuscript.docx; keep inventory titles only.
3) Ensure main-text citation to Supplementary Figures 13–14 before the Supplementary Materials section.

Requires PNGs referenced below (already in Tumor/ or parent New 2/).
"""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

TUMOR = Path(__file__).resolve().parent
NEW2 = TUMOR.parent
SUPP_DIR = TUMOR / "supplementary"
SUPP_FIG_DOC = SUPP_DIR / "Supplementary Figures.docx"
MANUSCRIPT = TUMOR / "NewManuscript.docx"

# Inventory titles match unify_supplementary_numbering.FIGURES (1–12) + short titles for manuscript (13–14).
TITLE_13_MANUSCRIPT = (
    "Supplementary Figure 13. Three-layer analytical decision schematic for the discovery cohort."
)
TITLE_14_MANUSCRIPT = (
    "Supplementary Figure 14. Kaplan–Meier overall survival for TP53+KRAS carriers versus "
    "others (integrated discovery cohort)."
)

FIGURES_DOC: list[tuple[str, Path, str | None, float]] = [
    (
        "Supplementary Figure 1. Triple-mutation additive deviations from independence (top positive and negative bars by stage)",
        TUMOR / "Stage_Triple_Additive_TopBars.png",
        "Highest positive and negative deviations from independence for mutation triples (excess/deficit versus independence).",
        6.5,
    ),
    (
        "Supplementary Figure 2. Univariate Cox forest plot for top triple-mutation combinations by stage",
        TUMOR / "Stage_Cox_Triples_TopForest.png",
        "Prespecified triple combinations ranked within stage strata; hazard ratios with confidence intervals.",
        6.5,
    ),
    (
        "Supplementary Figure 3. Diagnosis-to-collection timing × pairwise combination interaction volcano (DX2COLLECTION_YEAR ≥ 0)",
        TUMOR / "Stage_DX2Collection_Pair_Interaction_Volcano_DX_GE_0.png",
        "Adjusted interaction contrasts with dx_yr under DX ≥ 0 filter.",
        6.5,
    ),
    (
        "Supplementary Figure 4. Diagnosis-to-collection timing × triple-mutation interaction volcano (DX2COLLECTION_YEAR ≥ 0)",
        TUMOR / "Stage_DX2Collection_Triple_Interaction_Volcano_DX_GE_0.png",
        "Adjusted interaction contrasts for triples under DX ≥ 0 filter.",
        6.5,
    ),
    (
        "Supplementary Figure 5. TP53/KRAS patterns in short overall-survival subsets (exploratory sensitivity)",
        TUMOR / "Stage_12_TP53_KRAS_ShortSurvival_Comparison.png",
        "Exploratory subgroup panel; descriptive only.",
        6.5,
    ),
    (
        "Supplementary Figure 6. Diabetes-stratified exploratory gene analysis",
        TUMOR / "Stage_14_Diabetic_vs_NonDiabetic_Gene_Analysis.png",
        "Hypothesis-generating diabetes strata comparison.",
        6.5,
    ),
    (
        "Supplementary Figure 7. Triple-mutation timing interaction volcano under DX2COLLECTION_YEAR 0–5 years sensitivity filter",
        TUMOR / "Stage_DX2Collection_Triple_Interaction_Volcano_DX_0_TO_5.png",
        "Sensitivity restricting diagnosis-to-collection interval to 0–5 years.",
        6.5,
    ),
    (
        "Supplementary Figure 8. Multiple-comparison correction diagnostics for combination screens",
        TUMOR / "Stage_08_Multiple_Comparison_Correction_Analysis.png",
        "Diagnostics for multiplicity across prespecified pairwise/triple enumeration.",
        6.5,
    ),
    (
        "Supplementary Figure 9. Internal discovery–validation cohort diagnostics",
        TUMOR / "Stage_09_Validation_Cohort_Analysis.png",
        "Internal split reproducibility summaries for registry-derived summaries.",
        6.5,
    ),
    (
        "Supplementary Figure 10. Pathway-oriented functional validation summary",
        TUMOR / "Stage_10_Functional_Validation_Analysis.png",
        "Interpretive module grouping recurrent combinations.",
        6.5,
    ),
    (
        "Supplementary Figure 11. Deceased-versus-living comprehensive mutation enrichment overview",
        NEW2 / "01_Dead_Alive_Comprehensive_Analysis.png",
        "Cross-sectional deceased-versus-alive prevalence contrasts for context.",
        6.5,
    ),
    (
        "Supplementary Figure 12. Multivariable Cox internal validation summary by stage",
        TUMOR / "Stage_Multivariable_Cox_Validation.png",
        "Stability/discrimination summaries for penalized Cox fits within strata.",
        6.5,
    ),
    (
        "Supplementary Figure 13. Three-layer analytical decision schematic for the discovery cohort",
        TUMOR / "Figure_Decision_Schematic_ThreeLayers.png",
        "(i) Enrichment contrasts (deceased versus alive), "
        "(ii) stage-stratified Cox summaries for OS, "
        "(iii) diagnosis-to-biosample timing interactions (DX2COLLECTION YEAR surrogate). "
        "Public genomic atlases replicate only simplified univariate dual-mutation checks; external timing matching is not supported.",
        6.4,
    ),
    (
        "Supplementary Figure 14. Kaplan–Meier overall survival for TP53+KRAS versus all other patients "
        "(discovery cohort; OS_MONTHS>0; panels all prespecified stage groups and metastatic stratum)",
        TUMOR / "KM_TP53_KRAS_SuppFigure14_AB.png",
        "Two-sample log-rank p-values are printed on each panel (uncorrected). "
        "Numeric values from feasibility export: "
        "(A) p≈1.945×10⁻⁷; (B) p≈3.145×10⁻⁴. Interpret together with stage-stratified Cox in the main text.",
        6.8,
    ),
]


def add_figure(doc: Document, title: str, png: Path, legend: str | None, width_in: float) -> None:
    p = doc.add_paragraph()
    r = p.add_run(title.strip())
    r.bold = True
    if png.is_file():
        ip = doc.add_paragraph()
        ip.alignment = WD_ALIGN_PARAGRAPH.CENTER
        ip.add_run().add_picture(str(png), width=Inches(width_in))
    else:
        doc.add_paragraph(f"[Missing image file: {png}]")
    if legend:
        lp = doc.add_paragraph()
        lr = lp.add_run(legend)
        lr.italic = True
    doc.add_paragraph()


def write_supplementary_figures_docx() -> None:
    SUPP_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    t = doc.add_paragraph("Supplementary Figures")
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in t.runs:
        run.bold = True
        run.font.size = Pt(16)
    doc.add_paragraph()

    for title, png, leg, wi in FIGURES_DOC:
        add_figure(doc, title + ".", png, leg, wi)

    doc.save(str(SUPP_FIG_DOC))
    print(f"Wrote {SUPP_FIG_DOC}")


def para_is_image_only(p) -> bool:
    return ("pic:pic" in p._element.xml) and not p.text.strip()


def find_supplementary_materials_body_idx(doc: Document) -> int | None:
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip() == "Supplementary Materials":
            return i
    return None


def manuscript_body_has_sf13_cite(doc: Document) -> bool:
    end = find_supplementary_materials_body_idx(doc)
    if end is None:
        end = len(doc.paragraphs)
    blob = "\n".join(p.text for p in doc.paragraphs[:end])
    return bool(re.search(r"Supplementary Figures?\s+1[34]", blob, re.I))


def strip_sf13_sf14_inline_images_manuscript(doc: Document) -> None:
    delete_idx: list[int] = []
    for i, p in enumerate(doc.paragraphs):
        t = p.text.strip()
        if t.startswith("Supplementary Figure 13."):
            p.text = TITLE_13_MANUSCRIPT
            if i + 1 < len(doc.paragraphs) and para_is_image_only(doc.paragraphs[i + 1]):
                delete_idx.append(i + 1)
        elif t.startswith("Supplementary Figure 14."):
            p.text = TITLE_14_MANUSCRIPT
            if i + 1 < len(doc.paragraphs) and para_is_image_only(doc.paragraphs[i + 1]):
                delete_idx.append(i + 1)
    for j in sorted(set(delete_idx), reverse=True):
        el = doc.paragraphs[j]._element
        el.getparent().remove(el)


def add_methods_citation_for_sf13_14(doc: Document) -> None:
    if manuscript_body_has_sf13_cite(doc):
        return
    cite = (
        " Supplementary Figures 13 and 14 provide the three-layer analytic workflow schematic (enrichment, "
        "stage-stratified Cox summaries, timing interactions) and Kaplan–Meier overall survival contrasts comparing "
        "TP53+KRAS carriers against all other discovery-cohort patients (two-sample log-rank tests), using the same "
        "binary gene definitions as the stage-stratified Cox models."
    )
    idx = None
    for i, p in enumerate(doc.paragraphs):
        if p.text.startswith("OS was defined as time from diagnosis"):
            idx = i
            break
    if idx is None:
        print("WARN: Could not find OS Methods paragraph for citation insertion.", flush=True)
        return
    p = doc.paragraphs[idx]
    base = p.text.rstrip()
    needle = cite.lstrip().split(",")[0]  # "Supplementary Figures 13 and 14 provide"
    if cite.lstrip() in base.replace("\n", " ") or needle in base:
        return
    if base.endswith("."):
        p.text = base[:-1].rstrip() + cite
        if not p.text.endswith("."):
            p.text += "."
    else:
        p.text = base + cite
        if not p.text.endswith("."):
            p.text += "."


def main() -> int:
    if not MANUSCRIPT.is_file():
        print(f"Missing {MANUSCRIPT}")
        return 1

    missing = [(t, p) for t, p, _, _ in FIGURES_DOC if not p.is_file()]
    if missing:
        for t, p in missing:
            print(f"Missing PNG ({t[:50]}…): {p}")

    write_supplementary_figures_docx()

    doc = Document(str(MANUSCRIPT))
    strip_sf13_sf14_inline_images_manuscript(doc)
    add_methods_citation_for_sf13_14(doc)
    doc.save(str(MANUSCRIPT))

    doc2 = Document(str(MANUSCRIPT))
    if not manuscript_body_has_sf13_cite(doc2):
        print("WARN: citation check failed after save.", flush=True)
        return 2

    print(f"Updated {MANUSCRIPT} (SF13–14 titles only; Methods cites SF13–14).")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
