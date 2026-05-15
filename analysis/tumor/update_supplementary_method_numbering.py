#!/usr/bin/env python3
"""
Rename supplementary methods to Supplementary Method 1–9 in:
  - supplementary/Supplementary Methods.docx
  - NewManuscript.docx (inventory list + in-text § citations)
"""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document


METHOD_TITLES = [
    "Scope and relationship to the main manuscript",
    "Primary analytic cohort (discovery)",
    "Synergy and enrichment metrics (multiplicative, additive, and protective scores)",
    "Stage-stratified Cox and multivariable extensions",
    "DX2COLLECTION_YEAR × mutation-combination interaction models and sensitivity filters",
    "Internal validation (stratified split and cross-validation)",
    "External validation in public PDAC cohorts (cBioPortal)",
    "Overview of supplementary figure content",
    "Software and computational environment",
]

# Headings in Supplementary Methods.docx (short titles after renumber)
SUPP_METHODS_DOC_HEADINGS = [
    "Scope and relationship to the main manuscript",
    "Primary analytic cohort (discovery)",
    "Synergy and enrichment metrics (supplementary detail)",
    "Stage-stratified Cox and multivariable extensions",
    "DX2COLLECTION_YEAR × combination interaction models",
    "Internal validation",
    "External validation (cBioPortal)",
    "Supplementary figures (content summary)",
    "Software",
]


def update_supplementary_methods_docx(path: Path) -> None:
    doc = Document(str(path))
    heading_idx = 0
    for p in doc.paragraphs:
        t = p.text.strip()
        if not t:
            continue
        # Match old "N. Title" heading style (Heading 1)
        m = re.match(r"^(\d+)\.\s+(.+)$", t)
        if m and p.style.name.startswith("Heading") and heading_idx < len(SUPP_METHODS_DOC_HEADINGS):
            n = heading_idx + 1
            title = SUPP_METHODS_DOC_HEADINGS[heading_idx]
            new_text = f"Supplementary Method {n}. {title}"
            for r in p.runs:
                r.text = ""
            if p.runs:
                p.runs[0].text = new_text
            else:
                p.add_run(new_text)
            heading_idx += 1
    doc.save(str(path))
    print(f"Updated {heading_idx} method headings in {path}")


def replace_method_citations_in_paragraph(p) -> bool:
    """Replace Supplementary Methods §N -> Supplementary Method N in all runs."""
    changed = False
    full = p.text
    new_full = full
    new_full = re.sub(
        r"Supplementary Methods\s*§\s*(\d+)",
        r"Supplementary Method \1",
        new_full,
    )
    new_full = re.sub(
        r"Supplementary Methods,\s*Section\s*(\d+)",
        r"Supplementary Method \1",
        new_full,
    )
    if new_full != full:
        # Preserve first run formatting loosely: collapse to single run text update
        if p.runs:
            p.runs[0].text = new_full
            for r in p.runs[1:]:
                r.text = ""
        else:
            p.add_run(new_full)
        changed = True
    return changed


def update_manuscript_inventory(doc: Document) -> None:
    """Supplementary Methods numbered list at end (paragraphs after 'Supplementary Methods' heading)."""
    start = None
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip() == "Supplementary Methods" and i > 100:
            start = i + 1
            break
    if start is None:
        return
    for n in range(1, 10):
        idx = start + n - 1
        if idx >= len(doc.paragraphs):
            break
        p = doc.paragraphs[idx]
        t = p.text.strip()
        m = re.match(r"^(\d+)\.\s+(.+)$", t)
        if not m:
            continue
        title = METHOD_TITLES[n - 1] if n <= len(METHOD_TITLES) else m.group(2)
        new_text = f"Supplementary Method {n}. {title}"
        for r in p.runs:
            r.text = ""
        if p.runs:
            p.runs[0].text = new_text
        else:
            p.add_run(new_text)


def main() -> None:
    tumor = Path(__file__).resolve().parent
    supp_methods = tumor / "supplementary" / "Supplementary Methods.docx"
    manuscript = tumor / "NewManuscript.docx"

    update_supplementary_methods_docx(supp_methods)

    doc = Document(str(manuscript))
    n_cite = 0
    for p in doc.paragraphs:
        if replace_method_citations_in_paragraph(p):
            n_cite += 1
    update_manuscript_inventory(doc)
    doc.save(str(manuscript))
    print(f"Updated {n_cite} manuscript paragraphs with Supplementary Method N citations")
    print(f"Saved {manuscript}")


if __name__ == "__main__":
    main()
