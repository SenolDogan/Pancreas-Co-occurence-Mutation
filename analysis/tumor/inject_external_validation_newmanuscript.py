#!/usr/bin/env python3
"""
Insert external validation-only paragraphs into NewManuscript.docx with yellow highlight.
Does not alter existing paragraph text except appending one sentence to the Abstract.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph


def insert_paragraph_after(paragraph: Paragraph, *, style: str | None = None) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._element.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style:
        new_para.style = style
    return new_para


def add_yellow_run(paragraph: Paragraph, text: str) -> None:
    r = paragraph.add_run(text)
    r.font.highlight_color = WD_COLOR_INDEX.YELLOW


def yellow_heading(paragraph: Paragraph, text: str) -> None:
    add_yellow_run(paragraph, text)


def main() -> None:
    tumor = Path(__file__).resolve().parent
    path = tumor / "NewManuscript.docx"
    doc = Document(str(path))

    # --- Supplementary Table S2 (after S1 line) ---
    p157 = doc.paragraphs[157]
    ps2 = insert_paragraph_after(p157, style="Normal")
    add_yellow_run(
        ps2,
        "Supplementary Table S2. External public cohort replication summary (cBioPortal): studies paad_tcga_gdc, "
        "paad_tcga_pan_can_atlas_2018, pancreas_cptac_gdc, and pdac_msk_2024; columns list the cBioPortal "
        "\"all_cases_with_mutation_data\" sample-list identifier, cases with OS_MONTHS/OS_STATUS fields, Cox-eligible N "
        "after excluding OS_MONTHS≤0, TP53+KRAS co-mutation prevalence (%), univariate Cox HR (95% CI) for TP53+KRAS versus "
        "others, and two-sided p-value under the manuscript’s binary mutation definition (see Methods).",
    )

    # --- Discussion (after paragraph 110; before Limitations) ---
    p110 = doc.paragraphs[110]
    pd1 = insert_paragraph_after(p110, style="Heading 2")
    yellow_heading(pd1, "External validity and cohort independence (cBioPortal)")
    pd2 = insert_paragraph_after(pd1, style="Normal")
    add_yellow_run(
        pd2,
        "paad_tcga_gdc and paad_tcga_pan_can_atlas_2018 represent overlapping TCGA PAAD case harmonizations in cBioPortal "
        "and should be interpreted as one TCGA replication family rather than two independent patient populations. "
        "pdac_msk_2024 provides the closest co-mutation prevalence match to the primary cohort and the closest univariate "
        "hazard ratio magnitude, supporting robustness of an adverse OS association for TP53+KRAS under binary gene "
        "summaries despite platform and annotation differences. pancreas_cptac_gdc aligns co-mutation prevalence more "
        "closely than TCGA but yields a hazard ratio whose 95% confidence interval includes unity, indicating directional "
        "consistency without definitive statistical replication at conventional α=0.05 under the simplified external Cox "
        "specification.",
    )

    # --- Results (after TP53+KRAS backbone paragraph 63) ---
    p63 = doc.paragraphs[63]
    pr1 = insert_paragraph_after(p63, style="Heading 2")
    yellow_heading(pr1, "External replication in public cohorts (cBioPortal)")
    pr2 = insert_paragraph_after(pr1, style="Normal")
    add_yellow_run(
        pr2,
        "Across paad_tcga_gdc, paad_tcga_pan_can_atlas_2018, pancreas_cptac_gdc, and pdac_msk_2024, TP53+KRAS "
        "co-mutation prevalence was 43.2% (N=185), 48.9% (N=184), 65.2% (clinical N=161; Cox N=129 after OS_MONTHS>0 "
        "filtering), and 73.2% (N=2270), respectively—compared with ~71.6% in the primary integrated cohort. Univariate Cox "
        "hazard ratios for TP53+KRAS versus others were HR=2.007 (95% CI 1.347–2.989, p=0.000614) and HR=2.037 (95% CI "
        "1.353–3.069, p=0.000661) in the two TCGA-derived studies, HR=1.343 (95% CI 0.932–1.937, p=0.114) in "
        "pancreas_cptac_gdc, and HR=1.369 (95% CI 1.220–1.538, p=1.06×10⁻⁷) in pdac_msk_2024, matching the adverse hazard "
        "direction in the primary cohort (HR≈1.35) with the closest quantitative agreement in the large MSK cohort.",
    )

    # --- Methods (after internal validation paragraph 31) ---
    p31 = doc.paragraphs[31]
    pm1 = insert_paragraph_after(p31, style="Heading 2")
    yellow_heading(pm1, "External validation cohorts (cBioPortal)")
    pm2 = insert_paragraph_after(pm1, style="Normal")
    add_yellow_run(
        pm2,
        "Using the public cBioPortal API, we indexed four PDAC studies for a harmonized replication module: paad_tcga_gdc "
        "(TCGA Pancreatic Adenocarcinoma, GDC, 2025), paad_tcga_pan_can_atlas_2018 (TCGA PAAD, Pan-Cancer Atlas), "
        "pancreas_cptac_gdc (CPTAC Pancreatic Cancer, GDC, 2025), and pdac_msk_2024 (MSK PDAC, Nat Med 2024). For each "
        "study we used the cBioPortal \"all_cases_with_mutation_data\" sample list and the study’s extended mutation "
        "profile. Gene presence was coded as binary if ≥1 somatic mutation record was reported for that gene on the "
        "patient’s indexed tumor sample. OS_MONTHS and OS_STATUS were taken from cBioPortal patient clinical data; "
        "events were coded for deceased states (including TCGA-style 1:DECEASED tokens). We fitted the same univariate "
        "Cox proportional hazards model used as a descriptive backbone check elsewhere (TP53+KRAS both present versus "
        "all other patients) among patients with OS_MONTHS>0. This module does not recompute synergy scores, "
        "multiplicity-adjusted combination screens, or stage-stratified models in those external cohorts.",
    )

    # --- Abstract: append one yellow sentence to existing paragraph 10 ---
    p10 = doc.paragraphs[10]
    if "paad_tcga_gdc" in p10.text:
        raise SystemExit("Abstract already contains external validation sentence; aborting to avoid duplicate.")
    if p10.text and not p10.text.endswith(" "):
        add_yellow_run(p10, " ")
    add_yellow_run(
        p10,
        "In four publicly accessible cBioPortal cohorts (paad_tcga_gdc; paad_tcga_pan_can_atlas_2018; pancreas_cptac_gdc; "
        "pdac_msk_2024), the same binary TP53+KRAS definition reproduced a consistent adverse overall survival hazard "
        "direction in TCGA and MSK sources (Supplementary Table S2).",
    )

    doc.save(str(path))
    print(f"Updated: {path}")


if __name__ == "__main__":
    main()
