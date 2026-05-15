#!/usr/bin/env python3
"""Build Supplementary Methods.docx, Supplementary Figures.docx, Supplementary Tables.xlsx in Tumor/supplementary/."""

from __future__ import annotations

from pathlib import Path

import numpy as np
import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.size = Pt(14 if level == 1 else 12)


def add_para(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)


def add_figure(doc: Document, image_path: Path, title: str, legend: str | None = None, width_in: float = 6.5) -> None:
    p = doc.add_paragraph()
    r = p.add_run(title)
    r.bold = True
    if image_path.is_file():
        doc.add_picture(str(image_path), width=Inches(width_in))
    else:
        add_para(doc, f"[Missing image: {image_path.name}]")
    if legend:
        lp = doc.add_paragraph()
        lr = lp.add_run(legend)
        lr.italic = True
    doc.add_paragraph()


def build_supplementary_methods(tumor: Path, out: Path) -> None:
    doc = Document()
    title = doc.add_paragraph("Supplementary Methods")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.bold = True
        run.font.size = Pt(16)
    doc.add_paragraph()

    add_heading(doc, "1. Scope and relationship to the main manuscript", 1)
    add_para(
        doc,
        "Supplementary Methods document analysis extensions, sensitivity filters, internal and external "
        "validation modules, and timing-aware interaction models that support the primary Results without "
        "altering core definitions of overall survival (OS), binary mutation presence, or the three-group "
        "clinical stage scheme (Metastatic; Resectable; Borderline Resectable/Locally Advanced). "
        "Supplementary Tables and Supplementary Figures are provided as separate files "
        "(Supplementary_Tables.xlsx; Supplementary_Figures.docx).",
    )

    add_heading(doc, "2. Primary analytic cohort (discovery)", 1)
    add_para(
        doc,
        "The integrated discovery cohort (N=2330 patients) was derived from Tumor/Merged.xlsx by collapsing "
        "mutation-level rows to one record per PATIENT_ID. Binary mutation indicators were defined for ten "
        "prespecified genes (TP53, KRAS, CDKN2A, SMAD4, ARID1A, ATM, PIK3CA, BRAF, GNAS, RNF43): presence "
        "if Hugo_Symbol matched the gene at least once. OS_MONTHS and OS_STATUS followed the clinical dictionary "
        "used in stage scripts (OS_MONTHS > 0 for time-to-event models; event coded from 1:DECEASED). "
        "DX2COLLECTION_YEAR (dx_yr) denotes years from diagnosis to sample collection.",
    )

    add_heading(doc, "3. Synergy and enrichment metrics (supplementary detail)", 1)
    add_para(
        doc,
        "Multiplicative synergy compared observed joint mutation frequency to the product of marginal frequencies. "
        "Additive synergy was computed as p_AB − p_A·p_B (deviation from independence). Protective scores contrasted "
        "mutation prevalence between deceased and living patients. Dual screens covered 45 pairwise combinations; "
        "triple screens used prespecified higher-order contexts. Multiplicity control used Benjamini–Hochberg FDR "
        "as the primary framework, with Bonferroni/Holm reported as conservative sensitivity bounds.",
    )

    add_heading(doc, "4. Stage-stratified Cox and multivariable extensions", 1)
    add_para(
        doc,
        "Univariate Cox proportional hazards models for OS were fit within each stage stratum for single genes, "
        "pairs, and triples. Penalized multivariable Cox summaries and internal validation metrics are tabulated in "
        "Supplementary Tables (Stage_Multivariable_Cox_Report sheets). Stage_Mutation_Additive_OS_Report.xlsx "
        "contains gene frequency tests, additive pair/triple rankings, and Cox outputs by stage.",
    )

    add_heading(doc, "5. DX2COLLECTION_YEAR × combination interaction models", 1)
    add_para(
        doc,
        "Timing-interaction models tested whether diagnosis-to-collection interval modified OS associations for "
        "mutation combinations. Primary adjusted specification: Cox(OS_MONTHS ~ dx_yr + combo + dx_yr×combo + AGE + SEX + DISEASE_STATUS) "
        "within stage strata. Unadjusted models omitted clinical covariates when noted. Three dx_yr filters were applied: "
        "ALL (no filter), DX_GE_0 (exclude dx_yr < 0), and DX_0_TO_5 (0 ≤ dx_yr ≤ 5 years). Supplementary Table S1 "
        "consolidates adjusted interaction p-values and stability flags across filters (see "
        "Stage_DX2Collection_Interaction_Sensitivity_Comparison.xlsx). Full model outputs appear in "
        "Stage_DX2Collection_Combo_OS_Interaction.xlsx and Stage_DX2Collection_Combo_OS_Interaction_Sensitivity.xlsx.",
    )

    add_heading(doc, "6. Internal validation", 1)
    add_para(
        doc,
        "Internal validation on the discovery registry used stratified train/validation splits and cross-validation "
        "for descriptive synergy-style scores and multivariable Cox stability (stage_multivariable_cox_validation.py; "
        "Validation_Cohort_Analysis.xlsx in the parent project folder when available). These analyses assess "
        "reproducibility within the same merged cohort and are not external multi-center replication.",
    )

    add_heading(doc, "7. External validation (cBioPortal)", 1)
    add_para(
        doc,
        "External replication used the public cBioPortal REST API (May 2026) for four PDAC studies: paad_tcga_gdc "
        "(TCGA GDC 2025), paad_tcga_pan_can_atlas_2018 (TCGA Pan-Cancer Atlas), pancreas_cptac_gdc (CPTAC GDC 2025), "
        "and pdac_msk_2024 (MSK PDAC). For each study we used the cBioPortal category "
        "\"all_cases_with_mutation_data\" sample list and the study extended mutation profile. Gene presence was "
        "binary if ≥1 somatic mutation record existed for that entrezGeneId on the patient. OS_MONTHS and OS_STATUS "
        "were taken from patient-level clinical data; events were coded for deceased states. The harmonized external "
        "test was univariate Cox: TP53+KRAS both present versus all other patients among cases with OS_MONTHS > 0. "
        "This module did not recompute synergy scores, FDR combination screens, or stage-stratified models in external "
        "cohorts. Results are in Supplementary Table S2 and were generated by cbioportal_pdac_validation_report.py.",
    )

    add_heading(doc, "8. Supplementary figures (content summary)", 1)
    add_para(
        doc,
        "Supplementary Figure S1: top triple additive deviations from independence. S2: univariate Cox forest for "
        "top triples by stage. S3–S4: dx_yr×pair and dx_yr×triple interaction volcanoes (DX_GE_0). S5: short-OS "
        "TP53/KRAS sensitivity. S6: diabetes-stratified exploratory panel. S7: triple interaction volcano under "
        "DX_0_TO_5 sensitivity. Additional diagnostic panels (multiplicity correction, internal validation cohort, "
        "functional mapping) are included when image files are available (S8–S10).",
    )

    add_heading(doc, "9. Software", 1)
    add_para(
        doc,
        "Python 3.x with pandas, numpy, matplotlib, seaborn, scipy, scikit-learn, statsmodels, and lifelines "
        "(Cox models). cBioPortal API access via urllib. Word outputs via python-docx.",
    )

    doc.save(str(out))
    print(f"Wrote {out}")


def build_supplementary_figures(tumor: Path, out: Path) -> None:
    doc = Document()
    title = doc.add_paragraph("Supplementary Figures")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.bold = True
        run.font.size = Pt(16)
    doc.add_paragraph()

    figs = [
        (
            "Supplementary Figure S1. Triple additive top bars (vs independence)",
            tumor / "Stage_Triple_Additive_TopBars.png",
            "Highest positive and negative deviations from independence for mutation triples (excess co-occurrence).",
        ),
        (
            "Supplementary Figure S2. Cox triples top forest (univariate; stage-wise)",
            tumor / "Stage_Cox_Triples_TopForest.png",
            "Univariate Cox forest plot for top triple combinations within stage strata (HR with confidence intervals).",
        ),
        (
            "Supplementary Figure S3. dx_yr×pair interaction volcano (DX_GE_0)",
            tumor / "Stage_DX2Collection_Pair_Interaction_Volcano_DX_GE_0.png",
            "Interaction volcano under the sensitivity filter DX2COLLECTION_YEAR ≥ 0.",
        ),
        (
            "Supplementary Figure S4. dx_yr×triple interaction volcano (DX_GE_0)",
            tumor / "Stage_DX2Collection_Triple_Interaction_Volcano_DX_GE_0.png",
            "Triple interaction volcano under the sensitivity filter DX2COLLECTION_YEAR ≥ 0.",
        ),
        (
            "Supplementary Figure S5. Stage 12 short survival comparison",
            tumor / "Stage_12_TP53_KRAS_ShortSurvival_Comparison.png",
            "Exploratory comparison of TP53/KRAS patterns in short-survival subsets.",
        ),
        (
            "Supplementary Figure S6. Stage 14 diabetic vs non-diabetic",
            tumor / "Stage_14_Diabetic_vs_NonDiabetic_Gene_Analysis.png",
            "Exploratory diabetes-stratified gene analysis panel (hypothesis-generating).",
        ),
        (
            "Supplementary Figure S7. dx_yr×triple interaction volcano (DX_0_TO_5 sensitivity)",
            tumor / "Stage_DX2Collection_Triple_Interaction_Volcano_DX_0_TO_5.png",
            "Sensitivity analysis restricting DX2COLLECTION_YEAR to 0–5 years.",
        ),
        (
            "Supplementary Figure S8. Multiple comparison correction analysis",
            tumor / "Stage_08_Multiple_Comparison_Correction_Analysis.png",
            "Multiplicity diagnostic panel for combination screens.",
        ),
        (
            "Supplementary Figure S9. Validation cohort analysis",
            tumor / "Stage_09_Validation_Cohort_Analysis.png",
            "Internal train/validation cohort diagnostics.",
        ),
        (
            "Supplementary Figure S10. Functional validation analysis",
            tumor / "Stage_10_Functional_Validation_Analysis.png",
            "Pathway-oriented functional validation summary.",
        ),
        (
            "Supplementary Figure S11. Dead vs alive comprehensive synergy (archived overview)",
            tumor.parent / "01_Dead_Alive_Comprehensive_Analysis.png",
            "Cross-sectional deceased-versus-living enrichment overview retained for comparison with stage/OS figures.",
        ),
        (
            "Supplementary Figure S12. Stage multivariable Cox validation",
            tumor / "Stage_Multivariable_Cox_Validation.png",
            "Internal multivariable Cox validation summary by stage.",
        ),
    ]
    for title, path, legend in figs:
        add_figure(doc, path, title, legend=legend)
        if title.startswith("Supplementary Figure S7"):
            doc.add_page_break()

    doc.save(str(out))
    print(f"Wrote {out}")


def _safe_sheet_name(name: str, used: set[str]) -> str:
    s = name[:31].replace("/", "-").replace("\\", "-").replace("*", "").replace("?", "").replace(":", "")
    base = s
    i = 1
    while s in used:
        suffix = f"_{i}"
        s = (base[: 31 - len(suffix)] + suffix) if len(base) + len(suffix) > 31 else base + suffix
        i += 1
    used.add(s)
    return s


def build_supplementary_tables(tumor: Path, parent: Path, out: Path) -> None:
    used: set[str] = set()
    index_rows = []

    with pd.ExcelWriter(out, engine="openpyxl") as writer:
        # Index
        index_rows.append({"Table_ID": "S1", "Sheet": "S1_Timing_Interaction", "Description": "DX×combo sensitivity comparison (compact)"})
        index_rows.append({"Table_ID": "S2", "Sheet": "S2_External_cBioPortal", "Description": "External cohort TP53+KRAS replication summary"})
        index_rows.append({"Table_ID": "S3+", "Sheet": "see sheets", "Description": "Extended stage/additive/Cox outputs from Tumor xlsx reports"})

        # S1
        cmp_xlsx = tumor / "Stage_DX2Collection_Interaction_Sensitivity_Comparison.xlsx"
        if cmp_xlsx.is_file():
            df = pd.read_excel(cmp_xlsx, sheet_name="Comparison_AllFilters")
            for c in ["stable_FDR_lt_0p10", "p_int_adj_ALL"]:
                if c not in df.columns:
                    df[c] = np.nan
            df = df.sort_values(["stable_FDR_lt_0p10", "p_int_adj_ALL"], ascending=[False, True])
            sn = _safe_sheet_name("S1_Timing_Interaction", used)
            df.to_excel(writer, sheet_name=sn, index=False)

        # S2 external validation
        s2 = pd.DataFrame(
            [
                {
                    "cBioPortal_study_id": "paad_tcga_gdc",
                    "study_label": "TCGA PAAD (GDC 2025)",
                    "sample_list": "paad_tcga_gdc_sequenced",
                    "N_clinical_OS": 185,
                    "N_Cox_OS_gt0": 184,
                    "TP53_KRAS_pct": 43.2,
                    "Cox_HR_TP53_KRAS": 2.007,
                    "Cox_CI_low": 1.347,
                    "Cox_CI_high": 2.989,
                    "Cox_p": 0.000614,
                },
                {
                    "cBioPortal_study_id": "paad_tcga_pan_can_atlas_2018",
                    "study_label": "TCGA PAAD (Pan-Cancer Atlas)",
                    "sample_list": "paad_tcga_pan_can_atlas_2018_sequenced",
                    "N_clinical_OS": 184,
                    "N_Cox_OS_gt0": 183,
                    "TP53_KRAS_pct": 48.9,
                    "Cox_HR_TP53_KRAS": 2.037,
                    "Cox_CI_low": 1.353,
                    "Cox_CI_high": 3.069,
                    "Cox_p": 0.000661,
                },
                {
                    "cBioPortal_study_id": "pancreas_cptac_gdc",
                    "study_label": "CPTAC Pancreatic Cancer (GDC 2025)",
                    "sample_list": "pancreas_cptac_gdc_sequenced",
                    "N_clinical_OS": 161,
                    "N_Cox_OS_gt0": 129,
                    "TP53_KRAS_pct": 65.2,
                    "Cox_HR_TP53_KRAS": 1.343,
                    "Cox_CI_low": 0.932,
                    "Cox_CI_high": 1.937,
                    "Cox_p": 0.114,
                },
                {
                    "cBioPortal_study_id": "pdac_msk_2024",
                    "study_label": "MSK PDAC (Nat Med 2024)",
                    "sample_list": "pdac_msk_2024_sequenced",
                    "N_clinical_OS": 2270,
                    "N_Cox_OS_gt0": 2260,
                    "TP53_KRAS_pct": 73.2,
                    "Cox_HR_TP53_KRAS": 1.369,
                    "Cox_CI_low": 1.220,
                    "Cox_CI_high": 1.538,
                    "Cox_p": 1.06e-07,
                },
                {
                    "cBioPortal_study_id": "DISCOVERY",
                    "study_label": "Primary cohort (Merged.xlsx)",
                    "sample_list": "N/A",
                    "N_clinical_OS": 2330,
                    "N_Cox_OS_gt0": 2254,
                    "TP53_KRAS_pct": 71.6,
                    "Cox_HR_TP53_KRAS": 1.35,
                    "Cox_CI_low": 1.21,
                    "Cox_CI_high": 1.51,
                    "Cox_p": 2.15e-07,
                },
            ]
        )
        sn2 = _safe_sheet_name("S2_External_cBioPortal", used)
        s2.to_excel(writer, sheet_name=sn2, index=False)

        # Copy key sheets from Tumor reports
        copies = [
            (tumor / "Stage_Mutation_Additive_OS_Report.xlsx", "S3"),
            (tumor / "Stage_DX2Collection_Combo_OS_Interaction.xlsx", "S4"),
            (tumor / "Stage_DX2Collection_Combo_OS_Interaction_Sensitivity.xlsx", "S5"),
            (tumor / "Stage_Multivariable_Cox_Report.xlsx", "S6"),
            (tumor / "Stage_3Group_Comprehensive_Analysis.xlsx", "S7"),
            (tumor / "Stage_3Group_Top_Synergy_ByStage.xlsx", "S8"),
        ]
        for xlsx_path, prefix in copies:
            if not xlsx_path.is_file():
                continue
            xl = pd.ExcelFile(xlsx_path)
            for sh in xl.sheet_names:
                df = pd.read_excel(xlsx_path, sheet_name=sh)
                sn = _safe_sheet_name(f"{prefix}_{sh}", used)
                df.to_excel(writer, sheet_name=sn, index=False)
                index_rows.append(
                    {
                        "Table_ID": prefix,
                        "Sheet": sn,
                        "Description": f"From {xlsx_path.name} / {sh}",
                    }
                )

        # Parent-folder validation / significance if present
        extra = [
            (parent / "Validation_Cohort_Analysis.xlsx", "S9"),
            (parent / "Statistical_Significance_Analysis.xlsx", "S10"),
            (parent / "Dead_Alive_Comprehensive_Analysis.xlsx", "S11"),
        ]
        for xlsx_path, prefix in extra:
            if not xlsx_path.is_file():
                continue
            xl = pd.ExcelFile(xlsx_path)
            for sh in xl.sheet_names[:8]:  # cap sheets per file for workbook size
                df = pd.read_excel(xlsx_path, sheet_name=sh)
                sn = _safe_sheet_name(f"{prefix}_{sh}", used)
                df.to_excel(writer, sheet_name=sn, index=False)
                index_rows.append(
                    {"Table_ID": prefix, "Sheet": sn, "Description": f"From {xlsx_path.name} / {sh}"}
                )

        idx = pd.DataFrame(index_rows)
        idx.to_excel(writer, sheet_name=_safe_sheet_name("Table_Index", set()), index=False)

    print(f"Wrote {out}")


def main() -> None:
    tumor = Path(__file__).resolve().parent
    parent = tumor.parent
    supp = tumor / "supplementary"
    supp.mkdir(parents=True, exist_ok=True)

    build_supplementary_methods(tumor, supp / "Supplementary Methods.docx")
    build_supplementary_figures(tumor, supp / "Supplementary Figures.docx")
    build_supplementary_tables(tumor, parent, supp / "Supplementary Tables.xlsx")


if __name__ == "__main__":
    main()
