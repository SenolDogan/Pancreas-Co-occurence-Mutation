#!/usr/bin/env python3
"""
Place supplementary in-text citations so Methods 1–9, Figures 1–12, and Tables 1–11
each appear in ascending numeric order on first mention as the reader moves through
the manuscript (Introduction through Conclusions; abstract/background excluded).
"""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document

MANUSCRIPT = Path(__file__).resolve().parent / "NewManuscript.docx"

# Body starts at Introduction (skip structured abstract / background block).
BODY_START_PARA = 24

STRIP_PATTERNS = [
    r"\s*\(Supplementary Methods?[^)]*\)",
    r"\s*\(Supplementary Tables?[^)]*\)",
    r"\s*\(Supplementary Figures?[^)]*\)",
    r"\s*\(timing volcanoes:[^)]*\)",
    r"\s*\(see Supplementary Figure \d+\)",
    r"\s*\(see also Figure \d+\)",
    r"\s*\(see Supplementary Table \d+ for the compact comparison\)",
    r"\s*\(triple-level panels:[^)]*\)",
    r"\s*\(exploratory short-OS sensitivity:[^)]*\)",
    r"\s*\(exploratory panels in Supplementary Figures[^)]*\)",
    r"\s*\(archived overview[^)]*\)",
    r"\s*\(exploratory diabetes panel:[^)]*\)",
    r"\s*Additional detailed results and analyses are provided in Supplementary Tables[^.]*\.",
    r"\s*External validation: Supplementary Method[^.]*\.",
]

# (needle, fragment, optional_min_para) — only match at/after min_para when set.
CITATIONS: list[tuple[str, str, int | None]] = [
    # Methods 1 → 9 (strict ascending first appearance)
    (
        "Here, we present a large integrative analysis of recurrent mutation combinations in PDAC",
        " (Supplementary Method 1).",
        None,
    ),
    (
        "assembled as a single patient-level analytic table",
        " (Supplementary Method 2).",
        None,
    ),
    (
        "We defined three complementary synergy scores to characterize",
        " (Supplementary Method 3; Supplementary Figures 1–2).",
        None,
    ),
    (
        "Only combinations with ≥5 patients were included in the analysis to ensure statistical reliability",
        " [18] (Supplementary Methods 4 and 5; primary timing-interaction table: Supplementary Table 1).",
        None,
    ),
    (
        "Internal validation of descriptive synergy and protective score summaries",
        " (Supplementary Method 6).",
        None,
    ),
    (
        "Using the public cBioPortal API, we indexed four PDAC studies",
        " (Supplementary Method 7; Supplementary Table 2).",
        44,
    ),
    (
        "Contextual functional mapping was performed through pathway-oriented grouping",
        " (Supplementary Method 8).",
        None,
    ),
    (
        "All analyses were performed using Python 3.x with the following packages",
        " (Supplementary Method 9).",
        None,
    ),
    # Tables 2 → 11 and Figures 3 → 12 (Table 1 at statistical analysis)
    (
        "DX2COLLECTION YEAR measures the elapsed time between diagnosis and sequencing sample collection",
        " (Supplementary Tables 3–5; Supplementary Figures 3–4).",
        None,
    ),
    (
        "For the high-dimensional screen across dual mutation combinations",
        " (Supplementary Figures 5–10; Supplementary Tables 6–10).",
        None,
    ),
    (
        "The master analytic cohort comprised 2330 patients",
        " (Supplementary Figure 6; Supplementary Tables 3, 7, and 11).",
        None,
    ),
    (
        "Figure 1. Cross-sectional deceased-versus-living mutation enrichment overview",
        " (Supplementary Figure 11; Supplementary Table 11).",
        None,
    ),
    (
        "Figure 2. Integrated cohort and three-group stage landscape",
        " (Supplementary Table 7).",
        None,
    ),
    (
        "Table 2. Genetic mutation frequencies and lethality ratios",
        " (Supplementary Tables 3 and 11).",
        None,
    ),
    (
        "Table 3. Statistically significant dual mutation combinations",
        " (Supplementary Table 10).",
        None,
    ),
    (
        "Table 5. Triple mutation combinations: Lethality and protective scores",
        " (Supplementary Tables 3 and 8).",
        None,
    ),
    (
        "Figure 3 summarizes stage-stratified Cox results",
        " (Supplementary Figure 12).",
        None,
    ),
    (
        "Stratified 70/30 splits and cross-validation reproduced correlation structure",
        " ",
        None,
    ),
    (
        "To summarize the pairwise screen in a reviewer-friendly way",
        " (see also main Figure 7).",
        None,
    ),
    (
        "Pathway-oriented summaries grouped recurrent combinations under canonical PDAC modules",
        " ",
        None,
    ),
    (
        "Across paad_tcga_gdc, paad_tcga_pan_can_atlas_2018, pancreas_cptac_gdc, and pdac_msk_2024",
        " (Supplementary Table 2).",
        79,
    ),
    (
        "Table 1 summarizes the leading timing-interaction signals by stage",
        " ",
        None,
    ),
    (
        "particularly robust signal was observed in the metastatic subset for the TP53+KRAS+SMAD4 triple context",
        " (see Supplementary Table 1 for the compact comparison).",
        None,
    ),
]


def set_para_text(p, text: str) -> None:
    for r in p.runs:
        r.text = ""
    if p.runs:
        p.runs[0].text = text
    else:
        p.add_run(text)


def strip_supp(text: str) -> str:
    out = text
    for _ in range(5):
        prev = out
        for pat in STRIP_PATTERNS:
            out = re.sub(pat, "", out, flags=re.IGNORECASE)
        if out == prev:
            break
    return re.sub(r"\s{2,}", " ", out).strip()


def find_supp_idx(doc: Document) -> int:
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip() == "Supplementary Materials" and i > 50:
            return i
    raise RuntimeError("Supplementary Materials not found")


def ensure_insert(text: str, needle: str, frag: str) -> str:
    if not frag.strip() or needle not in text:
        return text
    key = re.sub(r"[^\w]", "", frag)[:35]
    if key and key in re.sub(r"[^\w]", "", text):
        return text
    pos = text.find(needle)
    m = re.search(
        r"\s*(\[(?:\d+[–-])?\d+[^\]]*\]|Supplementary Method|\(Supplementary)",
        text[pos + len(needle) :],
    )
    if m:
        ins = pos + len(needle) + m.start()
        return text[:ins].rstrip() + frag + text[ins:]
    return text.rstrip() + frag


def expand_ranges(nums: list[int], text: str, kind: str) -> list[int]:
    """Record first appearance of each supplementary number in a paragraph."""
    prefix = rf"Supplementary {kind}s? "
    for m in re.finditer(prefix + r"([^);.]+)", text):
        chunk = m.group(1)
        for part in re.split(r",| and ", chunk):
            part = part.strip()
            for rng in re.findall(r"(\d+)[–-](\d+)", part):
                lo, hi = int(rng[0]), int(rng[1])
                for n in range(lo, hi + 1):
                    if n not in nums:
                        nums.append(n)
            for solo in re.findall(r"\b(\d+)\b", part):
                n = int(solo)
                if n not in nums:
                    nums.append(n)
    return nums


def first_appearance_order(doc: Document, start: int, end: int, kind: str) -> list[int]:
    seen: list[int] = []
    for i in range(start, end):
        expand_ranges(seen, doc.paragraphs[i].text, kind)
    return seen


def monotonic(nums: list[int]) -> bool:
    return nums == sorted(nums)


def main() -> None:
    doc = Document(str(MANUSCRIPT))
    end = find_supp_idx(doc)

    for i in range(end):
        cleaned = strip_supp(doc.paragraphs[i].text)
        set_para_text(doc.paragraphs[i], cleaned)

    for needle, frag, min_para in CITATIONS:
        if not frag.strip():
            continue
        for i in range(BODY_START_PARA, end):
            if min_para is not None and i < min_para:
                continue
            p = doc.paragraphs[i]
            new = ensure_insert(p.text, needle, frag)
            if new != p.text:
                set_para_text(p, new)
                break

    for i in range(end):
        t = doc.paragraphs[i].text
        cleaned = re.sub(r"\.{2,}", ".", t)
        cleaned = re.sub(r"\s+\.", ".", cleaned)
        cleaned = re.sub(r":\s*\.", ".", cleaned)
        if cleaned != t:
            set_para_text(doc.paragraphs[i], cleaned)

    doc.save(str(MANUSCRIPT))
    doc = Document(str(MANUSCRIPT))
    end = find_supp_idx(doc)

    m_order = first_appearance_order(doc, BODY_START_PARA, end, "Method")
    f_order = first_appearance_order(doc, BODY_START_PARA, end, "Figure")
    t_order = first_appearance_order(doc, BODY_START_PARA, end, "Table")

    print("Saved:", MANUSCRIPT)
    print("Methods first appearance:", m_order)
    print("Figures first appearance:", f_order)
    print("Tables first appearance:", t_order)
    print("Methods monotonic:", monotonic(m_order))
    print("Figures monotonic:", monotonic(f_order))
    print("Tables monotonic:", monotonic(t_order))


if __name__ == "__main__":
    main()
