#!/usr/bin/env python3
"""
Nat Commun manuscript boost (NEW text highlighted yellow):
  1. Shortened structured abstract: KM log-rank and key findings in plain wording (no supplementary calls, no references).
  2. Writes Figure_Decision_Schematic_ThreeLayers.png (embed via move_nat_commun_figures_to_supplementary.py)
  3. Yellow emphasis on four cBioPortal cohorts (no duplicate Methods block)

Figures are not placed in main text: run nat_commun_boost_analyses.py and move_nat_commun_figures_to_supplementary.py
"""

from __future__ import annotations

import re
from pathlib import Path

import matplotlib.pyplot as plt
from matplotlib.patches import FancyArrowPatch, FancyBboxPatch
from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph

TUMOR = Path(__file__).resolve().parent
MANUSCRIPT = TUMOR / "NewManuscript.docx"
MARKER = "NAT_COMMUN_BOOST_V1"

SCHEMATIC_PNG = TUMOR / "Figure_Decision_Schematic_ThreeLayers.png"


SHORT_ABSTRACT = (
    "Background: Co-mutation patterns in pancreatic ductal adenocarcinoma (PDAC) may carry "
    "different clinical meaning depending on disease extent and whether associations are assessed "
    "as cross-sectional enrichment or stage-conditioned survival. Methods: We analyzed one integrated "
    "patient-level cohort (N=2,330) with binary calls across ten recurrently altered genes in three clinically "
    "prespecified disease-extent categories, jointly modeling deceased-versus-living contrasts with multiplicity "
    "control, stage-stratified Cox models for overall survival, and interactions between genotype indicators and "
    "diagnosis-to-biosample latency (years elapsed from diagnosis to mutation profiling). Results: TP53+KRAS "
    "remained the dominant backbone genotype (~72%). Kaplan–Meier separation between TP53+KRAS carriers and all "
    "other patients yielded two-sample log-rank p-values of approximately 2×10⁻⁷ across the combined stage strata "
    "and approximately 3×10⁻⁴ within metastatic disease alone, aligning descriptively with univariate Cox patterns in "
    "the same analytic file. Ranking of enrichment-based dual combinations diverged from stage-specific hazard summaries, "
    "emphasizing that prevalence-driven screens and proportional hazards formulations answer distinct questions; "
    "timing-informed models highlighted comparatively stronger metastatic interactions for selected higher-order genotype "
    "contexts. External consistency checks spanning TCGA-harmonized, proteomic oncology, and large clinicogenomic pancreatic "
    "cohorts identified directionally adverse univariate mortality associations for simplified binary TP53+KRAS contrasts, "
    "with prevalence and hazards numerically nearest to our integrated cohort in the largest tertiary sequencing atlas. "
    "Conclusions: Linking multiplicity-aware enrichment summaries to stage-conditioned hazards plus explicit ascertainment-timing "
    "sensitivity affords a graded clinical-genomic storyline beyond isolated mutation labels. Associations are correlational "
    "and clinically non-deterministic."
)

EMPHASIS_EXTERNAL = (
    "Editorial emphasis (four public cohorts). The external module should be read as a harmonized "
    "replication family rather than four independent discoveries: paad_tcga_gdc and paad_tcga_pan_can_atlas_2018 "
    "overlap as TCGA PAAD harmonizations; pancreas_cptac_gdc contributes proteogenomic PDAC with intermediate "
    "co-mutation prevalence; pdac_msk_2024 is the largest series and anchors quantitative agreement for "
    "TP53+KRAS (prevalence and univariate hazard). We therefore highlight MSK and the TCGA pair for backbone "
    "replication, CPTAC for directional consistency, and the extended five dual-pair × four-cohort Cox matrix "
    "(Supplementary Table 2; Validation_cBioPortal_FivePairs_FourStudies.xlsx) for prespecified dual signals "
    "beyond TP53+KRAS—without re-deriving synergy scores or stage-stratified models in external data."
)


def insert_paragraph_after(paragraph: Paragraph, *, style: str | None = None) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._element.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style:
        new_para.style = style
    return new_para


def add_yellow_run(paragraph: Paragraph, text: str) -> None:
    r = paragraph.add_run(text)
    r.font.highlight_color = WD_COLOR_INDEX.YELLOW


def set_para_yellow(paragraph: Paragraph, text: str) -> None:
    el = paragraph._element
    for child in list(el):
        if child.tag.endswith("}r"):
            el.remove(child)
    add_yellow_run(paragraph, text)


def find_para(doc: Document, pattern: str, start: int = 0) -> int | None:
    rx = re.compile(pattern, re.I)
    for i in range(start, len(doc.paragraphs)):
        if rx.search(doc.paragraphs[i].text):
            return i
    return None


def para_has_image(p: Paragraph) -> bool:
    return "pic:pic" in p._element.xml or "w:drawing" in p._element.xml


def already_applied(doc: Document) -> bool:
    for p in doc.paragraphs:
        if MARKER in p.text:
            return True
        # Kept editorial block from earlier runs
        if p.text.startswith("Editorial emphasis (four public cohorts)") and "five-pair" in p.text:
            return True
    return False


def make_schematic(path: Path) -> None:
    fig, ax = plt.subplots(figsize=(8.5, 4.2))
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 5)
    ax.axis("off")

    boxes = [
        (0.4, 2.8, 2.8, 1.5, "Layer 1\nEnrichment screen", "#E8F4FC"),
        (3.6, 2.8, 2.8, 1.5, "Layer 2\nStage-stratified Cox", "#FFF4E6"),
        (6.8, 2.8, 2.8, 1.5, "Layer 3\nTiming interaction\n(DX2COLLECTION)", "#E8F8E8"),
    ]
    for x, y, w, h, label, color in boxes:
        patch = FancyBboxPatch(
            (x, y),
            w,
            h,
            boxstyle="round,pad=0.05,rounding_size=0.15",
            linewidth=1.2,
            edgecolor="#333333",
            facecolor=color,
        )
        ax.add_patch(patch)
        ax.text(x + w / 2, y + h / 2, label, ha="center", va="center", fontsize=10, fontweight="bold")

    for x0, x1 in [(3.2, 3.6), (6.4, 6.8)]:
        ax.add_patch(
            FancyArrowPatch(
                (x0, 3.55),
                (x1, 3.55),
                arrowstyle="-|>",
                mutation_scale=14,
                linewidth=1.5,
                color="#444444",
            )
        )

    ax.text(
        5,
        1.35,
        "Distinct questions: prevalence vs hazard vs timing sensitivity\n"
        "External cBioPortal cohorts replicate univariate OS for prespecified pairs (not timing)",
        ha="center",
        va="center",
        fontsize=9,
        style="italic",
        color="#333333",
    )
    ax.text(5, 4.65, "Three-layer analytical decision schematic (discovery cohort)", ha="center", fontsize=11, fontweight="bold")
    fig.tight_layout()
    fig.savefig(path, dpi=200, bbox_inches="tight")
    plt.close(fig)


def main() -> int:
    if not MANUSCRIPT.is_file():
        print(f"Missing {MANUSCRIPT}")
        return 1

    doc = Document(str(MANUSCRIPT))
    if already_applied(doc):
        print("Manuscript already contains Nat Commun boost markers; skipping.")
        return 0

    # --- 1. Short abstract (yellow replacement) ---
    i_abs = find_para(doc, r"^Background:.*PDAC")
    if i_abs is None:
        i_abs = find_para(doc, r"^Background:")
    if i_abs is None:
        print("Could not find abstract body paragraph.", file=__import__("sys").stderr)
        return 1
    set_para_yellow(doc.paragraphs[i_abs], SHORT_ABSTRACT)

    # Schematic PNG for Supplementary Figure 13 (embed with move_nat_commun_figures_to_supplementary.py)
    make_schematic(SCHEMATIC_PNG)

    # --- 2. Yellow emphasis after existing external replication (no duplicate Methods) ---
    i_ext = find_para(doc, r"Across paad_tcga_gdc")
    if i_ext is None:
        print("Could not find external replication paragraph.", file=__import__("sys").stderr)
        return 1
    emph = insert_paragraph_after(doc.paragraphs[i_ext])
    add_yellow_run(emph, EMPHASIS_EXTERNAL)

    doc.save(str(MANUSCRIPT))
    print(f"Updated {MANUSCRIPT} (yellow abstract + external emphasis + schematic PNG for Supp Fig 13)")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
