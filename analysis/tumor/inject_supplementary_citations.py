#!/usr/bin/env python3
"""
Append supplementary cross-references to NewManuscript.docx only.
Does not modify existing paragraph text except by appending citation runs.
"""

from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document


# paragraph_index -> list of citation fragments to append (if not already present)
APPENDS: dict[int, list[str]] = {
    # Abstract block (single paragraph) — external validation already may cite S2
    10: [],  # handled separately: ensure S2 + Methods pointer if missing
    30: [" (Supplementary Methods §6; Supplementary Figure S9; Supplementary Tables S9)."],
    31: [" (Supplementary Methods §6)."],
    33: [" (Supplementary Methods §7; Supplementary Table S2)."],
    35: [" (Supplementary Methods §8; Supplementary Figure S10)."],
    37: [" (Supplementary Methods §9)."],
    43: [
        " (Supplementary Methods §5; Supplementary Table S1; Supplementary Tables S5; "
        "Supplementary Figures S3–S4 and S7)."
    ],
    49: [" (Supplementary Figure S8; Supplementary Tables S10)."],
    52: [
        " (Supplementary Tables S3, S7, and S11; Supplementary Methods §2; "
        "see Supplementary Tables.xlsx, Table_Index)."
    ],
    56: [" (archived overview also in Supplementary Figure S11)."],
    63: [
        " (Supplementary Tables S10–S11; Supplementary Figure S11; Supplementary Methods §3)."
    ],
    67: [" (Supplementary Table S2; Supplementary Methods §7)."],
    70: [" (Supplementary Tables S3 and S11)."],
    73: [" (Supplementary Tables S10)."],
    74: [" (Supplementary Tables S10)."],
    77: [" (Supplementary Tables S7–S8)."],
    81: [" (Supplementary Tables S3 and S8)."],
    82: [" (Supplementary Tables S3 and S8)."],
    85: [" (Supplementary Tables S6; Supplementary Figure S12)."],
    100: [
        " (Supplementary Methods §5; Supplementary Tables S1 and S5; "
        "Supplementary Figures S3–S4 and S7)."
    ],
    102: [" (Supplementary Table S1)."],
    104: [
        " (triple-level panels: Supplementary Figures S1–S2; "
        "timing volcanoes: Supplementary Figures S3–S4 and S7)."
    ],
    108: [" (Supplementary Figure S9; Supplementary Tables S9)."],
    110: [" (Supplementary Figure S10)."],
    115: [" (Supplementary Table S2; Supplementary Methods §7)."],
    116: [" (Supplementary Table S2)."],
    # Supplementary Information header — file bundle pointer
    156: [],
    157: [
        " Full supplementary files: Supplementary Methods.docx, Supplementary Figures.docx, "
        "and Supplementary Tables.xlsx (Tumor/supplementary/)."
    ],
    160: [
        " See Supplementary Methods.docx for the complete protocol (Sections 1–9)."
    ],
    161: [
        " See Supplementary Tables.xlsx (Table_Index) and Supplementary Figures.docx for all numbered items."
    ],
    163: [" (Supplementary Tables.xlsx, sheet S1_Timing_Interaction)."],
    164: [" (Supplementary Tables.xlsx, sheet S2_External_cBioPortal)."],
    167: [" (Supplementary Figures.docx)."],
    171: [" (Supplementary Figures.docx)."],
    175: [" (Supplementary Figures.docx)."],
    179: [" (Supplementary Figures.docx)."],
    183: [" (Supplementary Figures.docx)."],
    187: [" (Supplementary Figures.docx)."],
    191: [" (Supplementary Figures.docx)."],
}

ABSTRACT_EXTRA = (
    " External validation details are in Supplementary Methods §7 and Supplementary Table S2 "
    "(Supplementary Tables.xlsx)."
)


def _already_has(text: str, fragment: str) -> bool:
    # strip leading space/paren for key phrase check
    key = fragment.strip().lstrip("(").split(";")[0].split(",")[0][:40]
    return key in text


def append_citations(doc: Document) -> int:
    n = 0
    for idx, fragments in APPENDS.items():
        if idx >= len(doc.paragraphs):
            continue
        p = doc.paragraphs[idx]
        full = p.text
        for frag in fragments:
            if not frag or _already_has(full, frag):
                continue
            run = p.add_run(frag)
            full += frag
            n += 1
    return n


def patch_abstract(doc: Document) -> None:
    p = doc.paragraphs[10]
    t = p.text
    if "Supplementary Methods §7" in t and "Supplementary Table S2" in t:
        return
    if "Supplementary Table S2" in t and "Supplementary Methods" not in t:
        p.add_run(
            " External validation protocol and numeric results: Supplementary Methods §7; "
            "Supplementary Table S2 (Supplementary Tables.xlsx)."
        )
    elif "Supplementary Table S2" not in t:
        p.add_run(ABSTRACT_EXTRA)


def patch_table1_s1(doc: Document) -> None:
    """Ensure Table 1 caption cites Supplementary Table S1 explicitly."""
    for i in (53, 54):
        if i >= len(doc.paragraphs):
            continue
        p = doc.paragraphs[i]
        if "Supplementary Table S1" in p.text:
            continue
        if "Supplementary Table" in p.text:
            # only append S1 if generic mention without S1
            if not re.search(r"Supplementary Table S\d", p.text):
                p.add_run(" (numbered as Supplementary Table S1 in Supplementary Tables.xlsx).")


def main() -> int:
    path = Path(__file__).resolve().parent / "NewManuscript.docx"
    doc = Document(str(path))
    patch_abstract(doc)
    patch_table1_s1(doc)
    added = append_citations(doc)
    doc.save(str(path))
    print(f"Appended {added} citation fragment(s) to {path}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
